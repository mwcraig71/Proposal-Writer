import io
import os
from typing import Optional
from docx import Document
from docx.oxml.ns import qn
from pypdf import PdfReader
from openpyxl import load_workbook
import xml.etree.ElementTree as ET


def extract_text_from_pdf(file_content: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(file_content))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract ALL text from Word documents including text boxes, shapes, and tables.
    Uses XML parsing to ensure no content is missed.
    """
    try:
        doc = Document(io.BytesIO(file_content))
        all_text_parts = []
        
        # Method 1: Extract from document body XML directly to get ALL content
        # This captures text boxes, shapes, and other elements that doc.paragraphs misses
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
            'mc': 'http://schemas.openxmlformats.org/markup-compatibility/2006',
            'v': 'urn:schemas-microsoft-com:vml',
        }
        
        def extract_all_text_from_element(element):
            """Recursively extract all text from an XML element"""
            texts = []
            # Get text from 'w:t' elements (main text)
            for t_elem in element.iter(qn('w:t')):
                if t_elem.text:
                    texts.append(t_elem.text)
            return ''.join(texts)
        
        def process_element_recursive(element, depth=0):
            """Process element and all children recursively"""
            text_items = []
            
            # Check if this is a paragraph
            if element.tag == qn('w:p'):
                para_text = extract_all_text_from_element(element)
                if para_text.strip():
                    text_items.append(para_text.strip())
            
            # Check for text in drawing/shape elements
            elif 'drawing' in element.tag or 'txbxContent' in element.tag:
                for child in element:
                    text_items.extend(process_element_recursive(child, depth + 1))
            
            # Process children
            for child in element:
                text_items.extend(process_element_recursive(child, depth + 1))
            
            return text_items
        
        # Extract from main document body using recursive XML parsing
        body = doc.element.body
        xml_extracted = process_element_recursive(body)
        
        # Method 2: Standard paragraph extraction (backup)
        para_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                para_text.append(para.text.strip())
        
        # Method 3: Table extraction
        table_text = []
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    table_text.append(" | ".join(row_texts))
        
        # Method 4: Try to extract from textboxes in drawing elements
        textbox_text = []
        try:
            for element in body.iter():
                # Look for txbxContent which contains text box content
                if 'txbxContent' in element.tag:
                    for p in element.iter(qn('w:p')):
                        p_text = extract_all_text_from_element(p)
                        if p_text.strip():
                            textbox_text.append(p_text.strip())
                
                # Also check for text in VML shapes (legacy Word format)
                if element.tag and 'textbox' in element.tag.lower():
                    for p in element.iter(qn('w:p')):
                        p_text = extract_all_text_from_element(p)
                        if p_text.strip():
                            textbox_text.append(p_text.strip())
        except Exception:
            pass
        
        # Combine all extracted text, preferring the method that got the most content
        all_sources = [
            ('xml_recursive', xml_extracted),
            ('paragraphs', para_text),
            ('textboxes', textbox_text),
        ]
        
        # Use the source with the most content
        best_source = max(all_sources, key=lambda x: len('\n'.join(x[1])))
        all_text_parts = best_source[1]
        
        # Always add table text at the end
        if table_text:
            all_text_parts.append("\n=== TABLE CONTENT ===")
            all_text_parts.extend(table_text)
        
        # If we got very little text, try raw XML extraction as last resort
        if len('\n'.join(all_text_parts)) < 500:
            # Get raw text from all 'w:t' elements in document
            raw_texts = []
            for t_elem in body.iter(qn('w:t')):
                if t_elem.text:
                    raw_texts.append(t_elem.text)
            
            # Group consecutive texts into lines
            if raw_texts:
                all_text_parts = []
                current_line = []
                for text in raw_texts:
                    current_line.append(text)
                    # Check if this looks like end of a line/paragraph
                    if text.endswith('.') or text.endswith(':') or text.endswith('\n'):
                        all_text_parts.append(''.join(current_line))
                        current_line = []
                if current_line:
                    all_text_parts.append(''.join(current_line))
        
        result = "\n".join(all_text_parts)
        
        # Debug: if still short, log what we found
        if len(result) < 500:
            print(f"WARNING: Short extraction ({len(result)} chars). Methods tried:")
            for name, texts in all_sources:
                print(f"  {name}: {len(texts)} items, {len(''.join(texts))} chars")
            print(f"  tables: {len(table_text)} items")
        
        return result
        
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")


def extract_text_from_excel(file_content: bytes) -> str:
    try:
        wb = load_workbook(io.BytesIO(file_content), data_only=True)
        text_parts = []
        
        for sheet_name in wb.sheetnames:
            sheet = wb[sheet_name]
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            
            for row in sheet.iter_rows(values_only=True):
                row_values = [str(cell) if cell is not None else "" for cell in row]
                if any(v.strip() for v in row_values):
                    text_parts.append(" | ".join(row_values))
        
        return "\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to parse Excel: {str(e)}")


def extract_text_from_file(filename: str, file_content: bytes) -> str:
    ext = os.path.splitext(filename)[1].lower()
    
    if ext == '.pdf':
        return extract_text_from_pdf(file_content)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_content)
    elif ext in ['.xlsx', '.xls']:
        return extract_text_from_excel(file_content)
    elif ext == '.txt':
        return file_content.decode('utf-8', errors='ignore')
    else:
        raise ValueError(f"Unsupported file format: {ext}")
