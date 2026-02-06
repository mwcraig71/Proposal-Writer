import os
import json
from functools import wraps
from flask import render_template, request, jsonify, redirect, url_for, flash, send_file, send_from_directory, session
from werkzeug.utils import secure_filename
from main import app
from database import db
from models import (
    Firm, Employee, Project, EmployeeProjectLink, Proposal,
    ProposalSelectedEmployee, ProposalSelectedProject, ProposalEmployeeRelevantProject,
    ProjectFirmInvolvement, EmployeeProjectExperience, ProjectAlternateDescription, AISettings,
    ClientContact, ExperienceAlternateDescription, Certification, CertificationType,
    EmployeePhoto, ProjectPhoto, ProposalReference, ProposalIntelligence,
    FirmPhoto, ProposalSelectedFirmPhoto, MarketingPhoto, ProposalSelectedMarketingPhoto,
    EmployeeAlternateBio, FirmAlternateDescription, ProposalSavedResponse,
    Response, ProposalLinkedResponse, Reference, ProposalLinkedReference
)
from replit.object_storage import Client as ObjectStorageClient
import uuid
from document_parser import extract_text_from_file
from gemini_service import detect_document_type, parse_employee_resume, parse_project_sheet, parse_multiple_projects, parse_firm_info, find_matching_employee, combine_and_rewrite_text
from pdf_generator import generate_full_sf330, get_form_fields
import io

UPLOAD_FOLDER = 'uploads'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc', 'xlsx', 'xls', 'txt'}

# Simple password protection
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        app_password = os.environ.get('APP_PASSWORD')
        if app_password and not session.get('authenticated'):
            # Check if this is an API/AJAX request
            if request.path.startswith('/api/') or request.headers.get('X-Requested-With') == 'XMLHttpRequest' or request.is_json:
                return jsonify({'success': False, 'error': 'Authentication required. Please refresh the page and log in.'}), 401
            return redirect(url_for('login', next=request.url))
        return f(*args, **kwargs)
    return decorated_function


@app.route('/login', methods=['GET', 'POST'])
def login():
    app_password = os.environ.get('APP_PASSWORD')
    if not app_password:
        session['authenticated'] = True
        return redirect(url_for('index'))
    
    if request.method == 'POST':
        password = request.form.get('password', '')
        if password == app_password:
            session['authenticated'] = True
            next_url = request.args.get('next', url_for('index'))
            return redirect(next_url)
        else:
            flash('Incorrect password', 'error')
    
    return render_template('login.html')


@app.route('/logout')
def logout():
    session.pop('authenticated', None)
    flash('You have been logged out', 'success')
    return redirect(url_for('login'))

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route('/')
@login_required
def index():
    from datetime import date, timedelta
    from collections import defaultdict
    
    proposals = Proposal.query.order_by(Proposal.updated_at.desc()).limit(5).all()
    employees_count = Employee.query.count()
    projects_count = Project.query.count()
    firms_count = Firm.query.count()
    
    # Get expired and expiring certifications in next 6 months
    today = date.today()
    six_months_from_now = today + timedelta(days=180)
    
    expired_certs = Certification.query.filter(
        Certification.expiration_date < today
    ).order_by(Certification.expiration_date.desc()).all()
    
    expiring_certs = Certification.query.filter(
        Certification.expiration_date >= today,
        Certification.expiration_date <= six_months_from_now
    ).order_by(Certification.expiration_date.asc()).all()
    
    # Group by category
    expired_by_category = defaultdict(list)
    for cert in expired_certs:
        expired_by_category[cert.category or 'Other'].append(cert)
    
    expiring_by_category = defaultdict(list)
    for cert in expiring_certs:
        expiring_by_category[cert.category or 'Other'].append(cert)
    
    return render_template('index.html', 
                         proposals=proposals,
                         employees_count=employees_count,
                         projects_count=projects_count,
                         firms_count=firms_count,
                         expired_certs_by_category=dict(expired_by_category),
                         expiring_certs_by_category=dict(expiring_by_category),
                         expired_total=len(expired_certs),
                         expiring_total=len(expiring_certs))


@app.route('/upload', methods=['GET', 'POST'])
def upload():
    if request.method == 'GET':
        return render_template('upload.html')
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'File type not allowed. Supported: PDF, DOCX, XLSX, TXT'}), 400
    
    try:
        file_content = file.read()
        text = extract_text_from_file(file.filename, file_content)
        
        # Debug: Log extracted text for troubleshooting resume parsing
        print(f"DEBUG EXTRACTED TEXT (first 3000 chars):\n{text[:3000]}")
        print(f"DEBUG EXTRACTED TEXT LENGTH: {len(text)} characters")
        
        doc_type = detect_document_type(text)
        
        parsed_data = {}
        if doc_type == 'employee':
            parsed_data = parse_employee_resume(text)
        elif doc_type == 'project':
            parsed_data = parse_project_sheet(text)
        elif doc_type == 'projects':
            parsed_data = parse_multiple_projects(text)
        elif doc_type == 'firm':
            parsed_data = parse_firm_info(text)
        else:
            return jsonify({
                'doc_type': 'unknown',
                'raw_text': text[:5000],
                'message': 'Could not determine document type. Please specify manually.'
            })
        
        return jsonify({
            'doc_type': doc_type,
            'parsed_data': parsed_data,
            'raw_text': text[:2000]
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/upload-multi-projects', methods=['POST'])
def upload_multi_projects():
    """Force parse a document as multiple projects, bypassing auto-detection."""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'File type not allowed. Supported: PDF, DOCX, XLSX, TXT'}), 400
    
    try:
        file_content = file.read()
        text = extract_text_from_file(file.filename, file_content)
        
        parsed_data = parse_multiple_projects(text)
        
        return jsonify({
            'doc_type': 'projects',
            'parsed_data': parsed_data,
            'raw_text': text[:2000]
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


def safe_int(value):
    """Convert value to integer safely, extracting numbers from strings like '15 Years'"""
    if value is None:
        return None
    if isinstance(value, int):
        return value
    if isinstance(value, str):
        import re
        numbers = re.findall(r'\d+', value)
        if numbers:
            return int(numbers[0])
    try:
        return int(value)
    except (ValueError, TypeError):
        return None

@app.route('/save-parsed-data', methods=['POST'])
def save_parsed_data():
    data = request.json
    doc_type = data.get('doc_type')
    parsed_data = data.get('parsed_data')
    force_new = data.get('force_new', False)
    
    try:
        if doc_type == 'employee':
            name = parsed_data.get('name', 'Unknown')
            
            if not force_new:
                existing_id = find_matching_employee(name)
                if existing_id:
                    session['pending_employee_data'] = parsed_data
                    existing_emp = Employee.query.get(existing_id)
                    return jsonify({
                        'success': True,
                        'duplicate_found': True,
                        'existing_id': existing_id,
                        'existing_name': existing_emp.name if existing_emp else name,
                        'message': f'Found existing employee "{name}". You can merge with their data or save as a new person.'
                    })
            
            employee = Employee(
                name=name,
                title=parsed_data.get('title'),
                role=parsed_data.get('role'),
                years_experience_total=safe_int(parsed_data.get('years_experience_total')),
                years_experience_firm=safe_int(parsed_data.get('years_experience_firm')),
                education=parsed_data.get('education'),
                registrations=parsed_data.get('registrations'),
                training=parsed_data.get('training'),
                other_qualifications=parsed_data.get('other_qualifications')
            )
            db.session.add(employee)
            db.session.commit()
            
            bio_text = parsed_data.get('bio')
            if bio_text:
                from datetime import datetime
                alt_bio = EmployeeAlternateBio(
                    employee_id=employee.id,
                    label=f"Imported {datetime.now().strftime('%Y-%m-%d')}",
                    bio=bio_text
                )
                db.session.add(alt_bio)
                db.session.commit()
            
            project_experiences = parsed_data.get('project_experience', [])
            print(f"DEBUG: Found {len(project_experiences) if project_experiences else 0} project experiences for {name}")
            print(f"DEBUG: project_experience data: {project_experiences}")
            if project_experiences:
                for proj in project_experiences:
                    if proj.get('project_title'):
                        exp = EmployeeProjectExperience(
                            employee_id=employee.id,
                            project_title=proj.get('project_title'),
                            location=proj.get('location'),
                            owner_name=proj.get('owner_name'),
                            project_cost=proj.get('project_cost'),
                            year_completed=proj.get('year_completed'),
                            role_performed=proj.get('role_performed'),
                            brief_description=proj.get('brief_description'),
                            firm_name=proj.get('firm_name'),
                            is_current_firm=False
                        )
                        db.session.add(exp)
                db.session.commit()
            
            return jsonify({'success': True, 'id': employee.id, 'message': 'Employee saved successfully'})
        
        elif doc_type == 'project':
            project = Project(
                title=parsed_data.get('title', 'Untitled Project'),
                location=parsed_data.get('location'),
                year_completed_professional=parsed_data.get('year_completed_professional'),
                year_completed_construction=parsed_data.get('year_completed_construction'),
                owner_name=parsed_data.get('owner_name'),
                owner_contact_name=parsed_data.get('owner_contact_name'),
                owner_contact_phone=parsed_data.get('owner_contact_phone'),
                project_cost=parsed_data.get('project_cost'),
                project_delivery_method=parsed_data.get('project_delivery_method'),
                brief_description=parsed_data.get('brief_description'),
                relevance_writeup=parsed_data.get('relevance_writeup')
            )
            db.session.add(project)
            db.session.commit()
            
            team_members = parsed_data.get('team_members', [])
            for member in team_members:
                employee = Employee.query.filter_by(name=member.get('name')).first()
                if employee:
                    link = EmployeeProjectLink(
                        employee_id=employee.id,
                        project_id=project.id,
                        role_on_project=member.get('role_on_project')
                    )
                    db.session.add(link)
            
            db.session.commit()
            return jsonify({'success': True, 'id': project.id, 'message': 'Project saved successfully'})
        
        elif doc_type == 'firm':
            firm = Firm(
                name=parsed_data.get('name', 'Unknown Firm'),
                uei=parsed_data.get('uei'),
                street_address=parsed_data.get('street_address'),
                city=parsed_data.get('city'),
                state=parsed_data.get('state'),
                zip_code=parsed_data.get('zip_code'),
                country=parsed_data.get('country', 'USA'),
                year_established=parsed_data.get('year_established'),
                ownership_type=parsed_data.get('ownership_type'),
                is_small_business=parsed_data.get('is_small_business', False),
                small_business_categories=parsed_data.get('small_business_categories'),
                phone=parsed_data.get('phone'),
                fax=parsed_data.get('fax'),
                email=parsed_data.get('email'),
                point_of_contact_name=parsed_data.get('point_of_contact_name'),
                point_of_contact_title=parsed_data.get('point_of_contact_title')
            )
            db.session.add(firm)
            db.session.commit()
            return jsonify({'success': True, 'id': firm.id, 'message': 'Firm saved successfully'})
        
        elif doc_type == 'projects':
            projects_list = parsed_data.get('projects', [])
            saved_ids = []
            for proj_data in projects_list:
                contact_info = []
                if proj_data.get('primary_contact_email'):
                    contact_info.append(f"Primary Email: {proj_data.get('primary_contact_email')}")
                if proj_data.get('alternate_contact_name'):
                    alt_contact = f"Alternate Contact: {proj_data.get('alternate_contact_name')}"
                    if proj_data.get('alternate_contact_phone'):
                        alt_contact += f", {proj_data.get('alternate_contact_phone')}"
                    if proj_data.get('alternate_contact_email'):
                        alt_contact += f", {proj_data.get('alternate_contact_email')}"
                    contact_info.append(alt_contact)
                if proj_data.get('contract_number'):
                    contact_info.insert(0, f"Contract #: {proj_data.get('contract_number')}")
                
                project = Project(
                    title=proj_data.get('title') or proj_data.get('contract_number') or 'Untitled Project',
                    location=proj_data.get('location'),
                    year_completed_professional=proj_data.get('year_completed'),
                    owner_name=proj_data.get('client'),
                    owner_contact_name=proj_data.get('primary_contact_name'),
                    owner_contact_phone=proj_data.get('primary_contact_phone'),
                    project_cost=proj_data.get('contract_value'),
                    project_delivery_method=proj_data.get('delivery_method'),
                    brief_description=proj_data.get('description'),
                    relevance_writeup='\n'.join(contact_info) if contact_info else None
                )
                db.session.add(project)
                db.session.flush()
                saved_ids.append(project.id)
            
            db.session.commit()
            return jsonify({
                'success': True, 
                'ids': saved_ids, 
                'count': len(saved_ids),
                'message': f'{len(saved_ids)} projects saved successfully'
            })
        
        else:
            return jsonify({'error': 'Invalid document type'}), 400
            
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@app.route('/employees')
def employees():
    employees = Employee.query.order_by(Employee.name).all()
    firms = Firm.query.order_by(Firm.name).all()
    
    # Group employees by firm
    employees_by_firm = {}
    unassigned_employees = []
    for employee in employees:
        if employee.firm_id:
            if employee.firm_id not in employees_by_firm:
                employees_by_firm[employee.firm_id] = []
            employees_by_firm[employee.firm_id].append(employee)
        else:
            unassigned_employees.append(employee)
    
    # Define colors for each firm (cycle through if more firms than colors)
    firm_colors = [
        {'bg': 'bg-blue-50', 'border': 'border-blue-200', 'tab': 'bg-blue-600', 'tab_inactive': 'bg-blue-100 text-blue-700', 'badge': 'bg-blue-100 text-blue-800'},
        {'bg': 'bg-green-50', 'border': 'border-green-200', 'tab': 'bg-green-600', 'tab_inactive': 'bg-green-100 text-green-700', 'badge': 'bg-green-100 text-green-800'},
        {'bg': 'bg-purple-50', 'border': 'border-purple-200', 'tab': 'bg-purple-600', 'tab_inactive': 'bg-purple-100 text-purple-700', 'badge': 'bg-purple-100 text-purple-800'},
        {'bg': 'bg-orange-50', 'border': 'border-orange-200', 'tab': 'bg-orange-600', 'tab_inactive': 'bg-orange-100 text-orange-700', 'badge': 'bg-orange-100 text-orange-800'},
        {'bg': 'bg-pink-50', 'border': 'border-pink-200', 'tab': 'bg-pink-600', 'tab_inactive': 'bg-pink-100 text-pink-700', 'badge': 'bg-pink-100 text-pink-800'},
        {'bg': 'bg-teal-50', 'border': 'border-teal-200', 'tab': 'bg-teal-600', 'tab_inactive': 'bg-teal-100 text-teal-700', 'badge': 'bg-teal-100 text-teal-800'},
    ]
    
    # Build firm data with colors
    firm_data = []
    for i, firm in enumerate(firms):
        color = firm_colors[i % len(firm_colors)]
        firm_data.append({
            'firm': firm,
            'employees': employees_by_firm.get(firm.id, []),
            'color': color
        })
    
    return render_template('employees.html', employees=employees, firms=firms, firm_data=firm_data, 
                          unassigned_employees=unassigned_employees, firm_colors=firm_colors)


@app.route('/employees/scrape-website', methods=['POST'])
def scrape_employees_from_website():
    from web_scraper import scrape_team_page
    from gemini_service import parse_team_personnel
    
    data = request.json
    url = data.get('url', '').strip()
    
    if not url:
        return jsonify({'success': False, 'error': 'URL is required'})
    
    try:
        scraped = scrape_team_page(url)
        if not scraped.get('content'):
            return jsonify({'success': False, 'error': 'Could not fetch content from website'})
        
        parsed = parse_team_personnel(scraped['content'])
        personnel = parsed.get('personnel', [])
        
        return jsonify({
            'success': True,
            'personnel': personnel,
            'profiles_found': len(scraped.get('profile_links', []))
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/employees/save-batch', methods=['POST'])
def save_employees_batch():
    data = request.json
    personnel = data.get('personnel', [])
    firm_id = data.get('firm_id')
    
    saved_count = 0
    for person in personnel:
        name = person.get('name', '').strip()
        if not name:
            continue
        
        name_parts = name.split()
        first_name = name_parts[0] if name_parts else None
        last_name = name_parts[-1] if len(name_parts) > 1 else None
        middle_name = ' '.join(name_parts[1:-1]) if len(name_parts) > 2 else None
        
        employee = Employee(
            name=name,
            first_name=first_name,
            middle_name=middle_name,
            last_name=last_name,
            title=person.get('title'),
            role=person.get('role'),
            years_experience_total=person.get('years_experience'),
            education=person.get('education'),
            registrations=person.get('registrations'),
            bio=person.get('bio'),
            firm_id=int(firm_id) if firm_id else None
        )
        db.session.add(employee)
        saved_count += 1
    
    db.session.commit()
    return jsonify({'success': True, 'saved': saved_count})


@app.route('/employees/add', methods=['GET', 'POST'])
def add_employee():
    if request.method == 'POST':
        data = request.form
        first_name = data.get('first_name', '').strip()
        middle_name = data.get('middle_name', '').strip()
        last_name = data.get('last_name', '').strip()
        nickname = data.get('nickname', '').strip()
        
        # Build full name from components
        name_parts = [first_name, middle_name, last_name]
        full_name = ' '.join(p for p in name_parts if p)
        
        employee = Employee(
            name=full_name,
            first_name=first_name if first_name else None,
            middle_name=middle_name if middle_name else None,
            last_name=last_name if last_name else None,
            nickname=nickname if nickname else None,
            title=data.get('title'),
            role=data.get('role'),
            years_experience_total=int(data.get('years_experience_total') or 0) if data.get('years_experience_total') else None,
            years_experience_firm=int(data.get('years_experience_firm') or 0) if data.get('years_experience_firm') else None,
            education=data.get('education'),
            registrations=data.get('registrations'),
            training=data.get('training'),
            other_qualifications=data.get('other_qualifications'),
            firm_id=int(data.get('firm_id')) if data.get('firm_id') else None
        )
        db.session.add(employee)
        db.session.flush()  # Get the employee ID
        
        # Process checked certifications from the checklist
        checked_cert_ids = data.getlist('cert_types')
        for cert_type_id in checked_cert_ids:
            cert_type = CertificationType.query.get(int(cert_type_id))
            if cert_type:
                cert = Certification(
                    employee_id=employee.id,
                    name=cert_type.name,
                    category=cert_type.category,
                    cert_type=cert_type.cert_type,
                    status='pending'  # Mark as pending until details are filled
                )
                db.session.add(cert)
        
        # Process SPRAT certification
        if data.get('has_sprat'):
            sprat_level = data.get('sprat_level')
            cert = Certification(
                employee_id=employee.id,
                name='SPRAT',
                category='SPRAT',
                cert_type='certification',
                level=sprat_level if sprat_level else None,
                status='active'
            )
            db.session.add(cert)
        
        # Process PE License states
        pe_states = data.getlist('pe_states')
        for state in pe_states:
            cert = Certification(
                employee_id=employee.id,
                name='PE',
                category='PE License',
                cert_type='license',
                state=state,
                status='pending'  # Details to be filled later
            )
            db.session.add(cert)
        
        db.session.commit()
        return redirect(f'/employees/{employee.id}')
    
    firms = Firm.query.all()
    # Get certification types grouped by category
    cert_types = CertificationType.query.order_by(CertificationType.category, CertificationType.sort_order, CertificationType.name).all()
    cert_types_grouped = {}
    for ct in cert_types:
        category = ct.category or 'Other'
        if category not in cert_types_grouped:
            cert_types_grouped[category] = []
        cert_types_grouped[category].append(ct)
    
    return render_template('employee_add.html', firms=firms, cert_types_grouped=cert_types_grouped)


@app.route('/employees/<int:id>')
def employee_detail(id):
    from models import MarketingPhoto
    employee = Employee.query.get_or_404(id)
    projects = Project.query.join(EmployeeProjectLink).filter(EmployeeProjectLink.employee_id == id).all()
    project_experiences = EmployeeProjectExperience.query.filter_by(employee_id=id).order_by(EmployeeProjectExperience.year_completed.desc()).all()
    firms = Firm.query.all()
    employees = [{'id': e.id, 'name': e.name} for e in Employee.query.order_by(Employee.name).all()]
    
    # Get all projects for linking
    all_projects = Project.query.order_by(Project.title).all()
    all_projects_json = [{'id': p.id, 'title': p.title} for p in all_projects]
    
    # Serialize project experiences for JavaScript
    project_experiences_json = [{
        'id': exp.id,
        'project_title': exp.project_title,
        'role_performed': exp.role_performed,
        'year_completed': exp.year_completed,
        'location': exp.location,
        'owner_name': exp.owner_name,
        'brief_description': exp.brief_description,
        'sf330_include': exp.sf330_include,
        'is_current_firm': exp.is_current_firm,
        'firm_name': exp.firm_name,
        'project_cost': exp.project_cost,
        'linked_project_id': exp.linked_project_id,
        'linked_project_title': exp.linked_project.title if exp.linked_project else None,
        'selected_alt_description_id': exp.selected_alt_description_id,
        'resume_order': exp.resume_order,
        'active_description_label': exp.active_description_label,
        'alternate_descriptions': [{'id': ad.id, 'label': ad.label} for ad in exp.alternate_descriptions]
    } for exp in project_experiences]
    
    # Find marketing photos tagged with this employee's name
    employee_tag = f"#{employee.name.replace(' ', '')}"
    all_marketing = MarketingPhoto.query.all()
    marketing_photos = [p for p in all_marketing if employee_tag.lower() in (p.tags or '').lower()]
    
    return render_template('employee_detail.html', employee=employee, projects=projects, project_experiences=project_experiences,
                           project_experiences_json=project_experiences_json, firms=firms, employees=employees, 
                           marketing_photos=marketing_photos, all_projects_json=all_projects_json)


@app.route('/employee/<int:id>/download')
def download_employee_resume(id):
    """Download employee resume using template with placeholder replacement"""
    employee = Employee.query.get_or_404(id)
    firm = Firm.query.get(employee.firm_id) if employee.firm_id else None
    
    doc = None
    try:
        client = get_storage_client()
        template_bytes = client.download_as_bytes('templates/resume_template_custom.docx')
        if template_bytes:
            from docx import Document
            doc = Document(io.BytesIO(template_bytes))
    except:
        pass
    
    if doc is None:
        doc = create_default_resume_template()
    
    project_exp_str = ''
    numbered_experiences = EmployeeProjectExperience.query.filter_by(employee_id=id).filter(
        EmployeeProjectExperience.resume_order.isnot(None)
    ).order_by(EmployeeProjectExperience.resume_order.asc()).all()
    individual_project_placeholders = {}
    if numbered_experiences:
        exp_lines = []
        for idx, exp in enumerate(numbered_experiences, 1):
            line = exp.project_title or 'Untitled'
            if exp.role_performed:
                line += f" - {exp.role_performed}"
            if exp.location:
                line += f" | {exp.location}"
            if exp.year_completed:
                line += f" ({exp.year_completed})"
            if exp.active_description:
                line += f"\n{exp.active_description}"
            exp_lines.append(line)
            individual_project_placeholders[f'{{{{PROJECT_EXPERIENCE_{idx}}}}}'] = line
            individual_project_placeholders[f'{{{{PROJECT_EXPERIENCE_{idx}_TITLE}}}}'] = exp.project_title or ''
            individual_project_placeholders[f'{{{{PROJECT_EXPERIENCE_{idx}_ROLE}}}}'] = exp.role_performed or ''
            individual_project_placeholders[f'{{{{PROJECT_EXPERIENCE_{idx}_LOCATION}}}}'] = exp.location or ''
            individual_project_placeholders[f'{{{{PROJECT_EXPERIENCE_{idx}_YEAR}}}}'] = exp.year_completed or ''
            individual_project_placeholders[f'{{{{PROJECT_EXPERIENCE_{idx}_DESCRIPTION}}}}'] = exp.active_description or ''
        project_exp_str = '\n\n'.join(exp_lines)
    
    employee_location_parts = [employee.city or '', employee.state or '']
    employee_location = ', '.join(p for p in employee_location_parts if p)
    
    placeholders = {
        '{{EMPLOYEE_NAME}}': employee.display_name or employee.name or '',
        '{{EMPLOYEE_FIRST_NAME}}': employee.first_name or '',
        '{{EMPLOYEE_MIDDLE_NAME}}': employee.middle_name or '',
        '{{EMPLOYEE_LAST_NAME}}': employee.last_name or '',
        '{{EMPLOYEE_TITLE}}': employee.title or '',
        '{{EMPLOYEE_ROLE}}': employee.role or '',
        '{{EMPLOYEE_CITY}}': employee.city or '',
        '{{EMPLOYEE_STATE}}': employee.state or '',
        '{{EMPLOYEE_LOCATION}}': employee_location,
        '{{FIRM_NAME}}': firm.name if firm else '',
        '{{FIRM_CITY}}': firm.city if firm else '',
        '{{FIRM_STATE}}': firm.state if firm else '',
        '{{YEARS_EXPERIENCE_TOTAL}}': str(employee.years_experience_total) if employee.years_experience_total else '',
        '{{YEARS_EXPERIENCE_FIRM}}': str(employee.years_experience_firm) if employee.years_experience_firm else '',
        '{{EDUCATION}}': employee.education or '',
        '{{REGISTRATIONS}}': employee.registrations or '',
        '{{BIO}}': employee.bio or '',
        '{{TRAINING}}': employee.training or '',
        '{{OTHER_QUALIFICATIONS}}': employee.other_qualifications or '',
        '{{PROJECT_EXPERIENCE}}': project_exp_str,
    }
    placeholders.update(individual_project_placeholders)
    
    def replace_resume_placeholders_in_para(para, placeholders):
        for placeholder, value in placeholders.items():
            if placeholder in para.text:
                replaced = False
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
                        replaced = True
                if not replaced and para.runs:
                    full_text = para.text
                    new_text = full_text.replace(placeholder, value)
                    if full_text != new_text:
                        for i, run in enumerate(para.runs):
                            if i == 0:
                                run.text = new_text
                            else:
                                run.text = ''
    
    for para in doc.paragraphs:
        replace_resume_placeholders_in_para(para, placeholders)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_resume_placeholders_in_para(para, placeholders)
    
    import re
    project_exp_pattern = re.compile(r'\{\{PROJECT_EXPERIENCE_\d+(_[A-Z]+)?\}\}')
    def collect_empty_project_paras(paragraphs):
        to_remove = []
        in_project_zone = False
        for para in paragraphs:
            text = para.text.strip()
            if project_exp_pattern.search(text):
                cleaned = project_exp_pattern.sub('', text).strip(' -()')
                if not cleaned:
                    to_remove.append(para)
                    in_project_zone = True
            elif not text and in_project_zone:
                to_remove.append(para)
            else:
                in_project_zone = False
        return to_remove
    
    paras_to_remove = collect_empty_project_paras(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras_to_remove.extend(collect_empty_project_paras(cell.paragraphs))
    for para in paras_to_remove:
        p_element = para._element
        p_element.getparent().remove(p_element)
    
    safe_name = employee.display_name or employee.name or 'Employee'
    safe_name = ''.join(c for c in safe_name if c.isalnum() or c in ' _-')[:50].strip()
    filename = f"Resume_{safe_name.replace(' ', '_')}.docx"
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/employee/<int:id>/download-sf330')
def download_employee_sf330_resume(id):
    """Download employee resume in SF330 Section E format"""
    employee = Employee.query.get_or_404(id)
    firm = Firm.query.get(employee.firm_id) if employee.firm_id else None
    
    doc = None
    try:
        client = get_storage_client()
        template_bytes = client.download_as_bytes('templates/sf330_resume_template_custom.docx')
        if template_bytes:
            from docx import Document
            doc = Document(io.BytesIO(template_bytes))
    except:
        pass
    
    if doc is None:
        try:
            from docx import Document
            default_path = os.path.join(os.path.dirname(__file__), 'attached_assets', 'sf330_section_e_template.docx')
            if os.path.exists(default_path):
                doc = Document(default_path)
            else:
                doc = Document(os.path.join(os.path.dirname(__file__), 'attached_assets', '330_Section_E_Standards_template_1770398969209.docx'))
        except Exception as e:
            flash(f'SF330 Section E template not found: {str(e)}', 'error')
            return redirect(url_for('employee_detail', id=id))
    
    numbered_experiences = EmployeeProjectExperience.query.filter_by(employee_id=id).filter(
        EmployeeProjectExperience.resume_order.isnot(None)
    ).order_by(EmployeeProjectExperience.resume_order.asc()).all()
    
    project_exp_str = ''
    individual_project_placeholders = {}
    if numbered_experiences:
        exp_lines = []
        for idx, exp in enumerate(numbered_experiences, 1):
            line = exp.project_title or 'Untitled'
            if exp.role_performed:
                line += f" - {exp.role_performed}"
            if exp.location:
                line += f" | {exp.location}"
            if exp.year_completed:
                line += f" ({exp.year_completed})"
            if exp.active_description:
                line += f"\n{exp.active_description}"
            exp_lines.append(line)
            individual_project_placeholders[f'{{{{PROJECT_EXPERIENCE_{idx}}}}}'] = line
            individual_project_placeholders[f'{{{{PROJECT_EXPERIENCE_{idx}_TITLE}}}}'] = exp.project_title or ''
            individual_project_placeholders[f'{{{{PROJECT_EXPERIENCE_{idx}_ROLE}}}}'] = exp.role_performed or ''
            individual_project_placeholders[f'{{{{PROJECT_EXPERIENCE_{idx}_LOCATION}}}}'] = exp.location or ''
            individual_project_placeholders[f'{{{{PROJECT_EXPERIENCE_{idx}_YEAR}}}}'] = exp.year_completed or ''
            individual_project_placeholders[f'{{{{PROJECT_EXPERIENCE_{idx}_DESCRIPTION}}}}'] = exp.active_description or ''
        project_exp_str = '\n\n'.join(exp_lines)
    
    employee_location_parts = [employee.city or '', employee.state or '']
    employee_location = ', '.join(p for p in employee_location_parts if p)
    
    placeholders = {
        '{{EMPLOYEE_NAME}}': employee.display_name or employee.name or '',
        '{{EMPLOYEE_FIRST_NAME}}': employee.first_name or '',
        '{{EMPLOYEE_MIDDLE_NAME}}': employee.middle_name or '',
        '{{EMPLOYEE_LAST_NAME}}': employee.last_name or '',
        '{{EMPLOYEE_TITLE}}': employee.title or '',
        '{{EMPLOYEE_ROLE}}': employee.role or '',
        '{{EMPLOYEE_CITY}}': employee.city or '',
        '{{EMPLOYEE_STATE}}': employee.state or '',
        '{{EMPLOYEE_LOCATION}}': employee_location,
        '{{FIRM_NAME}}': firm.name if firm else '',
        '{{FIRM_CITY}}': firm.city if firm else '',
        '{{FIRM_STATE}}': firm.state if firm else '',
        '{{YEARS_EXPERIENCE_TOTAL}}': str(employee.years_experience_total) if employee.years_experience_total else '',
        '{{YEARS_EXPERIENCE_FIRM}}': str(employee.years_experience_firm) if employee.years_experience_firm else '',
        '{{EDUCATION}}': employee.education or '',
        '{{REGISTRATIONS}}': employee.registrations or '',
        '{{BIO}}': employee.bio or '',
        '{{TRAINING}}': employee.training or '',
        '{{OTHER_QUALIFICATIONS}}': employee.other_qualifications or '',
        '{{PROJECT_EXPERIENCE}}': project_exp_str,
    }
    placeholders.update(individual_project_placeholders)
    
    def replace_sf330_placeholders_in_para(para, phs):
        for placeholder, value in phs.items():
            if placeholder in para.text:
                replaced = False
                for run in para.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, value)
                        replaced = True
                if not replaced and para.runs:
                    full_text = para.text
                    new_text = full_text.replace(placeholder, value)
                    if full_text != new_text:
                        for i, run in enumerate(para.runs):
                            if i == 0:
                                run.text = new_text
                            else:
                                run.text = ''
    
    for para in doc.paragraphs:
        replace_sf330_placeholders_in_para(para, placeholders)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_sf330_placeholders_in_para(para, placeholders)
    
    import re
    project_exp_pattern = re.compile(r'\{\{PROJECT_EXPERIENCE_\d+(_[A-Z]+)?\}\}')
    def collect_empty_sf330_paras(paragraphs):
        to_remove = []
        in_project_zone = False
        for para in paragraphs:
            text = para.text.strip()
            if project_exp_pattern.search(text):
                cleaned = project_exp_pattern.sub('', text).strip(' -()')
                if not cleaned:
                    to_remove.append(para)
                    in_project_zone = True
            elif not text and in_project_zone:
                to_remove.append(para)
            else:
                in_project_zone = False
        return to_remove
    
    paras_to_remove = collect_empty_sf330_paras(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paras_to_remove.extend(collect_empty_sf330_paras(cell.paragraphs))
    for para in paras_to_remove:
        p_element = para._element
        p_element.getparent().remove(p_element)
    
    safe_name = employee.display_name or employee.name or 'Employee'
    safe_name = ''.join(c for c in safe_name if c.isalnum() or c in ' _-')[:50].strip()
    filename = f"SF330_Section_E_{safe_name.replace(' ', '_')}.docx"
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/employees/<int:id>', methods=['PUT'])
def update_employee(id):
    employee = Employee.query.get_or_404(id)
    data = request.json
    
    # Helper to convert empty strings to None for integer fields
    def to_int_or_none(value):
        if value is None or value == '':
            return None
        try:
            return int(value)
        except (ValueError, TypeError):
            return None
    
    # Handle name fields
    first_name = data.get('first_name', employee.first_name) or ''
    middle_name = data.get('middle_name', employee.middle_name) or ''
    last_name = data.get('last_name', employee.last_name) or ''
    nickname = data.get('nickname', employee.nickname)
    
    # Build full name from components
    name_parts = [first_name.strip(), middle_name.strip(), last_name.strip()]
    full_name = ' '.join(p for p in name_parts if p)
    
    employee.first_name = first_name.strip() if first_name.strip() else None
    employee.middle_name = middle_name.strip() if middle_name.strip() else None
    employee.last_name = last_name.strip() if last_name.strip() else None
    employee.nickname = nickname.strip() if nickname and nickname.strip() else None
    employee.name = full_name if full_name else data.get('name', employee.name)
    employee.city = data.get('city', employee.city) or None
    employee.state = data.get('state', employee.state) or None
    employee.title = data.get('title', employee.title) or None
    employee.role = data.get('role', employee.role) or None
    employee.years_experience_total = to_int_or_none(data.get('years_experience_total', employee.years_experience_total))
    employee.years_experience_firm = to_int_or_none(data.get('years_experience_firm', employee.years_experience_firm))
    employee.education = data.get('education', employee.education) or None
    employee.registrations = data.get('registrations', employee.registrations) or None
    employee.training = data.get('training', employee.training) or None
    employee.other_qualifications = data.get('other_qualifications', employee.other_qualifications) or None
    employee.bio = data.get('bio', employee.bio) or None
    employee.firm_id = to_int_or_none(data.get('firm_id', employee.firm_id))
    
    db.session.commit()
    return jsonify({'success': True})


@app.route('/employees/<int:id>/project-experience', methods=['POST'])
def add_project_experience(id):
    employee = Employee.query.get_or_404(id)
    data = request.json
    
    exp = EmployeeProjectExperience(
        employee_id=employee.id,
        project_title=data.get('project_title', ''),
        location=data.get('location'),
        owner_name=data.get('owner_name'),
        project_cost=data.get('project_cost'),
        year_completed=data.get('year_completed'),
        role_performed=data.get('role_performed'),
        brief_description=data.get('brief_description'),
        firm_name=data.get('firm_name'),
        is_current_firm=data.get('is_current_firm', False)
    )
    db.session.add(exp)
    db.session.commit()
    return jsonify({'success': True, 'id': exp.id})


@app.route('/employees/<int:id>/project-experience/<int:exp_id>', methods=['PUT'])
def update_project_experience(id, exp_id):
    exp = EmployeeProjectExperience.query.filter_by(id=exp_id, employee_id=id).first_or_404()
    data = request.json
    
    exp.project_title = data.get('project_title', exp.project_title)
    exp.location = data.get('location', exp.location)
    exp.owner_name = data.get('owner_name', exp.owner_name)
    exp.project_cost = data.get('project_cost', exp.project_cost)
    exp.year_completed = data.get('year_completed', exp.year_completed)
    exp.role_performed = data.get('role_performed', exp.role_performed)
    exp.brief_description = data.get('brief_description', exp.brief_description)
    exp.firm_name = data.get('firm_name', exp.firm_name)
    exp.is_current_firm = data.get('is_current_firm', exp.is_current_firm)
    
    db.session.commit()
    return jsonify({'success': True})


@app.route('/employees/<int:id>/project-experience/<int:exp_id>', methods=['DELETE'])
def delete_project_experience(id, exp_id):
    exp = EmployeeProjectExperience.query.filter_by(id=exp_id, employee_id=id).first_or_404()
    db.session.delete(exp)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/employees/<int:id>/alternate-bios', methods=['GET'])
@login_required
def get_alternate_bios(id):
    """Get all alternate bios for an employee"""
    employee = Employee.query.get_or_404(id)
    bios = EmployeeAlternateBio.query.filter_by(employee_id=id).order_by(EmployeeAlternateBio.created_at.desc()).all()
    return jsonify([{
        'id': b.id,
        'label': b.label,
        'bio': b.bio,
        'created_at': b.created_at.isoformat() if b.created_at else None
    } for b in bios])


@app.route('/employees/<int:id>/alternate-bios', methods=['POST'])
@login_required
def add_alternate_bio(id):
    """Add a new alternate bio for an employee"""
    employee = Employee.query.get_or_404(id)
    data = request.json
    
    alt_bio = EmployeeAlternateBio(
        employee_id=id,
        label=data.get('label', 'New Bio'),
        bio=data.get('bio', '')
    )
    db.session.add(alt_bio)
    db.session.commit()
    
    return jsonify({
        'success': True,
        'id': alt_bio.id,
        'label': alt_bio.label,
        'bio': alt_bio.bio
    })


@app.route('/employees/<int:id>/alternate-bios/<int:bio_id>', methods=['PUT'])
@login_required
def update_alternate_bio(id, bio_id):
    """Update an alternate bio"""
    alt_bio = EmployeeAlternateBio.query.filter_by(id=bio_id, employee_id=id).first_or_404()
    data = request.json
    
    if 'label' in data:
        alt_bio.label = data['label']
    if 'bio' in data:
        alt_bio.bio = data['bio']
    
    db.session.commit()
    return jsonify({'success': True})


@app.route('/employees/<int:id>/alternate-bios/<int:bio_id>', methods=['DELETE'])
@login_required
def delete_alternate_bio(id, bio_id):
    """Delete an alternate bio"""
    alt_bio = EmployeeAlternateBio.query.filter_by(id=bio_id, employee_id=id).first_or_404()
    db.session.delete(alt_bio)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/employees/<int:id>/generate-bio', methods=['POST'])
@login_required
def generate_bio_with_ai(id):
    """Generate an alternate bio using AI based on selected projects and user direction"""
    from gemini_service import generate_employee_bio
    
    employee = Employee.query.get_or_404(id)
    data = request.json
    
    project_ids = data.get('project_ids', [])
    direction = data.get('direction', '')
    
    selected_projects = []
    if project_ids:
        experiences = EmployeeProjectExperience.query.filter(
            EmployeeProjectExperience.id.in_(project_ids),
            EmployeeProjectExperience.employee_id == id
        ).all()
        
        for exp in experiences:
            selected_projects.append({
                'project_title': exp.project_title,
                'role_performed': exp.role_performed,
                'year_completed': exp.year_completed,
                'location': exp.location,
                'owner_name': exp.owner_name,
                'brief_description': exp.brief_description
            })
    
    employee_info = {
        'name': employee.name,
        'title': employee.title,
        'role': employee.role,
        'years_experience_total': employee.years_experience_total,
        'years_experience_firm': employee.years_experience_firm,
        'education': employee.education,
        'registrations': employee.registrations,
        'current_bio': employee.bio
    }
    
    try:
        generated_bio = generate_employee_bio(employee_info, selected_projects, direction)
        return jsonify({'success': True, 'bio': generated_bio})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/employees/<int:id>/format-registrations', methods=['POST'])
def format_registrations_ai(id):
    """Use AI to format professional registrations into SF330-compatible format"""
    from gemini_service import client
    
    employee = Employee.query.get_or_404(id)
    data = request.json
    raw_text = data.get('registrations', '').strip()
    
    if not raw_text:
        return jsonify({'success': False, 'error': 'No registration data provided'})
    
    prompt = """Format the following professional registrations/licenses into a standardized SF330-compatible format.

Rules:
1. Group by license type (PE, SE, etc.)
2. For each license type that has multiple states, combine all states into standard 2-letter abbreviations on ONE line, comma-separated
3. Format licenses with numbers as: TYPE State (#number)
4. Items that are NOT professional engineering/structural licenses (like certifications, training, pilot licenses, rope access, etc.) should be grouped under "Other" on a single line, separated by semicolons
5. Use standard US state abbreviations (Alabama=AL, Alaska=AK, Arizona=AZ, Arkansas=AR, California=CA, Colorado=CO, Connecticut=CT, Delaware=DE, District of Columbia=DC, Florida=FL, Georgia=GA, Hawaii=HI, Idaho=ID, Illinois=IL, Indiana=IN, Iowa=IA, Kansas=KS, Kentucky=KY, Louisiana=LA, Maine=ME, Maryland=MD, Massachusetts=MA, Michigan=MI, Minnesota=MN, Mississippi=MS, Missouri=MO, Montana=MT, Nebraska=NE, Nevada=NV, New Hampshire=NH, New Jersey=NJ, New Mexico=NM, New York=NY, North Carolina=NC, North Dakota=ND, Ohio=OH, Oklahoma=OK, Oregon=OR, Pennsylvania=PA, Puerto Rico=PR, Rhode Island=RI, South Carolina=SC, South Dakota=SD, Tennessee=TN, Texas=TX, Utah=UT, Vermont=VT, Virginia=VA, Washington=WA, West Virginia=WV, Wisconsin=WI, Wyoming=WY)
6. Each license type gets its own line
7. Sort state abbreviations alphabetically within each line

Example input:
SE #000958, Georgia
PE, Alabama
PE, Florida
PE, Georgia
Level 1 Rope Access Technician
FAA Part 107 UAS Remote Pilot

Example output:
SE Georgia (#000958)
PE AL, FL, GA
Other FAA Part 107 UAS Remote Pilot; SPRAT Level 1 Rope Access Tech

Now format this data:
""" + raw_text + """

Return ONLY the formatted text, no explanations or extra text."""

    try:
        response = client.models.generate_content(
            model='gemini-2.0-flash',
            contents=prompt
        )
        formatted = response.text.strip()
        return jsonify({'success': True, 'formatted': formatted})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/employees/<int:id>/experience/<int:exp_id>/toggle-sf330', methods=['POST'])
def toggle_sf330_include(id, exp_id):
    """Toggle SF330 inclusion flag for a project experience"""
    exp = EmployeeProjectExperience.query.filter_by(id=exp_id, employee_id=id).first_or_404()
    exp.sf330_include = not exp.sf330_include
    db.session.commit()
    return jsonify({'success': True, 'sf330_include': exp.sf330_include})


@app.route('/employees/<int:id>/update-resume-orders', methods=['POST'])
def update_resume_orders(id):
    """Update resume_order for multiple project experiences"""
    data = request.json
    orders = data.get('orders', {})
    try:
        for exp_id_str, order_val in orders.items():
            exp = EmployeeProjectExperience.query.filter_by(id=int(exp_id_str), employee_id=id).first()
            if exp:
                exp.resume_order = int(order_val) if order_val not in (None, '', 0) else None
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)})


@app.route('/employees/<int:id>/clear-resume-orders', methods=['POST'])
def clear_resume_orders(id):
    """Clear all resume_order values for an employee's project experiences"""
    try:
        exps = EmployeeProjectExperience.query.filter_by(employee_id=id).all()
        for exp in exps:
            exp.resume_order = None
        db.session.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)})


@app.route('/employees/<int:id>/get-experiences')
def get_experiences(id):
    """Get experiences by IDs for merge modal"""
    ids = request.args.get('ids', '').split(',')
    ids = [int(i) for i in ids if i.isdigit()]
    
    experiences = EmployeeProjectExperience.query.filter(
        EmployeeProjectExperience.id.in_(ids),
        EmployeeProjectExperience.employee_id == id
    ).all()
    
    return jsonify([{
        'id': e.id,
        'project_title': e.project_title,
        'location': e.location,
        'owner_name': e.owner_name,
        'project_cost': e.project_cost,
        'year_completed': e.year_completed,
        'role_performed': e.role_performed,
        'brief_description': e.brief_description,
        'firm_name': e.firm_name
    } for e in experiences])


@app.route('/api/ai-merge-field', methods=['POST'])
def ai_merge_field():
    """AI merge a single field from multiple values"""
    from gemini_service import merge_field_values
    
    data = request.json
    field_key = data.get('field_key')
    values = data.get('values', [])
    
    if len(values) < 2:
        return jsonify({'success': False, 'error': 'Need at least 2 values to merge'})
    
    try:
        merged_value = merge_field_values(field_key, values)
        return jsonify({'success': True, 'merged_value': merged_value})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/employees/<int:id>/save-merged-experience', methods=['POST'])
def save_merged_experience(id):
    """Save a manually merged experience"""
    data = request.json
    experience_ids = data.get('experience_ids', [])
    merged_data = data.get('merged_data', {})
    keep_originals = data.get('keep_originals', False)
    
    if len(experience_ids) < 2:
        return jsonify({'success': False, 'error': 'Need at least 2 experiences to merge'})
    
    experiences = EmployeeProjectExperience.query.filter(
        EmployeeProjectExperience.id.in_(experience_ids),
        EmployeeProjectExperience.employee_id == id
    ).all()
    
    if len(experiences) < 2:
        return jsonify({'success': False, 'error': 'Could not find all selected experiences'})
    
    try:
        new_exp = EmployeeProjectExperience(
            employee_id=id,
            project_title=merged_data.get('project_title', ''),
            location=merged_data.get('location', ''),
            owner_name=merged_data.get('owner_name', ''),
            project_cost=merged_data.get('project_cost', ''),
            year_completed=merged_data.get('year_completed', ''),
            role_performed=merged_data.get('role_performed', ''),
            brief_description=merged_data.get('brief_description', ''),
            firm_name=merged_data.get('firm_name', ''),
            is_current_firm=any(e.is_current_firm for e in experiences),
            sf330_include=any(e.sf330_include for e in experiences)
        )
        db.session.add(new_exp)
        
        if not keep_originals:
            for exp in experiences:
                db.session.delete(exp)
        
        db.session.commit()
        return jsonify({'success': True, 'merged_id': new_exp.id})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)})


@app.route('/api/experience/<int:exp_id>/alternate-descriptions', methods=['GET'])
def get_experience_alternate_descriptions(exp_id):
    """Get all alternate descriptions for an experience"""
    exp = EmployeeProjectExperience.query.get_or_404(exp_id)
    return jsonify([{
        'id': d.id,
        'label': d.label,
        'description': d.description
    } for d in exp.alternate_descriptions])


@app.route('/api/experience/<int:exp_id>/alternate-descriptions', methods=['POST'])
def add_experience_alternate_description(exp_id):
    """Add an alternate description to an experience"""
    exp = EmployeeProjectExperience.query.get_or_404(exp_id)
    data = request.json
    
    alt = ExperienceAlternateDescription(
        experience_id=exp_id,
        label=data.get('label', 'Version'),
        description=data.get('description', '')
    )
    db.session.add(alt)
    db.session.commit()
    return jsonify({'success': True, 'id': alt.id})


@app.route('/api/experience/alternate-descriptions/<int:alt_id>', methods=['PUT'])
def update_experience_alternate_description(alt_id):
    """Update an alternate description"""
    alt = ExperienceAlternateDescription.query.get_or_404(alt_id)
    data = request.json
    alt.label = data.get('label', alt.label)
    alt.description = data.get('description', alt.description)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/experience/alternate-descriptions/<int:alt_id>', methods=['DELETE'])
def delete_experience_alternate_description(alt_id):
    """Delete an alternate description"""
    alt = ExperienceAlternateDescription.query.get_or_404(alt_id)
    db.session.delete(alt)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/experience/<int:exp_id>/generate-alternate-description', methods=['POST'])
@login_required
def generate_alternate_experience_description(exp_id):
    """Generate an alternate description using AI"""
    from gemini_service import generate_alternate_project_writeup
    
    exp = EmployeeProjectExperience.query.get_or_404(exp_id)
    data = request.json or {}
    direction = data.get('direction', '')
    
    employee_name = exp.employee.name if exp.employee else 'Unknown'
    
    linked_project_desc = None
    if exp.linked_project_id and exp.linked_project:
        linked_project_desc = exp.linked_project.brief_description
    
    try:
        description = generate_alternate_project_writeup(
            project_title=exp.project_title or '',
            current_description=exp.brief_description or '',
            role=exp.role_performed or '',
            employee_name=employee_name,
            location=exp.location or '',
            owner_name=exp.owner_name or '',
            linked_project_description=linked_project_desc,
            direction=direction
        )
        return jsonify({'success': True, 'description': description})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/api/experience/<int:exp_id>/select-description', methods=['POST'])
@login_required
def select_experience_description(exp_id):
    """Set which description (main or alternate) to use for the resume"""
    exp = EmployeeProjectExperience.query.get_or_404(exp_id)
    data = request.json or {}
    alt_desc_id = data.get('alt_description_id')
    
    if alt_desc_id is None or alt_desc_id == '' or alt_desc_id == 'main':
        exp.selected_alt_description_id = None
    else:
        alt = ExperienceAlternateDescription.query.get(alt_desc_id)
        if alt and alt.experience_id == exp_id:
            exp.selected_alt_description_id = alt_desc_id
        else:
            return jsonify({'success': False, 'error': 'Invalid alternate description'}), 400
    
    db.session.commit()
    return jsonify({
        'success': True,
        'selected_id': exp.selected_alt_description_id,
        'active_description': exp.active_description,
        'active_label': exp.active_description_label
    })


@app.route('/api/projects/<int:project_id>/copy-to-resume', methods=['POST'])
def copy_project_to_resume(project_id):
    """Copy a project description to an employee's resume experience"""
    project = Project.query.get_or_404(project_id)
    data = request.json
    employee_ids = data.get('employee_ids', [])
    description_type = data.get('description_type', 'main')
    alt_desc_id = data.get('alternate_description_id')
    
    if not employee_ids:
        return jsonify({'error': 'No employees selected'}), 400
    
    description = project.brief_description
    if description_type == 'alternate' and alt_desc_id:
        alt = ProjectAlternateDescription.query.get(alt_desc_id)
        if alt:
            description = alt.description
    
    created = 0
    for emp_id in employee_ids:
        existing = EmployeeProjectExperience.query.filter_by(
            employee_id=emp_id,
            project_title=project.title
        ).first()
        
        if not existing:
            exp = EmployeeProjectExperience(
                employee_id=emp_id,
                project_title=project.title,
                location=project.location,
                owner_name=project.owner_name,
                project_cost=project.project_cost,
                year_completed=project.year_completed_professional,
                brief_description=description,
                is_current_firm=True
            )
            db.session.add(exp)
            created += 1
    
    db.session.commit()
    return jsonify({'success': True, 'created': created})


@app.route('/api/projects/search')
def search_projects():
    """Search projects by title for autocomplete"""
    query = request.args.get('q', '')
    if len(query) < 2:
        return jsonify([])
    
    projects = Project.query.filter(
        Project.title.ilike(f'%{query}%')
    ).order_by(Project.title).limit(10).all()
    
    return jsonify([{
        'id': p.id,
        'title': p.title,
        'location': p.location or '',
        'year_completed': p.year_completed_professional or '',
        'project_cost': p.project_cost or '',
        'owner_name': p.owner_name or ''
    } for p in projects])


@app.route('/api/experience/<int:exp_id>/copy-to-employees', methods=['POST'])
def copy_experience_to_employees(exp_id):
    """Copy a project experience to other employees"""
    exp = EmployeeProjectExperience.query.get_or_404(exp_id)
    data = request.json
    employee_ids = data.get('employee_ids', [])
    
    copied_count = 0
    for emp_id in employee_ids:
        if emp_id != exp.employee_id:
            existing = EmployeeProjectExperience.query.filter_by(
                employee_id=emp_id,
                project_title=exp.project_title
            ).first()
            if not existing:
                new_exp = EmployeeProjectExperience(
                    employee_id=emp_id,
                    project_title=exp.project_title,
                    location=exp.location,
                    owner_name=exp.owner_name,
                    project_cost=exp.project_cost,
                    year_completed=exp.year_completed,
                    role_performed=exp.role_performed,
                    brief_description=exp.brief_description,
                    firm_name=exp.firm_name,
                    is_current_firm=exp.is_current_firm
                )
                db.session.add(new_exp)
                copied_count += 1
    
    db.session.commit()
    return jsonify({'success': True, 'copied_count': copied_count})


@app.route('/api/experience/<int:exp_id>/add-to-projects', methods=['POST'])
def add_experience_to_projects(exp_id):
    """Add a resume project experience to the main projects database"""
    exp = EmployeeProjectExperience.query.get_or_404(exp_id)
    
    existing = Project.query.filter_by(title=exp.project_title).first()
    if existing:
        return jsonify({'success': False, 'error': 'A project with this title already exists', 'project_id': existing.id})
    
    project = Project(
        title=exp.project_title,
        location=exp.location,
        year_completed_professional=exp.year_completed,
        project_cost=exp.project_cost,
        owner_name=exp.owner_name,
        brief_description=exp.brief_description
    )
    db.session.add(project)
    db.session.commit()
    
    return jsonify({'success': True, 'project_id': project.id})


@app.route('/api/experience/rewrite-description', methods=['POST'])
def rewrite_experience_description():
    """AI rewrite of experience description using global settings and custom instructions"""
    data = request.json
    custom_instructions = data.get('custom_instructions', '')
    description_text = data.get('description', '')
    
    if not description_text:
        return jsonify({'error': 'No description to rewrite'}), 400
    
    from gemini_service import rewrite_description
    rewritten = rewrite_description(description_text, custom_instructions)
    
    return jsonify({'success': True, 'rewritten': rewritten})


@app.route('/employees/<int:id>', methods=['DELETE'])
def delete_employee(id):
    employee = Employee.query.get_or_404(id)
    EmployeeProjectExperience.query.filter_by(employee_id=id).delete()
    EmployeeProjectLink.query.filter_by(employee_id=id).delete()
    ProposalSelectedEmployee.query.filter_by(employee_id=id).delete()
    db.session.delete(employee)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/employees/merge-duplicates')
def merge_employees_page():
    ids = request.args.get('ids', '')
    id_list = [int(x) for x in ids.split(',') if x.isdigit()]
    if len(id_list) < 2:
        return redirect('/employees')
    
    employees = Employee.query.filter(Employee.id.in_(id_list)).all()
    if len(employees) < 2:
        return redirect('/employees')
    
    employees_data = [{
        'id': e.id,
        'name': e.name,
        'title': e.title,
        'role': e.role,
        'years_experience_firm': e.years_experience_firm,
        'years_experience_total': e.years_experience_total,
        'education': e.education,
        'registrations': e.registrations,
        'training': e.training,
        'other_qualifications': e.other_qualifications,
        'bio': e.bio
    } for e in employees]
    
    return render_template('employee_merge_duplicates.html', employees=employees, employees_json=employees_data)


@app.route('/employees/merge-duplicates', methods=['POST'])
def merge_employees():
    data = request.json
    primary_id = data.get('primary_id')
    merge_ids = data.get('merge_ids', [])
    merged_data = data.get('merged_data', {})
    
    primary = Employee.query.get_or_404(primary_id)
    
    integer_fields = ['years_experience_firm', 'years_experience_total']
    for key, value in merged_data.items():
        if hasattr(primary, key):
            if key in integer_fields:
                if value == '' or value is None:
                    value = None
                else:
                    try:
                        value = int(value)
                    except (ValueError, TypeError):
                        value = None
            setattr(primary, key, value)
    
    for merge_id in merge_ids:
        if merge_id == primary_id:
            continue
        merge_emp = Employee.query.get(merge_id)
        if not merge_emp:
            continue
        
        links_to_transfer = EmployeeProjectLink.query.filter_by(employee_id=merge_id).all()
        for link in links_to_transfer:
            existing = EmployeeProjectLink.query.filter_by(
                employee_id=primary_id, project_id=link.project_id
            ).first()
            if not existing:
                link.employee_id = primary_id
            else:
                db.session.delete(link)
        
        experiences_to_transfer = list(EmployeeProjectExperience.query.filter_by(employee_id=merge_id).all())
        for exp in experiences_to_transfer:
            existing = EmployeeProjectExperience.query.filter_by(
                employee_id=primary_id,
                project_title=exp.project_title
            ).first()
            if not existing:
                exp.employee_id = primary_id
                db.session.flush()
            else:
                db.session.delete(exp)
        
        alt_bios_to_transfer = list(EmployeeAlternateBio.query.filter_by(employee_id=merge_id).all())
        for alt_bio in alt_bios_to_transfer:
            alt_bio.employee_id = primary_id
            db.session.flush()
        
        ProposalSelectedEmployee.query.filter_by(employee_id=merge_id).delete()
        
        db.session.flush()
        db.session.delete(merge_emp)
    
    db.session.commit()
    return jsonify({'success': True, 'redirect': f'/employees/{primary_id}'})


def organize_projects_hierarchically(project_list):
    """Organize projects so contract projects come first with their task orders nested below"""
    contracts = [p for p in project_list if p.project_type != 'task_order']
    task_orders = [p for p in project_list if p.project_type == 'task_order']
    
    task_orders_by_parent = {}
    orphan_task_orders = []
    for to in task_orders:
        if to.parent_contract_id:
            if to.parent_contract_id not in task_orders_by_parent:
                task_orders_by_parent[to.parent_contract_id] = []
            task_orders_by_parent[to.parent_contract_id].append(to)
        else:
            orphan_task_orders.append(to)
    
    result = []
    for contract in sorted(contracts, key=lambda x: x.title):
        result.append({'project': contract, 'is_task_order': False, 'indent': 0})
        for to in sorted(task_orders_by_parent.get(contract.id, []), key=lambda x: x.title):
            result.append({'project': to, 'is_task_order': True, 'indent': 1})
    
    for to in sorted(orphan_task_orders, key=lambda x: x.title):
        result.append({'project': to, 'is_task_order': True, 'indent': 0})
    
    return result


@app.route('/projects')
def projects():
    all_projects = Project.query.order_by(Project.title).all()
    firms = Firm.query.order_by(Firm.name).all()
    
    hierarchical_projects = organize_projects_hierarchically(all_projects)
    
    # Group projects by firm
    projects_by_firm = {}
    unassigned_projects = []
    for project in all_projects:
        if project.firm_id:
            if project.firm_id not in projects_by_firm:
                projects_by_firm[project.firm_id] = []
            projects_by_firm[project.firm_id].append(project)
        else:
            unassigned_projects.append(project)
    
    # Define colors for each firm (cycle through if more firms than colors)
    firm_colors = [
        {'bg': 'bg-blue-50', 'border': 'border-blue-200', 'tab': 'bg-blue-600', 'tab_inactive': 'bg-blue-100 text-blue-700', 'badge': 'bg-blue-100 text-blue-800'},
        {'bg': 'bg-green-50', 'border': 'border-green-200', 'tab': 'bg-green-600', 'tab_inactive': 'bg-green-100 text-green-700', 'badge': 'bg-green-100 text-green-800'},
        {'bg': 'bg-purple-50', 'border': 'border-purple-200', 'tab': 'bg-purple-600', 'tab_inactive': 'bg-purple-100 text-purple-700', 'badge': 'bg-purple-100 text-purple-800'},
        {'bg': 'bg-orange-50', 'border': 'border-orange-200', 'tab': 'bg-orange-600', 'tab_inactive': 'bg-orange-100 text-orange-700', 'badge': 'bg-orange-100 text-orange-800'},
        {'bg': 'bg-pink-50', 'border': 'border-pink-200', 'tab': 'bg-pink-600', 'tab_inactive': 'bg-pink-100 text-pink-700', 'badge': 'bg-pink-100 text-pink-800'},
        {'bg': 'bg-teal-50', 'border': 'border-teal-200', 'tab': 'bg-teal-600', 'tab_inactive': 'bg-teal-100 text-teal-700', 'badge': 'bg-teal-100 text-teal-800'},
    ]
    
    # Build firm data with colors and hierarchical projects
    firm_data = []
    for i, firm in enumerate(firms):
        color = firm_colors[i % len(firm_colors)]
        firm_projects = projects_by_firm.get(firm.id, [])
        firm_hierarchical = organize_projects_hierarchically(firm_projects)
        firm_data.append({
            'firm': firm,
            'projects': firm_projects,
            'hierarchical_projects': firm_hierarchical,
            'color': color
        })
    
    unassigned_hierarchical = organize_projects_hierarchically(unassigned_projects)
    
    return render_template('projects.html', projects=all_projects, hierarchical_projects=hierarchical_projects,
                          firms=firms, firm_data=firm_data, 
                          unassigned_projects=unassigned_projects, unassigned_hierarchical=unassigned_hierarchical,
                          firm_colors=firm_colors)


@app.route('/projects/add', methods=['GET', 'POST'])
def add_project():
    if request.method == 'POST':
        data = request.form
        firm_id = data.get('firm_id')
        parent_contract_id = data.get('parent_contract_id')
        project = Project(
            title=data.get('title', ''),
            location=data.get('location'),
            year_completed_professional=data.get('year_completed_professional'),
            year_completed_construction=data.get('year_completed_construction'),
            owner_name=data.get('owner_name'),
            owner_contact_name=data.get('owner_contact_name'),
            owner_contact_phone=data.get('owner_contact_phone'),
            project_cost=data.get('project_cost'),
            project_delivery_method=data.get('project_delivery_method'),
            brief_description=data.get('brief_description'),
            relevance_writeup=data.get('relevance_writeup'),
            firm_id=int(firm_id) if firm_id else None,
            project_type=data.get('project_type', 'contract'),
            parent_contract_id=int(parent_contract_id) if parent_contract_id else None
        )
        db.session.add(project)
        db.session.commit()
        return redirect(f'/projects/{project.id}')
    
    firms = Firm.query.order_by(Firm.name).all()
    contract_projects = Project.query.filter_by(project_type='contract').order_by(Project.title).all()
    return render_template('project_add.html', firms=firms, contract_projects=contract_projects)


@app.route('/projects/<int:id>')
def project_detail(id):
    from models import MarketingPhoto
    project = Project.query.get_or_404(id)
    employee_links = EmployeeProjectLink.query.filter_by(project_id=id).all()
    all_employees = Employee.query.all()
    
    # Get personnel writeups linked to this project
    personnel_writeups = EmployeeProjectExperience.query.filter_by(linked_project_id=id).all()
    
    # Serialize personnel writeups for JavaScript
    personnel_writeups_json = [{
        'id': pw.id,
        'employee_id': pw.employee_id,
        'employee_name': pw.employee.name if pw.employee else 'Unknown',
        'project_title': pw.project_title,
        'role_performed': pw.role_performed,
        'year_completed': pw.year_completed,
        'brief_description': pw.brief_description,
        'location': pw.location,
        'owner_name': pw.owner_name,
        'project_cost': pw.project_cost,
        'firm_name': pw.firm_name,
        'is_current_firm': pw.is_current_firm
    } for pw in personnel_writeups]
    
    # Find marketing photos tagged with this project's name
    project_tag = f"#{project.title.replace(' ', '')}"
    all_marketing = MarketingPhoto.query.all()
    marketing_photos = [p for p in all_marketing if project_tag.lower() in (p.tags or '').lower()]
    
    firms = Firm.query.order_by(Firm.name).all()
    contract_projects = Project.query.filter(
        Project.project_type != 'task_order',
        Project.id != project.id
    ).order_by(Project.title).all()
    return render_template('project_detail.html', project=project, employee_links=employee_links, 
                           all_employees=all_employees, marketing_photos=marketing_photos,
                           personnel_writeups=personnel_writeups, personnel_writeups_json=personnel_writeups_json,
                           firms=firms, contract_projects=contract_projects)


@app.route('/projects/<int:id>', methods=['PUT'])
def update_project(id):
    project = Project.query.get_or_404(id)
    data = request.json
    
    project.title = data.get('title', project.title)
    project.location = data.get('location', project.location)
    project.year_completed_professional = data.get('year_completed_professional', project.year_completed_professional)
    project.year_completed_construction = data.get('year_completed_construction', project.year_completed_construction)
    project.owner_name = data.get('owner_name', project.owner_name)
    project.owner_contact_name = data.get('owner_contact_name', project.owner_contact_name)
    project.owner_contact_phone = data.get('owner_contact_phone', project.owner_contact_phone)
    project.owner_contact_email = data.get('owner_contact_email', project.owner_contact_email)
    project.project_cost = data.get('project_cost', project.project_cost)
    project.brief_description = data.get('brief_description', project.brief_description)
    project.relevance_writeup = data.get('relevance_writeup', project.relevance_writeup)
    project.is_with_other_firm = data.get('is_with_other_firm') in [True, 'true', 'True', '1', 1]
    project.other_firm_name = data.get('other_firm_name', project.other_firm_name) if project.is_with_other_firm else None
    
    if 'project_type' in data:
        project.project_type = data.get('project_type', 'contract')
    if 'parent_contract_id' in data:
        parent_id = data.get('parent_contract_id')
        project.parent_contract_id = int(parent_id) if parent_id else None
    
    # Handle firm_id assignment
    if 'firm_id' in data:
        firm_id = data.get('firm_id')
        project.firm_id = int(firm_id) if firm_id else None
    
    owner_contact_id = data.get('owner_contact_id')
    if owner_contact_id:
        project.owner_contact_id = int(owner_contact_id)
    else:
        project.owner_contact_id = None
        contact_name = data.get('owner_contact_name', '').strip()
        if contact_name:
            existing = ClientContact.query.filter(ClientContact.name.ilike(contact_name)).first()
            if existing:
                project.owner_contact_id = existing.id
            else:
                new_contact = ClientContact(
                    name=contact_name,
                    agency=data.get('owner_name', ''),
                    phone=data.get('owner_contact_phone', ''),
                    email=data.get('owner_contact_email', '')
                )
                db.session.add(new_contact)
                db.session.flush()
                project.owner_contact_id = new_contact.id
    
    db.session.commit()
    return jsonify({'success': True, 'contact_created': owner_contact_id is None and project.owner_contact_id is not None})


@app.route('/projects/<int:id>', methods=['DELETE'])
def delete_project(id):
    project = Project.query.get_or_404(id)
    EmployeeProjectLink.query.filter_by(project_id=id).delete()
    ProposalEmployeeRelevantProject.query.filter_by(project_id=id).delete()
    ProposalSelectedProject.query.filter_by(project_id=id).delete()
    db.session.delete(project)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/projects/<int:id>/download')
def download_project(id):
    """Download project as Word document - either SF330 Section F template or plain format"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from werkzeug.utils import secure_filename
    import io
    import os
    import re
    
    project = Project.query.get_or_404(id)
    format_type = request.args.get('format', 'plain')
    
    # Get firm information for the project
    firm = Firm.query.get(project.firm_id) if project.firm_id else None
    
    # Get team members linked to this project
    team_links = EmployeeProjectLink.query.filter_by(project_id=id).all()
    
    # Helper function to sanitize filename
    def make_safe_filename(name, prefix="Project"):
        if not name:
            return f"{prefix}_{id}"
        # Remove unsafe characters and limit length
        safe = secure_filename(name[:50])
        return safe if safe else f"{prefix}_{id}"
    
    if format_type == 'template':
        # First try to load custom template from object storage
        doc = None
        try:
            client = get_storage_client()
            template_bytes = client.download_as_bytes('templates/sf330_section_f_custom.docx')
            if template_bytes:
                doc = Document(io.BytesIO(template_bytes))
        except:
            pass
        
        # Fall back to attached SF330 template file
        if doc is None:
            template_path = 'attached_assets/SF330_Section_F_Template_(1)-mwc_1770394097919.docx'
            if os.path.exists(template_path):
                doc = Document(template_path)
            else:
                doc = create_default_sf330_template()
        
        # Build key personnel string
        key_personnel_str = ''
        if team_links:
            team_members = []
            for link in team_links:
                emp = Employee.query.get(link.employee_id)
                if emp:
                    role = f" - {link.role_on_project}" if link.role_on_project else ""
                    team_members.append(f"{emp.name}{role}")
            key_personnel_str = '\n'.join(team_members) if team_members else ''
        
        # Define all placeholder replacements (same tags as company template)
        placeholders = {
            '{{PROJECT_TITLE}}': project.title or '',
            '{{PROJECT_LOCATION}}': project.location or '',
            '{{PROJECT_COST}}': project.project_cost or '',
            '{{BRIEF_DESCRIPTION}}': project.brief_description or '',
            '{{DELIVERY_METHOD}}': project.project_delivery_method or '',
            '{{YEAR_COMPLETED_PROFESSIONAL}}': project.year_completed_professional or '',
            '{{YEAR_COMPLETED_CONSTRUCTION}}': project.year_completed_construction or '',
            '{{OWNER_NAME}}': project.owner_name or '',
            '{{OWNER_CONTACT}}': project.owner_contact_name or '',
            '{{OWNER_PHONE}}': project.owner_contact_phone or '',
            '{{OWNER_EMAIL}}': project.owner_contact_email or '',
            '{{FIRM_NAME}}': firm.name if firm else '',
            '{{FIRM_CITY}}': firm.city if firm else '',
            '{{FIRM_STATE}}': firm.state if firm else '',
            '{{FIRM_ROLE}}': 'Prime',
            '{{KEY_PERSONNEL}}': key_personnel_str,
        }
        
        # Helper function to replace placeholders in a paragraph (handles split runs)
        def replace_sf330_placeholders_in_para(para, placeholders):
            for placeholder, value in placeholders.items():
                if placeholder in para.text:
                    replaced = False
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)
                            replaced = True
                    if not replaced and para.runs:
                        full_text = para.text
                        new_text = full_text.replace(placeholder, value)
                        if full_text != new_text:
                            for i, run in enumerate(para.runs):
                                if i == 0:
                                    run.text = new_text
                                else:
                                    run.text = ''
        
        # Replace placeholders in paragraphs
        for para in doc.paragraphs:
            replace_sf330_placeholders_in_para(para, placeholders)
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_sf330_placeholders_in_para(para, placeholders)
        
        filename = f"SF330_Section_F_{make_safe_filename(project.title, 'Project')}.docx"
        
    elif format_type == 'company':
        # Use company template with placeholder replacement
        doc = None
        try:
            client = get_storage_client()
            template_bytes = client.download_as_bytes('templates/company_template_custom.docx')
            if template_bytes:
                doc = Document(io.BytesIO(template_bytes))
        except:
            pass
        
        # Fall back to default company template
        if doc is None:
            doc = create_default_company_template()
        
        # Build key personnel string
        key_personnel_str = ''
        if team_links:
            team_members = []
            for link in team_links:
                emp = Employee.query.get(link.employee_id)
                if emp:
                    role = f" ({link.role_on_project})" if link.role_on_project else ""
                    team_members.append(f"{emp.name}{role}")
            key_personnel_str = ', '.join(team_members) if team_members else 'Not specified'
        
        # Define all placeholder replacements
        placeholders = {
            '{{PROJECT_TITLE}}': project.title or '',
            '{{PROJECT_LOCATION}}': project.location or '',
            '{{PROJECT_COST}}': project.project_cost or '',
            '{{BRIEF_DESCRIPTION}}': project.brief_description or '',
            '{{DELIVERY_METHOD}}': project.project_delivery_method or '',
            '{{YEAR_COMPLETED_PROFESSIONAL}}': project.year_completed_professional or '',
            '{{YEAR_COMPLETED_CONSTRUCTION}}': project.year_completed_construction or '',
            '{{OWNER_NAME}}': project.owner_name or '',
            '{{OWNER_CONTACT}}': project.owner_contact_name or '',
            '{{OWNER_PHONE}}': project.owner_contact_phone or '',
            '{{OWNER_EMAIL}}': project.owner_contact_email or '',
            '{{FIRM_NAME}}': firm.name if firm else '',
            '{{FIRM_CITY}}': firm.city if firm else '',
            '{{FIRM_STATE}}': firm.state if firm else '',
            '{{FIRM_ROLE}}': 'Prime',
            '{{KEY_PERSONNEL}}': key_personnel_str,
        }
        
        # Helper function to replace placeholders in a paragraph (handles split runs)
        def replace_placeholders_in_para(para, placeholders):
            for placeholder, value in placeholders.items():
                if placeholder in para.text:
                    # First try to replace within individual runs
                    replaced = False
                    for run in para.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)
                            replaced = True
                    # If placeholder wasn't found in individual runs, it may be split across runs
                    # Rebuild paragraph text and replace
                    if not replaced and para.runs:
                        full_text = para.text
                        new_text = full_text.replace(placeholder, value)
                        if full_text != new_text:
                            # Clear all runs except first, put all text in first run
                            for i, run in enumerate(para.runs):
                                if i == 0:
                                    run.text = new_text
                                else:
                                    run.text = ''
        
        # Replace placeholders in paragraphs
        for para in doc.paragraphs:
            replace_placeholders_in_para(para, placeholders)
        
        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_placeholders_in_para(para, placeholders)
        
        filename = f"Company_{make_safe_filename(project.title, 'Project')}.docx"
        
    else:
        # Create plain Word document
        doc = Document()
        
        # Title
        title_para = doc.add_heading(project.title or 'Untitled Project', 0)
        
        # Basic Info Section
        doc.add_heading('Project Information', level=1)
        
        info_table = doc.add_table(rows=0, cols=2)
        info_table.style = 'Table Grid'
        
        def add_info_row(label, value):
            row = info_table.add_row()
            row.cells[0].paragraphs[0].add_run(label).bold = True
            row.cells[1].text = str(value) if value else ''
        
        add_info_row('Location', project.location)
        add_info_row('Project Cost', project.project_cost)
        add_info_row('Year Completed (Professional)', project.year_completed_professional)
        add_info_row('Year Completed (Construction)', project.year_completed_construction)
        add_info_row('Delivery Method', project.project_delivery_method)
        
        if firm:
            add_info_row('Firm', firm.name)
        
        if project.is_with_other_firm:
            add_info_row('Other Firm', project.other_firm_name)
        
        doc.add_paragraph()
        
        # Owner Information
        doc.add_heading('Owner Information', level=1)
        
        owner_table = doc.add_table(rows=0, cols=2)
        owner_table.style = 'Table Grid'
        
        def add_owner_row(label, value):
            row = owner_table.add_row()
            row.cells[0].paragraphs[0].add_run(label).bold = True
            row.cells[1].text = str(value) if value else ''
        
        add_owner_row('Owner Name', project.owner_name)
        add_owner_row('Contact Name', project.owner_contact_name)
        add_owner_row('Contact Phone', project.owner_contact_phone)
        add_owner_row('Contact Email', project.owner_contact_email)
        
        doc.add_paragraph()
        
        # Description
        doc.add_heading('Project Description', level=1)
        doc.add_paragraph(project.brief_description or 'No description provided.')
        
        # Team Members
        if team_links:
            doc.add_heading('Team Members', level=1)
            team_table = doc.add_table(rows=1, cols=2)
            team_table.style = 'Table Grid'
            
            hdr = team_table.rows[0].cells
            hdr[0].paragraphs[0].add_run('Name').bold = True
            hdr[1].paragraphs[0].add_run('Role').bold = True
            
            for link in team_links:
                emp = Employee.query.get(link.employee_id)
                if emp:
                    row = team_table.add_row()
                    row.cells[0].text = emp.name or ''
                    row.cells[1].text = link.role_on_project or ''
        
        filename = f"Project_{make_safe_filename(project.title, 'Download')}.docx"
    
    # Add project photos if available
    project_photos = ProjectPhoto.query.filter_by(project_id=project.id).all()
    if project_photos:
        from docx.shared import Inches, Pt, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn, nsmap
        from docx.oxml import OxmlElement
        from PIL import Image
        
        storage_client = get_storage_client()
        if storage_client:
            for photo in project_photos:
                try:
                    # Download photo from object storage
                    photo_data = storage_client.download_as_bytes(photo.storage_path)
                    if photo_data:
                        photo_stream = io.BytesIO(photo_data)
                        
                        # Validate it's an image and get dimensions
                        try:
                            img = Image.open(photo_stream)
                            orig_width, orig_height = img.size
                            photo_stream.seek(0)
                        except Exception:
                            print(f"Skipping non-image file: {photo.filename}")
                            continue
                        
                        # Calculate scaled dimensions - max height 2 inches, preserve aspect ratio
                        max_height_inches = 2.0
                        aspect_ratio = orig_width / orig_height
                        # Use 2 inches as max height regardless of original size
                        height_inches = max_height_inches
                        width_inches = height_inches * aspect_ratio
                        
                        # Cap width at 4 inches to prevent overflow
                        if width_inches > 4.0:
                            width_inches = 4.0
                            height_inches = width_inches / aspect_ratio
                        
                        # Add a paragraph for the image, right-aligned
                        img_para = doc.add_paragraph()
                        img_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        
                        # Add the image with explicit sizing
                        run = img_para.add_run()
                        picture = run.add_picture(photo_stream, width=Inches(width_inches), height=Inches(height_inches))
                        
                        # Add caption if available
                        if photo.caption:
                            caption_para = doc.add_paragraph()
                            caption_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                            caption_run = caption_para.add_run(photo.caption)
                            caption_run.italic = True
                            caption_run.font.size = Pt(9)
                except Exception as e:
                    print(f"Error adding photo {photo.filename}: {e}")
    
    # Save to BytesIO and return
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/api/personnel-writeups/<int:pw_id>', methods=['PUT'])
@login_required
def update_personnel_writeup(pw_id):
    """Update a personnel writeup (EmployeeProjectExperience)"""
    pw = EmployeeProjectExperience.query.get_or_404(pw_id)
    data = request.json
    
    if 'role_performed' in data:
        pw.role_performed = data['role_performed']
    if 'brief_description' in data:
        pw.brief_description = data['brief_description']
    if 'project_title' in data:
        pw.project_title = data['project_title']
    
    db.session.commit()
    return jsonify({
        'success': True,
        'id': pw.id,
        'role_performed': pw.role_performed,
        'brief_description': pw.brief_description
    })


@app.route('/api/personnel-writeups/<int:pw_id>/unlink', methods=['POST'])
@login_required
def unlink_personnel_writeup(pw_id):
    """Unlink a personnel writeup from its linked project"""
    pw = EmployeeProjectExperience.query.get_or_404(pw_id)
    pw.linked_project_id = None
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/personnel-writeups/<int:pw_id>/link', methods=['POST'])
@login_required
def link_personnel_writeup(pw_id):
    """Link a personnel writeup to a project"""
    pw = EmployeeProjectExperience.query.get_or_404(pw_id)
    data = request.json
    
    if not data:
        return jsonify({'success': False, 'error': 'Request body is required'}), 400
    
    project_id = data.get('project_id')
    if not project_id:
        return jsonify({'success': False, 'error': 'project_id is required'}), 400
    
    try:
        project_id = int(project_id)
    except (ValueError, TypeError):
        return jsonify({'success': False, 'error': 'project_id must be an integer'}), 400
    
    project = Project.query.get(project_id)
    if not project:
        return jsonify({'success': False, 'error': 'Project not found'}), 404
    
    pw.linked_project_id = project_id
    db.session.commit()
    
    return jsonify({
        'success': True,
        'linked_project_id': project_id,
        'linked_project_title': project.title
    })


@app.route('/api/personnel-writeups/<int:pw_id>/ai-enhance', methods=['POST'])
@login_required
def ai_enhance_personnel_writeup(pw_id):
    """Enhance a personnel writeup using the linked project's description"""
    from gemini_service import enhance_personnel_writeup
    
    pw = EmployeeProjectExperience.query.get_or_404(pw_id)
    data = request.json
    
    project_id = data.get('project_id')
    if not project_id:
        return jsonify({'success': False, 'error': 'project_id is required'}), 400
    
    instructions = data.get('instructions', '')
    
    project = Project.query.get(project_id)
    if not project:
        return jsonify({'success': False, 'error': 'Project not found'}), 404
    
    employee_name = pw.employee.name if pw.employee else 'Unknown'
    
    try:
        enhanced = enhance_personnel_writeup(
            personnel_writeup=pw.brief_description or '',
            project_description=project.brief_description or '',
            employee_name=employee_name,
            role=pw.role_performed or '',
            instructions=instructions
        )
        return jsonify({'success': True, 'enhanced': enhanced})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/projects/<int:id>/sync-personnel-names', methods=['POST'])
@login_required
def sync_personnel_project_names(id):
    """Sync all linked personnel writeup project names to match this project's title"""
    project = Project.query.get_or_404(id)
    
    writeups = EmployeeProjectExperience.query.filter_by(linked_project_id=id).all()
    count = 0
    
    for pw in writeups:
        pw.project_title = project.title
        count += 1
    
    db.session.commit()
    return jsonify({'success': True, 'count': count})


@app.route('/projects/<int:project_id>/team/<int:link_id>', methods=['PUT'])
def update_project_team_member(project_id, link_id):
    link = EmployeeProjectLink.query.get_or_404(link_id)
    if link.project_id != project_id:
        return jsonify({'success': False, 'error': 'Invalid link'}), 400
    
    data = request.json
    link.role_on_project = data.get('role_on_project', link.role_on_project)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/projects/<int:project_id>/team/<int:link_id>', methods=['DELETE'])
def remove_project_team_member(project_id, link_id):
    link = EmployeeProjectLink.query.get_or_404(link_id)
    if link.project_id != project_id:
        return jsonify({'success': False, 'error': 'Invalid link'}), 400
    
    db.session.delete(link)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/projects/merge-duplicates')
def merge_projects_page():
    ids = request.args.get('ids', '')
    id_list = [int(x) for x in ids.split(',') if x.isdigit()]
    if len(id_list) < 2:
        return redirect('/projects')
    
    projects = Project.query.filter(Project.id.in_(id_list)).all()
    if len(projects) < 2:
        return redirect('/projects')
    
    projects_data = [{
        'id': p.id,
        'title': p.title,
        'location': p.location,
        'year_completed_professional': p.year_completed_professional,
        'year_completed_construction': p.year_completed_construction,
        'owner_name': p.owner_name,
        'owner_contact_name': p.owner_contact_name,
        'owner_contact_phone': p.owner_contact_phone,
        'project_cost': p.project_cost,
        'project_delivery_method': p.project_delivery_method,
        'brief_description': p.brief_description,
        'relevance_writeup': p.relevance_writeup
    } for p in projects]
    
    return render_template('project_merge_duplicates.html', projects=projects, projects_json=projects_data)


@app.route('/projects/merge-duplicates', methods=['POST'])
def merge_projects():
    data = request.json
    primary_id = data.get('primary_id')
    merge_ids = data.get('merge_ids', [])
    merged_data = data.get('merged_data', {})
    
    primary = Project.query.get_or_404(primary_id)
    
    for key, value in merged_data.items():
        if hasattr(primary, key):
            setattr(primary, key, value)
    
    for merge_id in merge_ids:
        if merge_id == primary_id:
            continue
        merge_proj = Project.query.get(merge_id)
        if not merge_proj:
            continue
        
        for link in EmployeeProjectLink.query.filter_by(project_id=merge_id).all():
            existing = EmployeeProjectLink.query.filter_by(
                employee_id=link.employee_id, project_id=primary_id
            ).first()
            if not existing:
                link.project_id = primary_id
            else:
                db.session.delete(link)
        
        ProposalEmployeeRelevantProject.query.filter_by(project_id=merge_id).update({'project_id': primary_id})
        ProposalSelectedProject.query.filter_by(project_id=merge_id).delete()
        db.session.delete(merge_proj)
    
    db.session.commit()
    return jsonify({'success': True, 'redirect': f'/projects/{primary_id}'})


@app.route('/projects/<int:project_id>/link-employee', methods=['POST'])
def link_employee_to_project(project_id):
    data = request.json
    employee_id = data.get('employee_id')
    role = data.get('role')
    
    existing = EmployeeProjectLink.query.filter_by(
        employee_id=employee_id, 
        project_id=project_id
    ).first()
    
    if existing:
        existing.role_on_project = role
    else:
        link = EmployeeProjectLink(
            employee_id=employee_id,
            project_id=project_id,
            role_on_project=role
        )
        db.session.add(link)
    
    db.session.commit()
    return jsonify({'success': True})


@app.route('/projects/<int:project_id>/alternate-descriptions', methods=['POST'])
def add_alternate_description(project_id):
    data = request.json
    alt_desc = ProjectAlternateDescription(
        project_id=project_id,
        label=data.get('label', 'Alternate'),
        description=data.get('description', '')
    )
    db.session.add(alt_desc)
    db.session.commit()
    return jsonify({'success': True, 'id': alt_desc.id})


@app.route('/projects/<int:project_id>/alternate-descriptions/<int:alt_id>', methods=['PUT'])
def update_alternate_description(project_id, alt_id):
    alt_desc = ProjectAlternateDescription.query.filter_by(id=alt_id, project_id=project_id).first_or_404()
    data = request.json
    alt_desc.label = data.get('label', alt_desc.label)
    alt_desc.description = data.get('description', alt_desc.description)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/projects/<int:project_id>/alternate-descriptions/<int:alt_id>', methods=['DELETE'])
def delete_alternate_description(project_id, alt_id):
    alt_desc = ProjectAlternateDescription.query.filter_by(id=alt_id, project_id=project_id).first_or_404()
    db.session.delete(alt_desc)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/rewrite-description', methods=['POST'])
def rewrite_description():
    data = request.json
    description = data.get('description', '')
    custom_instructions = data.get('custom_instructions', '')
    
    if not description:
        return jsonify({'success': False, 'error': 'No description provided'})
    
    global_style = AISettings.get_value('ai_writing_style', '')
    global_tone = AISettings.get_value('ai_writing_tone', '')
    
    style_guidance = ""
    if global_style:
        style_guidance += f"\nGlobal style preference: {global_style}"
    if global_tone:
        style_guidance += f"\nGlobal tone preference: {global_tone}"
    if custom_instructions:
        style_guidance += f"\nAdditional instructions for this rewrite: {custom_instructions}"
    
    prompt = f"""You are a senior structural engineer with extensive experience in bridge inspection and rehabilitation.
Rewrite the following project description in a professional, technical tone appropriate for a federal SF330 proposal.
Focus on structural engineering aspects, bridge inspection methodologies, load ratings, condition assessments, and any rehabilitation or repair work.
Keep the same factual content but enhance the language to demonstrate technical expertise.
Keep the description concise (under 300 words) and suitable for Block 24 of SF330 Section F.
{style_guidance}

Original description:
{description}

Rewritten description (return ONLY the rewritten text, no explanations):"""
    
    try:
        from gemini_service import client
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt
        )
        rewritten = response.text.strip()
        return jsonify({'success': True, 'rewritten': rewritten})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/firms')
def firms():
    firms = Firm.query.order_by(Firm.name).all()
    return render_template('firms.html', firms=firms)


@app.route('/firms/scrape-website', methods=['POST'])
def scrape_firm_website():
    """Scrape a firm's website and extract company information using AI."""
    data = request.get_json()
    url = data.get('url', '').strip()
    
    if not url:
        return jsonify({'error': 'Please provide a website URL'}), 400
    
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url
    
    try:
        from web_scraper import get_firm_website_content
        from gemini_service import parse_firm_website
        
        website_content = get_firm_website_content(url)
        
        if not website_content:
            return jsonify({'error': 'Could not fetch content from this website. The site may be unavailable or blocking access.'}), 400
        
        firm_data = parse_firm_website(website_content)
        
        return jsonify({'success': True, 'data': firm_data})
    except Exception as e:
        return jsonify({'error': f'Error processing website: {str(e)}'}), 500


@app.route('/projects/scrape-website', methods=['POST'])
def scrape_projects_website():
    """Scrape a company portfolio website and extract project information using AI."""
    data = request.get_json()
    url = data.get('url', '').strip()
    
    if not url:
        return jsonify({'error': 'Please provide a website URL'}), 400
    
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url
    
    try:
        from web_scraper import scrape_portfolio_projects
        from gemini_service import parse_portfolio_projects
        
        scrape_result = scrape_portfolio_projects(url)
        
        if not scrape_result.get('content'):
            return jsonify({'error': 'Could not fetch content from this website. The site may be unavailable or blocking access.'}), 400
        
        parsed_data = parse_portfolio_projects(scrape_result['content'])
        projects = parsed_data.get('projects', [])
        
        return jsonify({
            'success': True,
            'projects': projects,
            'pages_scraped': scrape_result.get('pages_scraped', 0),
            'project_links': len(scrape_result.get('project_links', []))
        })
    except Exception as e:
        return jsonify({'error': f'Error processing website: {str(e)}'}), 500


@app.route('/projects/save-batch', methods=['POST'])
def save_batch_projects():
    """Save multiple projects at once (from website scrape or other sources)."""
    data = request.get_json()
    projects_data = data.get('projects', [])
    
    if not projects_data:
        return jsonify({'error': 'No projects provided'}), 400
    
    try:
        saved_count = 0
        for proj_data in projects_data:
            project = Project(
                title=proj_data.get('title') or 'Untitled Project',
                location=proj_data.get('location'),
                owner_name=proj_data.get('owner_name'),
                year_completed_professional=proj_data.get('year_completed_professional'),
                project_cost=proj_data.get('project_cost'),
                brief_description=proj_data.get('brief_description')
            )
            db.session.add(project)
            saved_count += 1
        
        db.session.commit()
        return jsonify({
            'success': True,
            'message': f'Successfully saved {saved_count} projects'
        })
    except Exception as e:
        db.session.rollback()
        return jsonify({'error': f'Error saving projects: {str(e)}'}), 500


@app.route('/firms/add', methods=['GET', 'POST'])
def add_firm():
    if request.method == 'POST':
        data = request.form
        firm = Firm(
            name=data.get('name', ''),
            uei=data.get('uei'),
            street_address=data.get('street_address'),
            city=data.get('city'),
            state=data.get('state'),
            zip_code=data.get('zip_code'),
            country=data.get('country', 'USA'),
            year_established=int(data.get('year_established')) if data.get('year_established') else None,
            ownership_type=data.get('ownership_type'),
            is_small_business=data.get('is_small_business') == 'on',
            small_business_categories=data.get('small_business_categories'),
            phone=data.get('phone'),
            fax=data.get('fax'),
            email=data.get('email'),
            point_of_contact_name=data.get('point_of_contact_name'),
            point_of_contact_title=data.get('point_of_contact_title')
        )
        db.session.add(firm)
        db.session.commit()
        return redirect(f'/firms/{firm.id}')
    
    return render_template('firm_add.html')


@app.route('/firms/<int:id>')
def firm_detail(id):
    firm = Firm.query.get_or_404(id)
    return render_template('firm_detail.html', firm=firm)


@app.route('/firms/<int:id>/edit', methods=['GET', 'POST'])
def edit_firm(id):
    firm = Firm.query.get_or_404(id)
    if request.method == 'POST':
        data = request.form
        firm.name = data.get('name', firm.name)
        firm.uei = data.get('uei')
        firm.street_address = data.get('street_address')
        firm.city = data.get('city')
        firm.state = data.get('state')
        firm.zip_code = data.get('zip_code')
        firm.country = data.get('country', 'USA')
        firm.year_established = int(data.get('year_established')) if data.get('year_established') else None
        firm.ownership_type = data.get('ownership_type')
        firm.is_small_business = data.get('is_small_business') == 'on'
        firm.small_business_categories = data.get('small_business_categories')
        firm.phone = data.get('phone')
        firm.fax = data.get('fax')
        firm.email = data.get('email')
        firm.point_of_contact_name = data.get('point_of_contact_name')
        firm.point_of_contact_title = data.get('point_of_contact_title')
        firm.bio = data.get('bio')
        firm.google_drive_folder_url = data.get('google_drive_folder_url')
        
        from datetime import datetime
        stat_fields = [
            ('stat_bridges_inspected', 'int'),
            ('stat_length_bridge_inspected', 'str'),
            ('stat_fcm_bridge_inspections', 'int'),
            ('stat_load_ratings_performed', 'int'),
            ('stat_critical_findings', 'int'),
            ('stat_timber_inspections', 'int')
        ]
        for field, field_type in stat_fields:
            new_val = data.get(field)
            old_val = getattr(firm, field)
            if field_type == 'int':
                new_val = int(new_val) if new_val else None
            else:
                new_val = new_val if new_val else None
            if new_val != old_val:
                setattr(firm, field, new_val)
                setattr(firm, f'{field}_updated', datetime.utcnow())
        
        db.session.commit()
        return redirect(f'/firms/{firm.id}')
    
    return render_template('firm_edit.html', firm=firm)


@app.route('/firms/<int:id>/delete', methods=['POST'])
def delete_firm(id):
    firm = Firm.query.get_or_404(id)
    if firm.employees:
        return jsonify({'error': 'Cannot delete firm with employees. Reassign or delete employees first.'}), 400
    if firm.proposals:
        return jsonify({'error': 'Cannot delete firm with proposals. Delete proposals first.'}), 400
    db.session.delete(firm)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/firms/<int:firm_id>/alternate-descriptions')
def get_firm_alternate_descriptions(firm_id):
    from models import FirmAlternateDescription
    alts = FirmAlternateDescription.query.filter_by(firm_id=firm_id).all()
    return jsonify([{
        'id': a.id,
        'label': a.label,
        'description': a.description
    } for a in alts])


@app.route('/api/firms/<int:firm_id>/alternate-descriptions', methods=['POST'])
def add_firm_alternate_description(firm_id):
    from models import FirmAlternateDescription
    Firm.query.get_or_404(firm_id)
    data = request.json
    alt = FirmAlternateDescription(
        firm_id=firm_id,
        label=data.get('label', 'Alternate'),
        description=data.get('description', '')
    )
    db.session.add(alt)
    db.session.commit()
    return jsonify({'success': True, 'id': alt.id})


@app.route('/api/firms/alternate-descriptions/<int:alt_id>', methods=['PUT'])
def update_firm_alternate_description(alt_id):
    from models import FirmAlternateDescription
    alt = FirmAlternateDescription.query.get_or_404(alt_id)
    data = request.json
    if 'label' in data:
        alt.label = data['label']
    if 'description' in data:
        alt.description = data['description']
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/firms/alternate-descriptions/<int:alt_id>', methods=['DELETE'])
def delete_firm_alternate_description(alt_id):
    from models import FirmAlternateDescription
    alt = FirmAlternateDescription.query.get_or_404(alt_id)
    db.session.delete(alt)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/proposals')
def proposals():
    search = request.args.get('search', '').strip()
    status_filter = request.args.get('status', '')
    
    query = Proposal.query
    
    if search:
        search_term = f'%{search}%'
        query = query.filter(
            db.or_(
                Proposal.tracking_number.ilike(search_term),
                Proposal.name.ilike(search_term),
                Proposal.contract_title.ilike(search_term),
                Proposal.solicitation_number.ilike(search_term)
            )
        )
    
    if status_filter:
        query = query.filter(Proposal.status == status_filter)
    
    proposals = query.order_by(Proposal.updated_at.desc()).all()
    return render_template('proposals.html', proposals=proposals, search=search, status_filter=status_filter)


@app.route('/proposals/parse-rfp', methods=['POST'])
def parse_rfp():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    try:
        from document_parser import extract_text_from_file
        from gemini_service import parse_rfp_rfq
        
        file_content = file.read()
        
        text = extract_text_from_file(file.filename, file_content)
        if not text:
            return jsonify({'error': 'Could not extract text from file'}), 400
        
        try:
            parsed_data = parse_rfp_rfq(text)
        except Exception as parse_err:
            print(f"RFP parsing error: {parse_err}")
            # Return empty parsed data with the raw text so user can still proceed
            parsed_data = {}
        
        if parsed_data is None:
            parsed_data = {}
            
        parsed_data['rfp_filename'] = file.filename
        parsed_data['rfp_content'] = file_content.hex()
        parsed_data['rfp_text'] = text[:50000]
        return jsonify({'success': True, 'data': parsed_data})
    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/proposals/new', methods=['GET', 'POST'])
def new_proposal():
    if request.method == 'GET':
        firms = Firm.query.all()
        from models import FirmAlternateDescription
        firm_alts = {}
        for firm in firms:
            alts = FirmAlternateDescription.query.filter_by(firm_id=firm.id).all()
            firm_alts[firm.id] = [{'id': a.id, 'label': a.label} for a in alts]
        return render_template('proposal_wizard_step1.html', firms=firms, firm_alts=firm_alts)
    
    data = request.form
    rfp_content = None
    rfp_filename = data.get('rfp_filename')
    
    # Handle file upload directly
    if 'rfp_file' in request.files:
        rfp_file = request.files['rfp_file']
        if rfp_file and rfp_file.filename:
            rfp_content = rfp_file.read()
            rfp_filename = rfp_file.filename
    
    proposal = Proposal(
        tracking_number=data.get('tracking_number'),
        name=data.get('name'),
        contract_title=data.get('contract_title'),
        contract_location=data.get('contract_location'),
        public_notice_date=data.get('public_notice_date'),
        solicitation_number=data.get('solicitation_number'),
        firm_id=data.get('firm_id') if data.get('firm_id') else None,
        firm_bio_alternate_id=data.get('firm_bio_alternate_id') if data.get('firm_bio_alternate_id') else None,
        rfp_filename=rfp_filename,
        rfp_content=rfp_content,
        rfp_text=data.get('rfp_text'),
        win_theme=data.get('win_theme')
    )
    db.session.add(proposal)
    db.session.commit()
    
    # Handle selected firm photos
    selected_photos = data.get('selected_firm_photos', '')
    if selected_photos and proposal.firm_id:
        from models import ProposalSelectedFirmPhoto, FirmPhoto
        photo_ids = [int(pid) for pid in selected_photos.split(',') if pid.strip()]
        for order, photo_id in enumerate(photo_ids):
            # Validate photo belongs to the selected firm
            photo = FirmPhoto.query.filter_by(id=photo_id, firm_id=proposal.firm_id).first()
            if photo:
                psfp = ProposalSelectedFirmPhoto(
                    proposal_id=proposal.id,
                    firm_photo_id=photo_id,
                    display_order=order
                )
                db.session.add(psfp)
    
    # Handle selected marketing photos
    selected_marketing = data.get('selected_marketing_photos', '')
    if selected_marketing:
        from models import ProposalSelectedMarketingPhoto, MarketingPhoto
        marketing_photo_ids = [int(pid) for pid in selected_marketing.split(',') if pid.strip()]
        for order, photo_id in enumerate(marketing_photo_ids):
            # Validate photo exists
            photo = MarketingPhoto.query.get(photo_id)
            if photo:
                psmp = ProposalSelectedMarketingPhoto(
                    proposal_id=proposal.id,
                    marketing_photo_id=photo_id,
                    display_order=order
                )
                db.session.add(psmp)
    
    # Handle reference document uploads (previous proposals)
    ref_count = int(data.get('ref_file_count', 0))
    for i in range(ref_count):
        ref_key = f'ref_file_{i}'
        if ref_key in request.files:
            ref_file = request.files[ref_key]
            if ref_file and ref_file.filename:
                file_content = ref_file.read()
                
                # Extract text from the reference document
                extracted_text = ""
                try:
                    extracted_text = extract_text_from_file(ref_file.filename, file_content)
                except Exception as e:
                    print(f"Error extracting text from reference doc: {e}")
                
                ref_doc = ProposalReference(
                    proposal_id=proposal.id,
                    filename=secure_filename(ref_file.filename),
                    file_content=file_content,
                    extracted_text=extracted_text,
                    file_size=len(file_content),
                    content_type=ref_file.content_type
                )
                db.session.add(ref_doc)
    
    # Handle intelligence document uploads
    from models import ProposalIntelligence
    intel_count = int(data.get('intel_file_count', 0))
    for i in range(intel_count):
        intel_key = f'intel_file_{i}'
        desc_key = f'intel_desc_{i}'
        if intel_key in request.files:
            intel_file = request.files[intel_key]
            if intel_file and intel_file.filename:
                file_content = intel_file.read()
                
                # Extract text from the intelligence document
                extracted_text = ""
                try:
                    extracted_text = extract_text_from_file(intel_file.filename, file_content)
                except Exception as e:
                    print(f"Error extracting text from intelligence doc: {e}")
                
                intel_doc = ProposalIntelligence(
                    proposal_id=proposal.id,
                    filename=secure_filename(intel_file.filename),
                    file_content=file_content,
                    extracted_text=extracted_text,
                    description=data.get(desc_key, ''),
                    file_size=len(file_content),
                    content_type=intel_file.content_type
                )
                db.session.add(intel_doc)
    
    db.session.commit()
    
    return redirect(url_for('proposal_step2', id=proposal.id))


@app.route('/proposals/<int:id>/step1', methods=['GET', 'POST'])
def proposal_step1_edit(id):
    proposal = Proposal.query.get_or_404(id)
    
    if request.method == 'GET':
        firms = Firm.query.all()
        from models import FirmAlternateDescription
        firm_alts = {}
        for firm in firms:
            alts = FirmAlternateDescription.query.filter_by(firm_id=firm.id).all()
            firm_alts[firm.id] = [{'id': a.id, 'label': a.label} for a in alts]
        return render_template('proposal_wizard_step1_edit.html', proposal=proposal, firms=firms, firm_alts=firm_alts)
    
    data = request.form
    
    proposal.tracking_number = data.get('tracking_number')
    proposal.name = data.get('name')
    proposal.contract_title = data.get('contract_title')
    proposal.contract_location = data.get('contract_location')
    proposal.public_notice_date = data.get('public_notice_date') or None
    proposal.solicitation_number = data.get('solicitation_number')
    proposal.firm_id = data.get('firm_id') if data.get('firm_id') else None
    proposal.firm_bio_alternate_id = data.get('firm_bio_alternate_id') if data.get('firm_bio_alternate_id') else None
    proposal.win_theme = data.get('win_theme')
    
    if 'rfp_file' in request.files:
        rfp_file = request.files['rfp_file']
        if rfp_file and rfp_file.filename:
            proposal.rfp_content = rfp_file.read()
            proposal.rfp_filename = rfp_file.filename
    
    if data.get('rfp_text'):
        proposal.rfp_text = data.get('rfp_text')
    
    from models import ProposalSelectedFirmPhoto, ProposalSelectedMarketingPhoto
    ProposalSelectedFirmPhoto.query.filter_by(proposal_id=id).delete()
    ProposalSelectedMarketingPhoto.query.filter_by(proposal_id=id).delete()
    
    selected_photos = data.get('selected_firm_photos', '')
    if selected_photos and proposal.firm_id:
        from models import FirmPhoto
        photo_ids = [int(pid) for pid in selected_photos.split(',') if pid.strip()]
        for order, photo_id in enumerate(photo_ids):
            photo = FirmPhoto.query.filter_by(id=photo_id, firm_id=proposal.firm_id).first()
            if photo:
                psfp = ProposalSelectedFirmPhoto(
                    proposal_id=proposal.id,
                    firm_photo_id=photo_id,
                    display_order=order
                )
                db.session.add(psfp)
    
    selected_marketing = data.get('selected_marketing_photos', '')
    if selected_marketing:
        from models import MarketingPhoto
        marketing_photo_ids = [int(pid) for pid in selected_marketing.split(',') if pid.strip()]
        for order, photo_id in enumerate(marketing_photo_ids):
            photo = MarketingPhoto.query.get(photo_id)
            if photo:
                psmp = ProposalSelectedMarketingPhoto(
                    proposal_id=proposal.id,
                    marketing_photo_id=photo_id,
                    display_order=order
                )
                db.session.add(psmp)
    
    db.session.commit()
    
    return redirect(url_for('proposal_step2', id=proposal.id))


@app.route('/api/proposals/with-orgcharts')
def api_proposals_with_orgcharts():
    proposals = Proposal.query.filter(Proposal.org_chart_data.isnot(None)).order_by(Proposal.created_at.desc()).all()
    return jsonify([{
        'id': p.id,
        'name': p.name,
        'tracking_number': p.tracking_number
    } for p in proposals])


@app.route('/api/proposals/<int:proposal_id>/orgchart-personnel')
def api_orgchart_personnel(proposal_id):
    proposal = Proposal.query.get_or_404(proposal_id)
    if not proposal.org_chart_data:
        return jsonify({'success': False, 'error': 'No org chart data'})
    
    try:
        org_data = json.loads(proposal.org_chart_data) if isinstance(proposal.org_chart_data, str) else proposal.org_chart_data
        nodes = org_data.get('nodes', [])
        
        personnel = []
        for node in nodes:
            node_data = node.get('data', {})
            staff_name = node_data.get('assignedStaff')
            role = node_data.get('role', '')
            staff_id = node_data.get('staffId')
            
            if staff_name and staff_id:
                personnel.append({
                    'employee_id': staff_id,
                    'name': staff_name,
                    'role': role
                })
        
        return jsonify({'success': True, 'personnel': personnel})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/proposals/<int:id>/step2', methods=['GET', 'POST'])
def proposal_step2(id):
    proposal = Proposal.query.get_or_404(id)
    
    if request.method == 'GET':
        employees = Employee.query.order_by(Employee.name).all()
        selected_ids = [se.employee_id for se in proposal.selected_employees]
        selected_roles = {se.employee_id: se.role_in_contract for se in proposal.selected_employees}
        proposals_with_orgcharts = Proposal.query.filter(Proposal.org_chart_data.isnot(None)).order_by(Proposal.created_at.desc()).all()
        return render_template('proposal_wizard_step2.html', 
                             proposal=proposal, 
                             employees=employees,
                             selected_ids=selected_ids,
                             selected_roles=selected_roles,
                             proposals_with_orgcharts=proposals_with_orgcharts)
    
    data = request.json
    employee_ids = data.get('employee_ids', [])
    roles = data.get('roles', {})
    
    ProposalSelectedEmployee.query.filter_by(proposal_id=id).delete()
    
    for idx, emp_id in enumerate(employee_ids):
        pse = ProposalSelectedEmployee(
            proposal_id=id,
            employee_id=emp_id,
            role_in_contract=roles.get(str(emp_id), ''),
            display_order=idx
        )
        db.session.add(pse)
    
    db.session.commit()
    return jsonify({'success': True, 'redirect': url_for('proposal_step3', id=id)})


@app.route('/proposals/<int:id>/step3', methods=['GET', 'POST'])
def proposal_step3(id):
    proposal = Proposal.query.get_or_404(id)
    
    if request.method == 'GET':
        projects = Project.query.order_by(Project.title).all()
        selected_ids = [sp.project_id for sp in proposal.selected_projects]
        selected_alt_descs = {sp.project_id: sp.alternate_description_id for sp in proposal.selected_projects if sp.alternate_description_id}
        return render_template('proposal_wizard_step3.html', 
                             proposal=proposal, 
                             projects=projects,
                             selected_ids=selected_ids,
                             selected_alt_descs=selected_alt_descs)
    
    data = request.json
    project_ids = data.get('project_ids', [])
    writeups = data.get('writeups', {})
    desc_versions = data.get('desc_versions', {})
    
    ProposalSelectedProject.query.filter_by(proposal_id=id).delete()
    
    for idx, proj_id in enumerate(project_ids[:10]):
        alt_desc_id = desc_versions.get(str(proj_id))
        if alt_desc_id:
            alt_desc = ProjectAlternateDescription.query.filter_by(id=alt_desc_id, project_id=proj_id).first()
            alt_desc_id = alt_desc.id if alt_desc else None
        psp = ProposalSelectedProject(
            proposal_id=id,
            project_id=proj_id,
            display_order=idx,
            custom_writeup=writeups.get(str(proj_id), ''),
            alternate_description_id=alt_desc_id
        )
        db.session.add(psp)
    
    db.session.commit()
    return jsonify({'success': True, 'redirect': url_for('proposal_step4', id=id)})


@app.route('/proposals/<int:id>/step4')
def proposal_step4(id):
    proposal = Proposal.query.get_or_404(id)
    
    selected_employees = ProposalSelectedEmployee.query.filter_by(proposal_id=id)\
        .order_by(ProposalSelectedEmployee.display_order).all()
    selected_projects = ProposalSelectedProject.query.filter_by(proposal_id=id)\
        .order_by(ProposalSelectedProject.display_order).all()
    
    matrix = {}
    for pse in selected_employees:
        matrix[pse.employee_id] = {}
        for psp in selected_projects:
            link = EmployeeProjectLink.query.filter_by(
                employee_id=pse.employee_id,
                project_id=psp.project_id
            ).first()
            matrix[pse.employee_id][psp.project_id] = link is not None
    
    return render_template('proposal_wizard_step4.html',
                         proposal=proposal,
                         selected_employees=selected_employees,
                         selected_projects=selected_projects,
                         matrix=matrix)


@app.route('/proposals/<int:id>/rfp-download')
def download_rfp(id):
    """Download the stored RFP file"""
    proposal = Proposal.query.get_or_404(id)
    if not proposal.rfp_content or not proposal.rfp_filename:
        return jsonify({'error': 'No RFP file stored'}), 404
    
    from flask import send_file
    from io import BytesIO
    
    file_ext = proposal.rfp_filename.rsplit('.', 1)[-1].lower() if '.' in proposal.rfp_filename else 'pdf'
    mime_types = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'txt': 'text/plain'
    }
    
    return send_file(
        BytesIO(proposal.rfp_content),
        mimetype=mime_types.get(file_ext, 'application/octet-stream'),
        as_attachment=True,
        download_name=proposal.rfp_filename
    )


@app.route('/proposals/<int:id>/reference/<int:ref_id>/download')
def download_reference(id, ref_id):
    """Download a reference document"""
    ref_doc = ProposalReference.query.filter_by(id=ref_id, proposal_id=id).first_or_404()
    
    from io import BytesIO
    
    file_ext = ref_doc.filename.rsplit('.', 1)[-1].lower() if '.' in ref_doc.filename else 'pdf'
    mime_types = {
        'pdf': 'application/pdf',
        'docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'doc': 'application/msword',
        'txt': 'text/plain'
    }
    
    return send_file(
        BytesIO(ref_doc.file_content),
        mimetype=mime_types.get(file_ext, 'application/octet-stream'),
        as_attachment=True,
        download_name=ref_doc.filename
    )


@app.route('/proposals/<int:id>/reference/<int:ref_id>/delete', methods=['POST'])
def delete_reference(id, ref_id):
    """Delete a reference document"""
    ref_doc = ProposalReference.query.filter_by(id=ref_id, proposal_id=id).first_or_404()
    db.session.delete(ref_doc)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/proposals/<int:id>/reference/upload', methods=['POST'])
def upload_reference(id):
    """Upload additional reference documents after proposal creation"""
    proposal = Proposal.query.get_or_404(id)
    
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    ref_file = request.files['file']
    if not ref_file or not ref_file.filename:
        return jsonify({'error': 'No file selected'}), 400
    
    file_content = ref_file.read()
    
    # Validate file type
    if not allowed_file(ref_file.filename):
        return jsonify({'error': 'File type not allowed. Only PDF and Word documents are accepted.'}), 400
    
    # Check file size (max 50MB)
    if len(file_content) > 50 * 1024 * 1024:
        return jsonify({'error': 'File too large. Maximum size is 50MB.'}), 400
    
    extracted_text = ""
    try:
        extracted_text = extract_text_from_file(ref_file.filename, file_content)
    except Exception as e:
        print(f"Error extracting text from reference doc: {e}")
    
    ref_doc = ProposalReference(
        proposal_id=proposal.id,
        filename=secure_filename(ref_file.filename),
        file_content=file_content,
        extracted_text=extracted_text,
        file_size=len(file_content),
        content_type=ref_file.content_type
    )
    db.session.add(ref_doc)
    db.session.commit()
    
    return jsonify({
        'success': True,
        'id': ref_doc.id,
        'filename': ref_doc.filename,
        'file_size': ref_doc.file_size
    })


@app.route('/proposals/<int:id>/generate-outline', methods=['POST'])
def generate_proposal_outline(id):
    """Generate AI proposal outline based on RFP and all linked data"""
    proposal = Proposal.query.get_or_404(id)
    data = request.json
    custom_instructions = data.get('custom_instructions', '')
    
    from models import FirmAlternateDescription, ProposalLinkedResponse, ProposalLinkedReference
    from gemini_service import generate_proposal_outline_ai
    
    firm_bio = ''
    if proposal.firm:
        firm_bio = proposal.firm.bio or ''
        if proposal.firm_bio_alternate_id:
            alt = FirmAlternateDescription.query.get(proposal.firm_bio_alternate_id)
            if alt:
                firm_bio = alt.description or firm_bio
    
    employees_data = []
    for pse in proposal.selected_employees:
        emp = pse.employee
        employees_data.append({
            'name': emp.name,
            'title': emp.title,
            'role_in_contract': pse.role_in_contract,
            'years_experience': emp.years_experience_total,
            'bio': emp.bio,
            'education': emp.education,
            'registrations': emp.registrations
        })
    
    projects_data = []
    for psp in proposal.selected_projects:
        proj = psp.project
        projects_data.append({
            'title': proj.title,
            'location': proj.location,
            'owner': proj.owner_name,
            'description': proj.brief_description
        })
    
    linked_responses = []
    for link in ProposalLinkedResponse.query.filter_by(proposal_id=id).all():
        resp = link.response
        linked_responses.append({
            'question': resp.question,
            'response': resp.response
        })
    
    linked_references = []
    for link in ProposalLinkedReference.query.filter_by(proposal_id=id).all():
        ref = link.reference
        linked_references.append({
            'project_name': ref.project_name,
            'final_score': ref.final_score,
            'client': ref.client
        })
    
    try:
        outline = generate_proposal_outline_ai(
            rfp_text=proposal.rfp_text or '',
            firm_name=proposal.firm.name if proposal.firm else '',
            firm_bio=firm_bio,
            employees=employees_data,
            projects=projects_data,
            contract_title=proposal.contract_title or '',
            solicitation_number=proposal.solicitation_number or '',
            custom_instructions=custom_instructions,
            org_chart_data=proposal.org_chart_data or '',
            org_chart_notes=proposal.org_chart_notes or '',
            linked_responses=linked_responses,
            linked_references=linked_references
        )
        
        proposal.proposal_outline = outline
        proposal.proposal_outline_instructions = custom_instructions
        db.session.commit()
        
        return jsonify({'success': True, 'outline': outline})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})


@app.route('/proposals/<int:id>/save-outline', methods=['POST'])
def save_proposal_outline(id):
    """Save edited proposal outline"""
    proposal = Proposal.query.get_or_404(id)
    data = request.json
    
    proposal.proposal_outline = data.get('outline', '')
    db.session.commit()
    
    return jsonify({'success': True})


@app.route('/proposals/<int:id>/generate-cover-letter', methods=['POST'])
def generate_cover_letter(id):
    """Generate AI cover letter using RFP + firm + staff + project data"""
    proposal = Proposal.query.get_or_404(id)
    data = request.json
    custom_instructions = data.get('custom_instructions', '')
    
    from models import AISettings, FirmAlternateDescription
    from gemini_service import generate_cover_letter_ai
    
    style = AISettings.get_value('writing_style', '')
    tone = AISettings.get_value('writing_tone', '')
    
    firm_bio = ''
    if proposal.firm:
        firm_bio = proposal.firm.bio or ''
        if proposal.firm_bio_alternate_id:
            alt = FirmAlternateDescription.query.get(proposal.firm_bio_alternate_id)
            if alt:
                firm_bio = alt.description or firm_bio
    
    employees_data = []
    for pse in proposal.selected_employees:
        emp = pse.employee
        employees_data.append({
            'name': emp.name,
            'title': emp.title,
            'role_in_contract': pse.role_in_contract,
            'years_experience': emp.years_experience_total,
            'education': emp.education,
            'registrations': emp.registrations
        })
    
    projects_data = []
    for psp in proposal.selected_projects:
        proj = psp.project
        projects_data.append({
            'title': proj.title,
            'location': proj.location,
            'owner': proj.owner_name,
            'description': proj.brief_description
        })
    
    # Gather reference proposal text from uploaded previous proposals
    reference_proposals_text = ""
    for ref_doc in proposal.reference_documents:
        if ref_doc.extracted_text:
            reference_proposals_text += f"\n\n--- Reference: {ref_doc.filename} ---\n{ref_doc.extracted_text}"
    
    result = generate_cover_letter_ai(
        rfp_text=proposal.rfp_text or '',
        firm_name=proposal.firm.name if proposal.firm else '',
        firm_bio=firm_bio,
        employees=employees_data,
        projects=projects_data,
        contract_title=proposal.contract_title or '',
        solicitation_number=proposal.solicitation_number or '',
        style=style,
        tone=tone,
        custom_instructions=custom_instructions,
        reference_proposals=reference_proposals_text,
        org_chart_data=proposal.org_chart_data or '',
        org_chart_notes=proposal.org_chart_notes or '',
        proposal_outline=proposal.proposal_outline or ''
    )
    
    proposal.cover_letter = result.get('cover_letter', '')
    proposal.written_sections = result.get('written_sections', '')
    db.session.commit()
    
    return jsonify({'success': True, 'cover_letter': proposal.cover_letter, 'written_sections': proposal.written_sections})


@app.route('/proposals/<int:id>/employee/<int:emp_id>/relevant-projects', methods=['GET', 'POST'])
def employee_relevant_projects(id, emp_id):
    proposal = Proposal.query.get_or_404(id)
    pse = ProposalSelectedEmployee.query.filter_by(proposal_id=id, employee_id=emp_id).first_or_404()
    
    if request.method == 'GET':
        employee = Employee.query.get(emp_id)
        selected_project_ids = [rp.project_id for rp in pse.relevant_projects if rp.project_id]
        selected_experience_ids = [rp.experience_id for rp in pse.relevant_projects if rp.experience_id]
        
        linked_project_ids = [link.project_id for link in EmployeeProjectLink.query.filter_by(employee_id=emp_id).all()]
        firm_projects = Project.query.filter(Project.id.in_(linked_project_ids)).order_by(Project.title).all() if linked_project_ids else []
        
        personnel_experiences = EmployeeProjectExperience.query.filter_by(employee_id=emp_id).order_by(EmployeeProjectExperience.project_title).all()
        
        return render_template('employee_relevant_projects.html',
                             proposal=proposal,
                             employee=employee,
                             firm_projects=firm_projects,
                             personnel_experiences=personnel_experiences,
                             selected_project_ids=selected_project_ids,
                             selected_experience_ids=selected_experience_ids)
    
    data = request.json
    project_ids = data.get('project_ids', [])
    experience_ids = data.get('experience_ids', [])
    
    total_items = project_ids[:5] if len(project_ids) >= 5 else project_ids
    remaining_slots = 5 - len(total_items)
    exp_to_add = experience_ids[:remaining_slots] if remaining_slots > 0 else []
    
    ProposalEmployeeRelevantProject.query.filter_by(proposal_selected_employee_id=pse.id).delete()
    
    display_order = 0
    for proj_id in total_items:
        perp = ProposalEmployeeRelevantProject(
            proposal_selected_employee_id=pse.id,
            project_id=proj_id,
            display_order=display_order
        )
        db.session.add(perp)
        display_order += 1
    
    for exp_id in exp_to_add:
        perp = ProposalEmployeeRelevantProject(
            proposal_selected_employee_id=pse.id,
            experience_id=exp_id,
            display_order=display_order
        )
        db.session.add(perp)
        display_order += 1
    
    db.session.commit()
    return jsonify({'success': True})


@app.route('/proposals/<int:id>/delete', methods=['POST'])
def delete_proposal(id):
    proposal = Proposal.query.get_or_404(id)
    
    ProposalSelectedEmployee.query.filter_by(proposal_id=id).delete()
    ProposalSelectedProject.query.filter_by(proposal_id=id).delete()
    ProposalSelectedFirmPhoto.query.filter_by(proposal_id=id).delete()
    ProposalSelectedMarketingPhoto.query.filter_by(proposal_id=id).delete()
    ProposalReference.query.filter_by(proposal_id=id).delete()
    ProposalIntelligence.query.filter_by(proposal_id=id).delete()
    
    db.session.delete(proposal)
    db.session.commit()
    
    flash(f'Proposal "{proposal.name}" has been deleted.', 'success')
    return redirect('/proposals')


@app.route('/proposals/<int:id>/generate-pdf')
def generate_proposal_pdf(id):
    proposal = Proposal.query.get_or_404(id)
    
    selected_employees = ProposalSelectedEmployee.query.filter_by(proposal_id=id)\
        .order_by(ProposalSelectedEmployee.display_order).all()
    selected_projects = ProposalSelectedProject.query.filter_by(proposal_id=id)\
        .order_by(ProposalSelectedProject.display_order).all()
    
    employee_project_matrix = {}
    for pse in selected_employees:
        employee_project_matrix[pse.employee_id] = set()
        for psp in selected_projects:
            link = EmployeeProjectLink.query.filter_by(
                employee_id=pse.employee_id,
                project_id=psp.project_id
            ).first()
            if link:
                employee_project_matrix[pse.employee_id].add(psp.project_id)
    
    proposal_data = {
        'contract_title': proposal.contract_title,
        'contract_location': proposal.contract_location,
        'public_notice_date': proposal.public_notice_date,
        'solicitation_number': proposal.solicitation_number,
        'firm': {
            'name': proposal.firm.name if proposal.firm else '',
            'street_address': proposal.firm.street_address if proposal.firm else '',
            'city': proposal.firm.city if proposal.firm else '',
            'state': proposal.firm.state if proposal.firm else '',
            'zip_code': proposal.firm.zip_code if proposal.firm else '',
            'year_established': proposal.firm.year_established if proposal.firm else '',
            'uei': proposal.firm.uei if proposal.firm else '',
            'ownership_type': proposal.firm.ownership_type if proposal.firm else '',
            'point_of_contact_name': proposal.firm.point_of_contact_name if proposal.firm else '',
            'phone': proposal.firm.phone if proposal.firm else '',
            'email': proposal.firm.email if proposal.firm else '',
        } if proposal.firm else {},
        'employees': [{
            'id': pse.employee_id,
            'name': pse.employee.name,
            'role_in_contract': pse.role_in_contract or pse.employee.role,
            'years_experience_total': pse.employee.years_experience_total,
            'years_experience_firm': pse.employee.years_experience_firm,
            'firm_name': pse.employee.firm.name if pse.employee.firm else '',
            'education': pse.employee.education,
            'registrations': pse.employee.registrations,
            'training': pse.employee.training,
            'other_qualifications': pse.employee.other_qualifications,
            'project_experiences': [{
                'project_title': exp.project_title,
                'location': exp.location,
                'owner_name': exp.owner_name,
                'project_cost': exp.project_cost,
                'year_completed': exp.year_completed,
                'role_performed': exp.role_performed,
                'brief_description': exp.brief_description,
                'firm_name': exp.firm_name,
            } for exp in pse.employee.project_experiences],
        } for pse in selected_employees],
        'projects': [{
            'id': psp.project_id,
            'title': psp.project.title,
            'location': psp.project.location,
            'year_completed_professional': psp.project.year_completed_professional,
            'year_completed_construction': psp.project.year_completed_construction,
            'owner_name': psp.project.owner_name,
            'owner_contact_name': psp.project.owner_contact_name,
            'owner_contact_phone': psp.project.owner_contact_phone,
            'brief_description': psp.alternate_description.description if psp.alternate_description else psp.project.brief_description,
            'custom_writeup': psp.custom_writeup or psp.project.relevance_writeup or '',
        } for psp in selected_projects],
        'employee_project_matrix': employee_project_matrix
    }
    
    try:
        pdf_bytes = generate_full_sf330(proposal_data)
        
        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"SF330_{proposal.name.replace(' ', '_')}.pdf"
        )
    except Exception as e:
        flash(f'Error generating PDF: {str(e)}', 'error')
        return redirect(url_for('proposal_step4', id=id))


@app.route('/proposals/<int:id>/finalize', methods=['POST'])
def finalize_proposal(id):
    proposal = Proposal.query.get_or_404(id)
    proposal.status = 'finalized'
    db.session.commit()
    return jsonify({'success': True})


@app.route('/proposals/<int:id>/generate-word')
def generate_proposal_word(id):
    """Generate SF330 as Word document"""
    from word_generator import generate_section_a_c, generate_section_e, generate_section_f, generate_section_g, generate_section_h, generate_part_ii, combine_documents_as_zip
    
    proposal = Proposal.query.get_or_404(id)
    
    selected_employees = ProposalSelectedEmployee.query.filter_by(proposal_id=id)\
        .order_by(ProposalSelectedEmployee.display_order).all()
    selected_projects = ProposalSelectedProject.query.filter_by(proposal_id=id)\
        .order_by(ProposalSelectedProject.display_order).all()
    
    employee_project_matrix = {}
    for pse in selected_employees:
        employee_project_matrix[pse.employee_id] = set()
        for psp in selected_projects:
            link = EmployeeProjectLink.query.filter_by(
                employee_id=pse.employee_id,
                project_id=psp.project_id
            ).first()
            if link:
                employee_project_matrix[pse.employee_id].add(psp.project_id)
    
    documents = {}
    
    firms_data = [{'firm': proposal.firm, 'role': 'Prime', 'is_prime': True}] if proposal.firm else []
    doc_a = generate_section_a_c(proposal, firms_data)
    buffer = io.BytesIO()
    doc_a.save(buffer)
    documents['section_a_c'] = buffer.getvalue()
    
    section_e_docs = []
    for pse in selected_employees:
        employee = pse.employee
        if not employee:
            continue
        numbered_exps = [e for e in (employee.project_experiences if employee else []) if e.resume_order is not None]
        numbered_exps.sort(key=lambda e: e.resume_order or 0)
        experiences = numbered_exps
        
        class EmpWrapper:
            pass
        emp = EmpWrapper()
        emp.name = employee.name
        emp.proposal_role = pse.role_in_contract or employee.role
        emp.years_experience = employee.years_experience_total
        emp.years_with_firm = employee.years_experience_firm
        emp.education = employee.education
        emp.registrations = employee.registrations
        emp.training = employee.training
        emp.other_qualifications = employee.other_qualifications
        
        class ExpWrapper:
            pass
        exp_list = []
        for e in experiences[:5]:
            exp = ExpWrapper()
            exp.title = e.project_title
            exp.location = e.location
            exp.year_completed = e.year_completed
            exp.description = e.active_description or e.brief_description
            exp.role = e.role_performed
            all_descs = []
            if e.brief_description:
                all_descs.append(('Main Description', e.brief_description))
            for alt in e.alternate_descriptions:
                if alt.description:
                    all_descs.append((alt.label, alt.description))
            exp.all_descriptions = all_descs
            exp_list.append(exp)
        
        doc_e = generate_section_e(emp, proposal, exp_list)
        buffer = io.BytesIO()
        doc_e.save(buffer)
        section_e_docs.append({'name': employee.name, 'data': buffer.getvalue()})
    
    documents['section_e'] = section_e_docs
    
    section_f_docs = []
    for idx, psp in enumerate(selected_projects):
        project = psp.project
        if not project:
            continue
        doc_f = generate_section_f(project, proposal, idx + 1)
        buffer = io.BytesIO()
        doc_f.save(buffer)
        section_f_docs.append({'name': project.title, 'data': buffer.getvalue()})
    documents['section_f'] = section_f_docs
    
    employees_with_roles = [
        {'employee': pse.employee, 'role': pse.role_in_contract or (pse.employee.role if pse.employee else '')}
        for pse in selected_employees if pse.employee
    ]
    projects = [psp.project for psp in selected_projects if psp.project]
    doc_g = generate_section_g(employees_with_roles, projects, employee_project_matrix)
    buffer = io.BytesIO()
    doc_g.save(buffer)
    documents['section_g'] = buffer.getvalue()
    
    doc_h = generate_section_h(proposal)
    buffer = io.BytesIO()
    doc_h.save(buffer)
    documents['section_h_i'] = buffer.getvalue()
    
    if proposal.firm:
        doc_ii = generate_part_ii(proposal.firm)
        buffer = io.BytesIO()
        doc_ii.save(buffer)
        documents['part_ii'] = buffer.getvalue()
    
    zip_data = combine_documents_as_zip(documents)
    
    filename = f"SF330_{proposal.tracking_number or proposal.id}_{proposal.name or 'proposal'}.zip"
    filename = filename.replace(' ', '_').replace('/', '-')
    
    return send_file(
        io.BytesIO(zip_data),
        mimetype='application/zip',
        as_attachment=True,
        download_name=filename
    )


@app.route('/proposals/<int:id>/generate-word-simple')
def generate_proposal_word_simple(id):
    """Generate SF330 as simple Word document (no template)"""
    from word_generator import generate_simple_sf330
    
    proposal = Proposal.query.get_or_404(id)
    
    selected_employees = ProposalSelectedEmployee.query.filter_by(proposal_id=id)\
        .order_by(ProposalSelectedEmployee.display_order).all()
    selected_projects = ProposalSelectedProject.query.filter_by(proposal_id=id)\
        .order_by(ProposalSelectedProject.display_order).all()
    
    # Build employee project matrix
    employee_project_matrix = {}
    for pse in selected_employees:
        if pse.employee:
            for link in EmployeeProjectLink.query.filter_by(employee_id=pse.employee_id).all():
                employee_project_matrix[(pse.employee_id, link.project_id)] = True
    
    # Build employees data with experiences
    employees_data = []
    for pse in selected_employees:
        if pse.employee:
            experiences = list(pse.employee.project_experiences) if pse.employee else []
            employees_data.append({
                'employee': pse.employee,
                'role': pse.role_in_contract or pse.employee.role,
                'experiences': experiences
            })
    
    # Build projects data with descriptions
    projects_data = []
    for psp in selected_projects:
        if psp.project:
            # Get alternate description if selected
            desc = psp.project.brief_description
            if psp.alternate_description_id:
                alt = ProjectAlternateDescription.query.get(psp.alternate_description_id)
                if alt:
                    desc = alt.description
            projects_data.append({
                'project': psp.project,
                'description': desc
            })
    
    # Build firms data
    firms_data = [{'firm': proposal.firm, 'role': 'Prime'}] if proposal.firm else []
    
    # Generate document
    doc_bytes = generate_simple_sf330(proposal, employees_data, projects_data, firms_data, employee_project_matrix)
    
    filename = f"SF330_{proposal.tracking_number or proposal.id}_{proposal.name or 'proposal'}_simple.docx"
    filename = filename.replace(' ', '_').replace('/', '-')
    
    return send_file(
        io.BytesIO(doc_bytes),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )


@app.route('/api/employees')
def api_employees():
    employees = Employee.query.order_by(Employee.name).all()
    return jsonify([{
        'id': e.id,
        'name': e.name,
        'title': e.title,
        'role': e.role
    } for e in employees])


@app.route('/api/projects')
def api_projects():
    projects = Project.query.order_by(Project.title).all()
    return jsonify([{
        'id': p.id,
        'title': p.title,
        'location': p.location,
        'owner_name': p.owner_name
    } for p in projects])


@app.route('/api/firms')
def api_firms():
    firms = Firm.query.order_by(Firm.name).all()
    return jsonify([{
        'id': f.id,
        'name': f.name,
        'city': f.city,
        'state': f.state
    } for f in firms])


@app.route('/employees/compare/<int:existing_id>')
def employee_compare(existing_id):
    existing = Employee.query.get_or_404(existing_id)
    
    parsed_new = session.get('pending_employee_data', {})
    
    if not parsed_new:
        new_data = request.args.get('new_data')
        if new_data:
            try:
                parsed_new = json.loads(new_data)
            except:
                parsed_new = {}
    
    existing_experiences = EmployeeProjectExperience.query.filter_by(employee_id=existing_id).all()
    
    existing_data = {
        'name': existing.name,
        'title': existing.title,
        'role': existing.role,
        'years_experience_total': existing.years_experience_total,
        'years_experience_firm': existing.years_experience_firm,
        'education': existing.education,
        'registrations': existing.registrations,
        'training': existing.training,
        'other_qualifications': existing.other_qualifications,
        'project_experience': [{
            'project_title': exp.project_title,
            'location': exp.location,
            'owner_name': exp.owner_name,
            'project_cost': exp.project_cost,
            'year_completed': exp.year_completed,
            'role_performed': exp.role_performed,
            'brief_description': exp.brief_description,
            'firm_name': exp.firm_name
        } for exp in existing_experiences]
    }
    
    return render_template('employee_compare.html', 
                          existing=existing,
                          existing_data=existing_data,
                          new_data=parsed_new)


@app.route('/employees/merge/<int:existing_id>', methods=['POST'])
def employee_merge(existing_id):
    existing = Employee.query.get_or_404(existing_id)
    data = request.json
    merged_data = data.get('merged_data', {})
    new_project_experiences = data.get('new_project_experiences', [])
    
    try:
        if merged_data.get('name'):
            existing.name = merged_data['name']
        if merged_data.get('title'):
            existing.title = merged_data['title']
        if merged_data.get('role'):
            existing.role = merged_data['role']
        if merged_data.get('years_experience_total') is not None:
            existing.years_experience_total = merged_data['years_experience_total']
        if merged_data.get('years_experience_firm') is not None:
            existing.years_experience_firm = merged_data['years_experience_firm']
        if merged_data.get('education'):
            existing.education = merged_data['education']
        if merged_data.get('registrations'):
            existing.registrations = merged_data['registrations']
        if merged_data.get('training'):
            existing.training = merged_data['training']
        if merged_data.get('other_qualifications'):
            existing.other_qualifications = merged_data['other_qualifications']
        
        for proj in new_project_experiences:
            if proj.get('project_title'):
                query = EmployeeProjectExperience.query.filter_by(
                    employee_id=existing_id,
                    project_title=proj.get('project_title')
                )
                if proj.get('owner_name'):
                    query = query.filter_by(owner_name=proj.get('owner_name'))
                if proj.get('firm_name'):
                    query = query.filter_by(firm_name=proj.get('firm_name'))
                existing_proj = query.first()
                
                if not existing_proj:
                    exp = EmployeeProjectExperience(
                        employee_id=existing_id,
                        project_title=proj.get('project_title'),
                        location=proj.get('location'),
                        owner_name=proj.get('owner_name'),
                        project_cost=proj.get('project_cost'),
                        year_completed=proj.get('year_completed'),
                        role_performed=proj.get('role_performed'),
                        brief_description=proj.get('brief_description'),
                        firm_name=proj.get('firm_name'),
                        is_current_firm=False
                    )
                    db.session.add(exp)
        
        db.session.commit()
        
        if 'pending_employee_data' in session:
            del session['pending_employee_data']
        
        return jsonify({'success': True, 'id': existing.id, 'message': 'Employee updated successfully'})
    except Exception as e:
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ai-combine-text', methods=['POST'])
def api_ai_combine_text():
    data = request.json
    text1 = data.get('text1', '')
    text2 = data.get('text2', '')
    field_name = data.get('field_name', 'description')
    
    try:
        combined = combine_and_rewrite_text(text1, text2, field_name)
        return jsonify({'success': True, 'combined_text': combined})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/ai-combine', methods=['POST'])
def api_ai_combine():
    data = request.json
    values = data.get('values', [])
    field_name = data.get('field_name', 'description')
    
    if len(values) < 2:
        return jsonify({'error': 'Need at least 2 values to combine'}), 400
    
    try:
        combined = values[0]
        for i in range(1, len(values)):
            combined = combine_and_rewrite_text(combined, values[i], field_name)
        return jsonify({'combined': combined})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/settings')
def settings():
    ai_style = AISettings.get_value('ai_writing_style', '')
    ai_tone = AISettings.get_value('ai_writing_tone', '')
    
    # Check if a custom template exists in object storage
    has_custom_template = False
    has_company_template = False
    try:
        client = get_storage_client()
        obj = client.download_as_bytes('templates/sf330_section_f_custom.docx')
        has_custom_template = obj is not None
    except:
        pass
    
    try:
        client = get_storage_client()
        obj = client.download_as_bytes('templates/company_template_custom.docx')
        has_company_template = obj is not None
    except:
        pass
    
    has_resume_template = False
    try:
        client = get_storage_client()
        obj = client.download_as_bytes('templates/resume_template_custom.docx')
        has_resume_template = obj is not None
    except:
        pass
    
    has_sf330_resume_template = False
    try:
        client = get_storage_client()
        obj = client.download_as_bytes('templates/sf330_resume_template_custom.docx')
        has_sf330_resume_template = obj is not None
    except:
        pass
    
    return render_template('settings.html', ai_style=ai_style, ai_tone=ai_tone, 
                           has_custom_template=has_custom_template, has_company_template=has_company_template,
                           has_resume_template=has_resume_template, has_sf330_resume_template=has_sf330_resume_template)


@app.route('/settings', methods=['POST'])
def save_settings():
    ai_style = request.form.get('ai_writing_style', '')
    ai_tone = request.form.get('ai_writing_tone', '')
    
    AISettings.set_value('ai_writing_style', ai_style)
    AISettings.set_value('ai_writing_tone', ai_tone)
    
    flash('AI settings saved successfully!', 'success')
    return redirect(url_for('settings'))


@app.route('/settings/template/export')
def export_template():
    """Download the SF330 Section F template"""
    
    # First check if there's a custom template in object storage
    try:
        client = get_storage_client()
        template_bytes = client.download_as_bytes('templates/sf330_section_f_custom.docx')
        if template_bytes:
            return send_file(
                io.BytesIO(template_bytes),
                as_attachment=True,
                download_name='SF330_Section_F_Template.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
    except:
        pass
    
    # Fall back to attached SF330 template file
    import os
    template_path = 'attached_assets/SF330_Section_F_Template_(1)-mwc_1770394097919.docx'
    if os.path.exists(template_path):
        return send_file(
            template_path,
            as_attachment=True,
            download_name='SF330_Section_F_Template.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    
    # Last resort: generate default template programmatically
    doc = create_default_sf330_template()
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name='SF330_Section_F_Template.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/settings/template/import', methods=['POST'])
def import_template():
    """Upload a custom SF330 Section F template"""
    if 'template' not in request.files:
        flash('No file uploaded', 'error')
        return redirect(url_for('settings'))
    
    file = request.files['template']
    if file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('settings'))
    
    if not file.filename.lower().endswith('.docx'):
        flash('Please upload a .docx file', 'error')
        return redirect(url_for('settings'))
    
    try:
        # Read the file content
        file_content = file.read()
        
        # Validate it's a valid Word document
        from docx import Document
        doc = Document(io.BytesIO(file_content))
        
        # Store in object storage
        client = get_storage_client()
        client.upload_from_bytes('templates/sf330_section_f_custom.docx', file_content)
        
        flash('Custom template uploaded successfully! It will be used for all future project downloads.', 'success')
    except Exception as e:
        flash(f'Error uploading template: {str(e)}', 'error')
    
    return redirect(url_for('settings'))


@app.route('/settings/template/reset', methods=['POST'])
def reset_template():
    """Remove custom template and reset to default"""
    try:
        client = get_storage_client()
        client.delete('templates/sf330_section_f_custom.docx')
        flash('Template reset to default successfully!', 'success')
    except Exception as e:
        flash(f'Error resetting template: {str(e)}', 'error')
    
    return redirect(url_for('settings'))


def create_default_sf330_template():
    """Create a default SF330 Section F template with placeholders"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    header_para = doc.add_paragraph()
    header_run = header_para.add_run('F. Example projects which best illustrate proposed team\'s qualifications for this contract:')
    header_run.bold = True
    
    doc.add_paragraph()
    
    p21 = doc.add_paragraph()
    p21.add_run('21. Title and location: ').bold = True
    p21.add_run('{{PROJECT_TITLE}} - {{PROJECT_LOCATION}}')
    
    p22a = doc.add_paragraph()
    p22a.add_run('22a. Year completed - Professional services: ').bold = True
    p22a.add_run('{{YEAR_COMPLETED_PROFESSIONAL}}')
    
    p22b = doc.add_paragraph()
    p22b.add_run('22b. Year completed - Construction: ').bold = True
    p22b.add_run('{{YEAR_COMPLETED_CONSTRUCTION}}')
    
    p23a = doc.add_paragraph()
    p23a.add_run('23a. Project owner: ').bold = True
    p23a.add_run('{{OWNER_NAME}}')
    
    p23b = doc.add_paragraph()
    p23b.add_run('23b. Point of contact name: ').bold = True
    p23b.add_run('{{OWNER_CONTACT}}')
    
    p23c = doc.add_paragraph()
    p23c.add_run('23c. Point of contact telephone number: ').bold = True
    p23c.add_run('{{OWNER_PHONE}}')
    
    p23d = doc.add_paragraph()
    p23d.add_run('23d. Point of contact email: ').bold = True
    p23d.add_run('{{OWNER_EMAIL}}')
    
    p24 = doc.add_paragraph()
    p24.add_run('24. Brief description and specific role: ').bold = True
    desc_para = doc.add_paragraph()
    desc_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    desc_para.add_run('{{BRIEF_DESCRIPTION}}')
    
    cost_p = doc.add_paragraph()
    cost_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cost_p.add_run('{{PROJECT_COST}}')
    
    team_p = doc.add_paragraph()
    team_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    team_p.add_run('{{KEY_PERSONNEL}}')
    
    p25 = doc.add_paragraph()
    p25.add_run('25. Firms involved: ').bold = True
    p25.add_run('{{FIRM_NAME}} - {{FIRM_CITY}}, {{FIRM_STATE}} - {{FIRM_ROLE}}')
    
    return doc


def create_default_company_template():
    """Create a default company template with placeholders"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    # Title
    title = doc.add_heading('{{PROJECT_TITLE}}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Location
    loc_para = doc.add_paragraph()
    loc_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    loc_para.add_run('{{PROJECT_LOCATION}}').italic = True
    
    doc.add_paragraph()
    
    # Project Information Section
    doc.add_heading('Project Information', level=1)
    
    info_table = doc.add_table(rows=7, cols=2)
    info_table.style = 'Table Grid'
    
    info_rows = [
        ('Project Cost:', '{{PROJECT_COST}}'),
        ('Year Completed (Professional):', '{{YEAR_COMPLETED_PROFESSIONAL}}'),
        ('Year Completed (Construction):', '{{YEAR_COMPLETED_CONSTRUCTION}}'),
        ('Delivery Method:', '{{DELIVERY_METHOD}}'),
        ('Firm:', '{{FIRM_NAME}}'),
        ('Firm Location:', '{{FIRM_CITY}}, {{FIRM_STATE}}'),
        ('Key Personnel:', '{{KEY_PERSONNEL}}'),
    ]
    
    for i, (label, value) in enumerate(info_rows):
        row = info_table.rows[i]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].text = value
    
    doc.add_paragraph()
    
    # Owner Information Section
    doc.add_heading('Owner/Client Information', level=1)
    
    owner_table = doc.add_table(rows=4, cols=2)
    owner_table.style = 'Table Grid'
    
    owner_rows = [
        ('Owner/Client:', '{{OWNER_NAME}}'),
        ('Contact Name:', '{{OWNER_CONTACT}}'),
        ('Phone:', '{{OWNER_PHONE}}'),
        ('Email:', '{{OWNER_EMAIL}}'),
    ]
    
    for i, (label, value) in enumerate(owner_rows):
        row = owner_table.rows[i]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].text = value
    
    doc.add_paragraph()
    
    # Description Section
    doc.add_heading('Project Description', level=1)
    doc.add_paragraph('{{BRIEF_DESCRIPTION}}')
    
    cost_para = doc.add_paragraph()
    cost_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cost_para.add_run('{{PROJECT_COST}}')
    
    team_para = doc.add_paragraph()
    team_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    team_para.add_run('{{KEY_PERSONNEL}}')
    
    doc.add_paragraph()
    
    # Photos placeholder
    photos_para = doc.add_paragraph()
    photos_para.add_run('[Project photos will be inserted here if available]').italic = True
    photos_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return doc


def create_default_resume_template():
    """Create a default resume template with placeholders"""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = Document()
    
    header = doc.add_heading('{{EMPLOYEE_NAME}}', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('{{EMPLOYEE_TITLE}}')
    run.italic = True
    
    doc.add_paragraph()
    
    doc.add_heading('Personnel Information', level=1)
    
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Table Grid'
    
    info_rows = [
        ('Role:', '{{EMPLOYEE_ROLE}}'),
        ('Firm:', '{{FIRM_NAME}}'),
        ('Years Experience (Total):', '{{YEARS_EXPERIENCE_TOTAL}}'),
        ('Years Experience (Firm):', '{{YEARS_EXPERIENCE_FIRM}}'),
        ('Firm Location:', '{{FIRM_CITY}}, {{FIRM_STATE}}'),
    ]
    
    for i, (label, value) in enumerate(info_rows):
        row = info_table.rows[i]
        row.cells[0].paragraphs[0].add_run(label).bold = True
        row.cells[1].text = value
    
    doc.add_paragraph()
    
    doc.add_heading('Education', level=1)
    doc.add_paragraph('{{EDUCATION}}')
    
    doc.add_heading('Registrations / Certifications', level=1)
    doc.add_paragraph('{{REGISTRATIONS}}')
    
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph('{{BIO}}')
    
    doc.add_heading('Training', level=1)
    doc.add_paragraph('{{TRAINING}}')
    
    doc.add_heading('Other Qualifications', level=1)
    doc.add_paragraph('{{OTHER_QUALIFICATIONS}}')
    
    doc.add_heading('Project Experience', level=1)
    for i in range(1, 11):
        doc.add_paragraph(f'{{{{PROJECT_EXPERIENCE_{i}_TITLE}}}} - {{{{PROJECT_EXPERIENCE_{i}_ROLE}}}}')
        doc.add_paragraph(f'{{{{PROJECT_EXPERIENCE_{i}_LOCATION}}}} ({{{{PROJECT_EXPERIENCE_{i}_YEAR}}}})')
        doc.add_paragraph(f'{{{{PROJECT_EXPERIENCE_{i}_DESCRIPTION}}}}')
        doc.add_paragraph('')
    
    return doc


@app.route('/settings/company-template/export')
def export_company_template():
    """Download the company template"""
    
    # First check if there's a custom company template in object storage
    try:
        client = get_storage_client()
        template_bytes = client.download_as_bytes('templates/company_template_custom.docx')
        if template_bytes:
            return send_file(
                io.BytesIO(template_bytes),
                as_attachment=True,
                download_name='Company_Project_Template.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
    except:
        pass
    
    # Create default template
    doc = create_default_company_template()
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return send_file(
        buffer,
        as_attachment=True,
        download_name='Company_Project_Template.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/settings/company-template/import', methods=['POST'])
def import_company_template():
    """Upload a custom company template"""
    if 'template' not in request.files:
        flash('No file uploaded', 'error')
        return redirect(url_for('settings'))
    
    file = request.files['template']
    if file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('settings'))
    
    if not file.filename.lower().endswith('.docx'):
        flash('Please upload a .docx file', 'error')
        return redirect(url_for('settings'))
    
    try:
        # Read the file content
        file_content = file.read()
        
        # Validate it's a valid Word document
        from docx import Document
        doc = Document(io.BytesIO(file_content))
        
        # Store in object storage
        client = get_storage_client()
        client.upload_from_bytes('templates/company_template_custom.docx', file_content)
        
        flash('Custom company template uploaded successfully!', 'success')
    except Exception as e:
        flash(f'Error uploading template: {str(e)}', 'error')
    
    return redirect(url_for('settings'))


@app.route('/settings/company-template/reset', methods=['POST'])
def reset_company_template():
    """Remove custom company template and reset to default"""
    try:
        client = get_storage_client()
        client.delete('templates/company_template_custom.docx')
        flash('Company template reset to default successfully!', 'success')
    except Exception as e:
        flash(f'Error resetting template: {str(e)}', 'error')
    
    return redirect(url_for('settings'))


@app.route('/settings/resume-template/export')
def export_resume_template():
    """Download the resume template"""
    try:
        client = get_storage_client()
        template_bytes = client.download_as_bytes('templates/resume_template_custom.docx')
        if template_bytes:
            return send_file(
                io.BytesIO(template_bytes),
                as_attachment=True,
                download_name='Resume_Template.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
    except:
        pass
    
    doc = create_default_resume_template()
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(
        buffer,
        as_attachment=True,
        download_name='Resume_Template.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/settings/resume-template/import', methods=['POST'])
def import_resume_template():
    """Upload a custom resume template"""
    if 'template' not in request.files:
        flash('No file uploaded', 'error')
        return redirect(url_for('settings'))
    
    file = request.files['template']
    if file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('settings'))
    
    if not file.filename.lower().endswith('.docx'):
        flash('Please upload a .docx file', 'error')
        return redirect(url_for('settings'))
    
    try:
        file_content = file.read()
        from docx import Document
        doc = Document(io.BytesIO(file_content))
        
        client = get_storage_client()
        client.upload_from_bytes('templates/resume_template_custom.docx', file_content)
        
        flash('Custom resume template uploaded successfully!', 'success')
    except Exception as e:
        flash(f'Error uploading template: {str(e)}', 'error')
    
    return redirect(url_for('settings'))


@app.route('/settings/resume-template/reset', methods=['POST'])
def reset_resume_template():
    """Remove custom resume template and reset to default"""
    try:
        client = get_storage_client()
        client.delete('templates/resume_template_custom.docx')
        flash('Resume template reset to default successfully!', 'success')
    except Exception as e:
        flash(f'Error resetting template: {str(e)}', 'error')
    
    return redirect(url_for('settings'))


@app.route('/settings/sf330-resume-template/export')
def export_sf330_resume_template():
    """Download the SF330 Section E resume template"""
    try:
        client = get_storage_client()
        template_bytes = client.download_as_bytes('templates/sf330_resume_template_custom.docx')
        if template_bytes:
            return send_file(
                io.BytesIO(template_bytes),
                as_attachment=True,
                download_name='SF330_Section_E_Resume_Template.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
    except:
        pass
    
    default_path = os.path.join(os.path.dirname(__file__), 'attached_assets', 'sf330_section_e_template.docx')
    if not os.path.exists(default_path):
        default_path = os.path.join(os.path.dirname(__file__), 'attached_assets', '330_Section_E_Standards_template_1770398969209.docx')
    
    try:
        with open(default_path, 'rb') as f:
            return send_file(
                io.BytesIO(f.read()),
                as_attachment=True,
                download_name='SF330_Section_E_Resume_Template.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
    except Exception as e:
        flash(f'Error exporting SF330 resume template: {str(e)}', 'error')
        return redirect(url_for('settings'))


@app.route('/settings/sf330-resume-template/import', methods=['POST'])
def import_sf330_resume_template():
    """Upload a custom SF330 Section E resume template"""
    if 'template' not in request.files:
        flash('No file uploaded', 'error')
        return redirect(url_for('settings'))
    
    file = request.files['template']
    if file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('settings'))
    
    if not file.filename.lower().endswith('.docx'):
        flash('Please upload a .docx file', 'error')
        return redirect(url_for('settings'))
    
    try:
        file_content = file.read()
        from docx import Document
        doc = Document(io.BytesIO(file_content))
        
        client = get_storage_client()
        client.upload_from_bytes('templates/sf330_resume_template_custom.docx', file_content)
        
        flash('Custom SF330 resume template uploaded successfully!', 'success')
    except Exception as e:
        flash(f'Error uploading template: {str(e)}', 'error')
    
    return redirect(url_for('settings'))


@app.route('/settings/sf330-resume-template/reset', methods=['POST'])
def reset_sf330_resume_template():
    """Remove custom SF330 resume template and reset to default"""
    try:
        client = get_storage_client()
        client.delete('templates/sf330_resume_template_custom.docx')
        flash('SF330 resume template reset to default successfully!', 'success')
    except Exception as e:
        flash(f'Error resetting template: {str(e)}', 'error')
    
    return redirect(url_for('settings'))


@app.route('/data/export')
def export_data():
    from data_export import export_all_data
    from datetime import datetime
    
    data = export_all_data()
    json_str = json.dumps(data, indent=2, ensure_ascii=False)
    
    filename = f"sf330_data_export_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
    
    return send_file(
        io.BytesIO(json_str.encode('utf-8')),
        mimetype='application/json',
        as_attachment=True,
        download_name=filename
    )


@app.route('/data/import', methods=['POST'])
def import_data():
    from data_export import import_all_data
    
    if 'file' not in request.files:
        flash('No file uploaded', 'error')
        return redirect(url_for('settings'))
    
    file = request.files['file']
    if file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('settings'))
    
    if not file.filename.endswith('.json'):
        flash('Please upload a JSON file', 'error')
        return redirect(url_for('settings'))
    
    try:
        content = file.read().decode('utf-8')
        data = json.loads(content)
        
        clear_existing = request.form.get('clear_existing') == 'on'
        results = import_all_data(data, clear_existing=clear_existing)
        
        if results['success']:
            summary = ', '.join([f"{k}: {v}" for k, v in results['imported'].items() if v > 0])
            flash(f'Data imported successfully! {summary}', 'success')
        else:
            flash(f'Import failed: {", ".join(results["errors"])}', 'error')
    except json.JSONDecodeError:
        flash('Invalid JSON file', 'error')
    except Exception as e:
        flash(f'Import error: {str(e)}', 'error')
    
    return redirect(url_for('settings'))


@app.route('/contacts')
def contacts():
    all_contacts = ClientContact.query.order_by(ClientContact.name).all()
    return render_template('contacts.html', contacts=all_contacts)


@app.route('/contacts/add', methods=['GET', 'POST'])
def add_contact():
    if request.method == 'GET':
        return render_template('contact_form.html', contact=None)
    
    contact = ClientContact(
        name=request.form.get('name'),
        agency=request.form.get('agency'),
        role=request.form.get('role'),
        phone=request.form.get('phone'),
        email=request.form.get('email'),
        physical_street=request.form.get('physical_street'),
        physical_city=request.form.get('physical_city'),
        physical_state=request.form.get('physical_state'),
        physical_zip=request.form.get('physical_zip'),
        mailing_street=request.form.get('mailing_street'),
        mailing_city=request.form.get('mailing_city'),
        mailing_state=request.form.get('mailing_state'),
        mailing_zip=request.form.get('mailing_zip')
    )
    db.session.add(contact)
    db.session.commit()
    flash('Contact added successfully!', 'success')
    return redirect(url_for('contacts'))


@app.route('/contacts/<int:id>', methods=['GET', 'POST'])
def edit_contact(id):
    contact = ClientContact.query.get_or_404(id)
    
    if request.method == 'GET':
        return render_template('contact_form.html', contact=contact)
    
    contact.name = request.form.get('name')
    contact.agency = request.form.get('agency')
    contact.role = request.form.get('role')
    contact.phone = request.form.get('phone')
    contact.email = request.form.get('email')
    contact.physical_street = request.form.get('physical_street')
    contact.physical_city = request.form.get('physical_city')
    contact.physical_state = request.form.get('physical_state')
    contact.physical_zip = request.form.get('physical_zip')
    contact.mailing_street = request.form.get('mailing_street')
    contact.mailing_city = request.form.get('mailing_city')
    contact.mailing_state = request.form.get('mailing_state')
    contact.mailing_zip = request.form.get('mailing_zip')
    
    db.session.commit()
    flash('Contact updated successfully!', 'success')
    return redirect(url_for('contacts'))


@app.route('/contacts/<int:id>/delete', methods=['POST'])
def delete_contact(id):
    contact = ClientContact.query.get_or_404(id)
    db.session.delete(contact)
    db.session.commit()
    flash('Contact deleted successfully!', 'success')
    return redirect(url_for('contacts'))


@app.route('/contacts/merge', methods=['GET'])
def merge_contacts_page():
    ids = request.args.get('ids', '')
    id_list = [int(x) for x in ids.split(',') if x.isdigit()]
    if len(id_list) < 2:
        return redirect('/contacts')
    
    contacts = ClientContact.query.filter(ClientContact.id.in_(id_list)).all()
    if len(contacts) < 2:
        return redirect('/contacts')
    
    contacts_data = [{
        'id': c.id,
        'name': c.name,
        'agency': c.agency,
        'role': c.role,
        'phone': c.phone,
        'email': c.email,
        'physical_street': c.physical_street,
        'physical_city': c.physical_city,
        'physical_state': c.physical_state,
        'physical_zip': c.physical_zip,
        'mailing_street': c.mailing_street,
        'mailing_city': c.mailing_city,
        'mailing_state': c.mailing_state,
        'mailing_zip': c.mailing_zip
    } for c in contacts]
    
    return render_template('contact_merge.html', contacts=contacts, contacts_json=contacts_data)


@app.route('/contacts/merge', methods=['POST'])
def merge_contacts():
    data = request.json
    primary_id = data.get('primary_id')
    merge_ids = data.get('merge_ids', [])
    merged_data = data.get('merged_data', {})
    
    primary = ClientContact.query.get_or_404(primary_id)
    
    for key, value in merged_data.items():
        if hasattr(primary, key):
            setattr(primary, key, value if value else None)
    
    for merge_id in merge_ids:
        if merge_id == primary_id:
            continue
        merge_contact = ClientContact.query.get(merge_id)
        if merge_contact:
            db.session.delete(merge_contact)
    
    db.session.commit()
    return jsonify({'success': True, 'redirect': f'/contacts/{primary_id}'})


@app.route('/api/contacts/search')
def search_contacts():
    query = request.args.get('q', '')
    if len(query) < 2:
        return jsonify([])
    
    contacts = ClientContact.query.filter(
        ClientContact.name.ilike(f'%{query}%')
    ).order_by(ClientContact.name).limit(10).all()
    
    return jsonify([{
        'id': c.id,
        'name': c.name,
        'agency': c.agency or '',
        'role': c.role or '',
        'phone': c.phone or '',
        'email': c.email or ''
    } for c in contacts])


# ==================== CERTIFICATIONS ====================

@app.route('/certifications')
def certifications():
    from datetime import date
    employees = Employee.query.order_by(Employee.name).all()
    return render_template('certifications.html', employees=employees, today=date.today())


@app.route('/certifications/<int:employee_id>')
def employee_certifications(employee_id):
    from datetime import date
    employee = Employee.query.get_or_404(employee_id)
    certs = Certification.query.filter_by(employee_id=employee_id).order_by(Certification.category, Certification.name).all()
    return render_template('employee_certifications.html', employee=employee, certifications=certs, today=date.today())


@app.route('/certifications/<int:employee_id>/add', methods=['POST'])
def add_certification(employee_id):
    employee = Employee.query.get_or_404(employee_id)
    
    expiration_date = None
    exp_str = request.form.get('expiration_date')
    if exp_str:
        try:
            from datetime import datetime as dt
            expiration_date = dt.strptime(exp_str, '%Y-%m-%d').date()
        except ValueError:
            pass
    
    cert = Certification(
        employee_id=employee_id,
        cert_type=request.form.get('cert_type', 'certification'),
        category=request.form.get('category'),
        name=request.form.get('name'),
        state=request.form.get('state'),
        level=request.form.get('level'),
        status=request.form.get('status'),
        expiration_date=expiration_date,
        license_number=request.form.get('license_number'),
        notes=request.form.get('notes')
    )
    
    pdf_file = request.files.get('pdf_file')
    if pdf_file and pdf_file.filename:
        cert.pdf_filename = secure_filename(pdf_file.filename)
        cert.pdf_content = pdf_file.read()
    
    db.session.add(cert)
    db.session.commit()
    flash('Certification added successfully!', 'success')
    return redirect(url_for('employee_certifications', employee_id=employee_id))


@app.route('/certifications/<int:cert_id>/update', methods=['POST'])
def update_certification(cert_id):
    cert = Certification.query.get_or_404(cert_id)
    
    cert.cert_type = request.form.get('cert_type', cert.cert_type)
    cert.category = request.form.get('category', cert.category)
    cert.name = request.form.get('name', cert.name)
    cert.state = request.form.get('state', cert.state)
    cert.level = request.form.get('level', cert.level)
    cert.status = request.form.get('status', cert.status)
    cert.license_number = request.form.get('license_number', cert.license_number)
    cert.notes = request.form.get('notes', cert.notes)
    
    exp_str = request.form.get('expiration_date')
    if exp_str:
        try:
            from datetime import datetime as dt
            cert.expiration_date = dt.strptime(exp_str, '%Y-%m-%d').date()
        except ValueError:
            pass
    elif 'expiration_date' in request.form:
        cert.expiration_date = None
    
    db.session.commit()
    flash('Certification updated successfully!', 'success')
    return redirect(url_for('employee_certifications', employee_id=cert.employee_id))


@app.route('/certifications/<int:cert_id>/delete', methods=['POST'])
def delete_certification(cert_id):
    cert = Certification.query.get_or_404(cert_id)
    employee_id = cert.employee_id
    db.session.delete(cert)
    db.session.commit()
    flash('Certification deleted successfully!', 'success')
    return redirect(url_for('employee_certifications', employee_id=employee_id))


@app.route('/certifications/<int:cert_id>/pdf', methods=['GET'])
def view_certification_pdf(cert_id):
    cert = Certification.query.get_or_404(cert_id)
    if not cert.pdf_content:
        flash('No PDF file attached to this certification.', 'warning')
        return redirect(url_for('employee_certifications', employee_id=cert.employee_id))
    
    return send_file(
        io.BytesIO(cert.pdf_content),
        mimetype='application/pdf',
        as_attachment=False,
        download_name=cert.pdf_filename or f'{cert.name}.pdf'
    )


@app.route('/certifications/<int:cert_id>/pdf/upload', methods=['POST'])
def upload_certification_pdf(cert_id):
    cert = Certification.query.get_or_404(cert_id)
    
    pdf_file = request.files.get('pdf_file')
    if pdf_file and pdf_file.filename:
        cert.pdf_filename = secure_filename(pdf_file.filename)
        cert.pdf_content = pdf_file.read()
        db.session.commit()
        flash('PDF uploaded successfully!', 'success')
    else:
        flash('No file selected.', 'warning')
    
    return redirect(url_for('employee_certifications', employee_id=cert.employee_id))


@app.route('/certifications/<int:cert_id>/pdf/delete', methods=['POST'])
def delete_certification_pdf(cert_id):
    cert = Certification.query.get_or_404(cert_id)
    cert.pdf_filename = None
    cert.pdf_content = None
    db.session.commit()
    flash('PDF deleted successfully!', 'success')
    return redirect(url_for('employee_certifications', employee_id=cert.employee_id))


@app.route('/certifications/import-csv', methods=['POST'])
def import_certifications_csv():
    import csv
    from datetime import datetime as dt
    
    file = request.files.get('csv_file')
    if not file:
        flash('No file uploaded.', 'error')
        return redirect(url_for('certifications'))
    
    content = file.read().decode('utf-8-sig')
    lines = content.strip().split('\n')
    reader = csv.reader(lines)
    rows = list(reader)
    
    if len(rows) < 3:
        flash('Invalid CSV format.', 'error')
        return redirect(url_for('certifications'))
    
    # Row 2 has names starting at column 4 (index 4)
    header_row = rows[1]
    name_columns = {}
    for idx, cell in enumerate(header_row):
        if idx >= 4 and cell.strip():
            # Convert "Last, First" to "First Last"
            parts = cell.strip().split(',')
            if len(parts) == 2:
                name = f"{parts[1].strip()} {parts[0].strip()}"
            else:
                name = cell.strip()
            name_columns[idx] = name
    
    # Create or find employees
    employee_map = {}
    for idx, name in name_columns.items():
        emp = Employee.query.filter(Employee.name.ilike(name)).first()
        if not emp:
            emp = Employee(name=name)
            db.session.add(emp)
            db.session.flush()
        employee_map[idx] = emp
    
    # Parse certifications
    current_category = None
    
    for row_idx, row in enumerate(rows[2:], start=2):
        if len(row) < 2:
            continue
        
        col0 = row[0].strip() if len(row) > 0 else ''
        col1 = row[1].strip() if len(row) > 1 else ''
        
        # Detect category headers
        if col0 and not col1:
            if 'National' in col0 or 'Highway' in col0 or 'Institute' in col0:
                current_category = 'NHI'
            elif 'Safety' in col0:
                current_category = 'Safety'
            elif 'SPRAT' in col0:
                current_category = 'SPRAT'
            elif 'Drone' in col0:
                current_category = 'Drone'
            elif 'PE' in col0 or 'COA' in col0:
                current_category = 'PE License'
            continue
        
        # Certification name in column 1
        cert_name = col1 if col1 else col0
        if not cert_name or cert_name in ['registered', 'N/A', 'Able to pursue']:
            continue
        
        # Check for PE state rows (state in column 0)
        state = None
        if current_category == 'PE License' and col0 and col0 not in ['registered', 'N/A']:
            state = col0
            cert_name = 'PE'
        
        # Process each employee column
        for col_idx, emp in employee_map.items():
            if col_idx >= len(row):
                continue
            
            value = row[col_idx].strip() if row[col_idx] else ''
            if not value or value == '----':
                continue
            
            # Determine status and expiration
            status = None
            expiration_date = None
            level = None
            
            if value.lower() in ['yes', 'no']:
                status = value.lower()
            elif '/' in value:
                # Date format M/D/YYYY
                try:
                    expiration_date = dt.strptime(value, '%m/%d/%Y').date()
                    status = 'active' if expiration_date > dt.now().date() else 'expired'
                except ValueError:
                    status = value
            elif value.lower() == 'registered':
                status = 'registered'
            elif value.isdigit():
                level = value
                status = 'active'
            else:
                status = value
            
            # Skip 'no' values
            if status == 'no':
                continue
            
            # Check if certification already exists
            existing = Certification.query.filter_by(
                employee_id=emp.id,
                name=cert_name,
                state=state
            ).first()
            
            if not existing:
                cert = Certification(
                    employee_id=emp.id,
                    cert_type='license' if current_category == 'PE License' else 'training' if current_category in ['NHI', 'Safety'] else 'certification',
                    category=current_category,
                    name=cert_name,
                    state=state,
                    level=level,
                    status=status,
                    expiration_date=expiration_date
                )
                db.session.add(cert)
    
    db.session.commit()
    flash(f'CSV imported successfully! Created {len(employee_map)} personnel records with certifications.', 'success')
    return redirect(url_for('certifications'))


@app.route('/certification-types/seed', methods=['POST'])
def seed_certification_types():
    """Seed CertificationType table from existing Certification records"""
    # Get unique name/category combinations from existing certifications
    existing_certs = db.session.query(
        Certification.name, 
        Certification.category, 
        Certification.cert_type
    ).distinct().all()
    
    added = 0
    for cert in existing_certs:
        if not cert.name:
            continue
        # Check if already exists
        existing = CertificationType.query.filter_by(
            name=cert.name,
            category=cert.category
        ).first()
        if not existing:
            cert_type = CertificationType(
                name=cert.name,
                category=cert.category or 'Other',
                cert_type=cert.cert_type or 'certification',
                has_levels=(cert.category == 'SPRAT'),
                has_expiration=(cert.category != 'NHI')  # NHI training doesn't typically expire
            )
            db.session.add(cert_type)
            added += 1
    
    db.session.commit()
    flash(f'Seeded {added} certification types from existing records.', 'success')
    return redirect(url_for('certifications'))


@app.route('/api/certification-types')
def get_certification_types():
    """Get all certification types grouped by category for the checklist"""
    types = CertificationType.query.order_by(CertificationType.category, CertificationType.sort_order, CertificationType.name).all()
    
    grouped = {}
    for ct in types:
        category = ct.category or 'Other'
        if category not in grouped:
            grouped[category] = []
        grouped[category].append({
            'id': ct.id,
            'name': ct.name,
            'category': ct.category,
            'cert_type': ct.cert_type,
            'has_levels': ct.has_levels,
            'has_expiration': ct.has_expiration
        })
    
    return jsonify(grouped)


@app.route('/certification-types/add', methods=['POST'])
def add_certification_type():
    """Add a new certification type to the master list"""
    data = request.json
    
    # Check if already exists
    existing = CertificationType.query.filter_by(
        name=data.get('name'),
        category=data.get('category')
    ).first()
    
    if existing:
        return jsonify({'success': False, 'message': 'Certification type already exists', 'id': existing.id})
    
    cert_type = CertificationType(
        name=data.get('name'),
        category=data.get('category', 'Other'),
        cert_type=data.get('cert_type', 'certification'),
        has_levels=data.get('has_levels', False),
        has_expiration=data.get('has_expiration', True)
    )
    db.session.add(cert_type)
    db.session.commit()
    
    return jsonify({
        'success': True,
        'id': cert_type.id,
        'name': cert_type.name,
        'category': cert_type.category
    })


# ============ Photo Upload Routes ============

def get_storage_client():
    """Get object storage client - returns None if not configured"""
    try:
        # Get bucket ID from environment variable
        bucket_id = os.environ.get('DEFAULT_OBJECT_STORAGE_BUCKET_ID')
        if bucket_id:
            return ObjectStorageClient(bucket_id=bucket_id)
        return ObjectStorageClient()
    except Exception as e:
        print(f"Object storage not available: {e}")
        return None

ALLOWED_IMAGE_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'webp'}

def allowed_image_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_IMAGE_EXTENSIONS


@app.route('/employees/<int:id>/photos', methods=['POST'])
def upload_employee_photo(id):
    """Upload a photo for an employee"""
    employee = Employee.query.get_or_404(id)
    
    if 'photo' not in request.files:
        return jsonify({'success': False, 'error': 'No photo file provided'}), 400
    
    file = request.files['photo']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'}), 400
    
    if not allowed_image_file(file.filename):
        return jsonify({'success': False, 'error': 'Invalid file type. Allowed: PNG, JPG, JPEG, GIF, WEBP'}), 400
    
    client = get_storage_client()
    if not client:
        return jsonify({'success': False, 'error': 'Object storage not configured'}), 500
    
    # Generate unique storage path
    ext = file.filename.rsplit('.', 1)[1].lower()
    unique_id = str(uuid.uuid4())
    storage_path = f"employees/{id}/photos/{unique_id}.{ext}"
    
    try:
        # Upload to object storage
        file_data = file.read()
        client.upload_from_bytes(storage_path, file_data)
        
        # Create database record
        photo = EmployeePhoto(
            employee_id=id,
            filename=secure_filename(file.filename),
            storage_path=storage_path,
            caption=request.form.get('caption', ''),
            file_size=len(file_data),
            content_type=file.content_type
        )
        db.session.add(photo)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'photo': {
                'id': photo.id,
                'filename': photo.filename,
                'caption': photo.caption,
                'url': f'/photos/employee/{photo.id}'
            }
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/employees/<int:id>/photos/<int:photo_id>', methods=['DELETE'])
def delete_employee_photo(id, photo_id):
    """Delete a photo for an employee"""
    photo = EmployeePhoto.query.filter_by(id=photo_id, employee_id=id).first_or_404()
    
    client = get_storage_client()
    if client:
        try:
            client.delete(photo.storage_path)
        except Exception as e:
            print(f"Error deleting from storage: {e}")
    
    db.session.delete(photo)
    db.session.commit()
    
    return jsonify({'success': True})


@app.route('/photos/employee/<int:photo_id>')
def serve_employee_photo(photo_id):
    """Serve an employee photo from object storage"""
    photo = EmployeePhoto.query.get_or_404(photo_id)
    
    client = get_storage_client()
    if not client:
        return "Object storage not configured", 500
    
    try:
        data = client.download_as_bytes(photo.storage_path)
        return send_file(
            io.BytesIO(data),
            mimetype=photo.content_type or 'image/jpeg',
            download_name=photo.filename
        )
    except Exception as e:
        return f"Error loading photo: {e}", 404


@app.route('/projects/<int:id>/photos', methods=['POST'])
def upload_project_photo(id):
    """Upload a photo for a project"""
    project = Project.query.get_or_404(id)
    
    if 'photo' not in request.files:
        return jsonify({'success': False, 'error': 'No photo file provided'}), 400
    
    file = request.files['photo']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'}), 400
    
    if not allowed_image_file(file.filename):
        return jsonify({'success': False, 'error': 'Invalid file type. Allowed: PNG, JPG, JPEG, GIF, WEBP'}), 400
    
    client = get_storage_client()
    if not client:
        return jsonify({'success': False, 'error': 'Object storage not configured'}), 500
    
    # Generate unique storage path
    ext = file.filename.rsplit('.', 1)[1].lower()
    unique_id = str(uuid.uuid4())
    storage_path = f"projects/{id}/photos/{unique_id}.{ext}"
    
    try:
        # Upload to object storage
        file_data = file.read()
        client.upload_from_bytes(storage_path, file_data)
        
        # Create database record
        photo = ProjectPhoto(
            project_id=id,
            filename=secure_filename(file.filename),
            storage_path=storage_path,
            caption=request.form.get('caption', ''),
            file_size=len(file_data),
            content_type=file.content_type
        )
        db.session.add(photo)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'photo': {
                'id': photo.id,
                'filename': photo.filename,
                'caption': photo.caption,
                'url': f'/photos/project/{photo.id}'
            }
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/projects/<int:id>/photos/<int:photo_id>', methods=['DELETE'])
def delete_project_photo(id, photo_id):
    """Delete a photo for a project"""
    photo = ProjectPhoto.query.filter_by(id=photo_id, project_id=id).first_or_404()
    
    client = get_storage_client()
    if client:
        try:
            client.delete(photo.storage_path)
        except Exception as e:
            print(f"Error deleting from storage: {e}")
    
    db.session.delete(photo)
    db.session.commit()
    
    return jsonify({'success': True})


@app.route('/photos/project/<int:photo_id>')
def serve_project_photo(photo_id):
    """Serve a project photo from object storage"""
    photo = ProjectPhoto.query.get_or_404(photo_id)
    
    client = get_storage_client()
    if not client:
        return "Object storage not configured", 500
    
    try:
        data = client.download_as_bytes(photo.storage_path)
        return send_file(
            io.BytesIO(data),
            mimetype=photo.content_type or 'image/jpeg',
            download_name=photo.filename
        )
    except Exception as e:
        return f"Error loading photo: {e}", 404


@app.route('/employees/<int:id>/photos/<int:photo_id>/caption', methods=['PUT'])
def update_employee_photo_caption(id, photo_id):
    """Update the caption for an employee photo"""
    photo = EmployeePhoto.query.filter_by(id=photo_id, employee_id=id).first_or_404()
    data = request.json
    photo.caption = data.get('caption', '')
    db.session.commit()
    return jsonify({'success': True})


@app.route('/employees/<int:id>/photos/<int:photo_id>/set-primary', methods=['POST'])
def set_employee_primary_photo(id, photo_id):
    """Set a photo as the primary photo for an employee's resume"""
    photo = EmployeePhoto.query.filter_by(id=photo_id, employee_id=id).first_or_404()
    
    # Clear any existing primary photo for this employee
    EmployeePhoto.query.filter_by(employee_id=id, is_primary=True).update({'is_primary': False})
    
    # Set this photo as primary
    photo.is_primary = True
    db.session.commit()
    
    return jsonify({'success': True})


@app.route('/projects/<int:id>/photos/<int:photo_id>/caption', methods=['PUT'])
def update_project_photo_caption(id, photo_id):
    """Update the caption for a project photo"""
    photo = ProjectPhoto.query.filter_by(id=photo_id, project_id=id).first_or_404()
    data = request.json
    photo.caption = data.get('caption', '')
    db.session.commit()
    return jsonify({'success': True})


@app.route('/projects/<int:id>/photos/<int:photo_id>/set-primary', methods=['POST'])
def set_project_primary_photo(id, photo_id):
    """Set a photo as the primary photo for a project"""
    photo = ProjectPhoto.query.filter_by(id=photo_id, project_id=id).first_or_404()
    
    # Clear any existing primary photo for this project
    ProjectPhoto.query.filter_by(project_id=id, is_primary=True).update({'is_primary': False})
    
    # Set this photo as primary
    photo.is_primary = True
    db.session.commit()
    
    return jsonify({'success': True})


# ============ Firm Photo Routes ============

@app.route('/firms/<int:id>/photos', methods=['POST'])
def upload_firm_photo(id):
    """Upload a photo for a firm"""
    from models import FirmPhoto
    firm = Firm.query.get_or_404(id)
    
    if 'photo' not in request.files:
        return jsonify({'success': False, 'error': 'No photo file provided'}), 400
    
    file = request.files['photo']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'}), 400
    
    if not allowed_image_file(file.filename):
        return jsonify({'success': False, 'error': 'Invalid file type. Allowed: PNG, JPG, JPEG, GIF, WEBP'}), 400
    
    client = get_storage_client()
    if not client:
        return jsonify({'success': False, 'error': 'Object storage not configured'}), 500
    
    ext = file.filename.rsplit('.', 1)[1].lower()
    unique_id = str(uuid.uuid4())
    storage_path = f"firms/{id}/photos/{unique_id}.{ext}"
    
    try:
        file_data = file.read()
        client.upload_from_bytes(storage_path, file_data)
        
        photo = FirmPhoto(
            firm_id=id,
            filename=secure_filename(file.filename),
            storage_path=storage_path,
            caption=request.form.get('caption', ''),
            file_size=len(file_data),
            content_type=file.content_type
        )
        db.session.add(photo)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'photo': {
                'id': photo.id,
                'filename': photo.filename,
                'caption': photo.caption,
                'storage_path': photo.storage_path
            }
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/firms/<int:id>/photos/<int:photo_id>', methods=['DELETE'])
def delete_firm_photo(id, photo_id):
    """Delete a firm photo"""
    from models import FirmPhoto
    photo = FirmPhoto.query.filter_by(id=photo_id, firm_id=id).first_or_404()
    
    client = get_storage_client()
    if client:
        try:
            client.delete(photo.storage_path)
        except Exception as e:
            print(f"Error deleting from storage: {e}")
    
    db.session.delete(photo)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/firms/<int:id>/photos/<int:photo_id>/caption', methods=['PUT'])
def update_firm_photo_caption(id, photo_id):
    """Update a firm photo caption"""
    from models import FirmPhoto
    photo = FirmPhoto.query.filter_by(id=photo_id, firm_id=id).first_or_404()
    data = request.json
    photo.caption = data.get('caption', '')
    db.session.commit()
    return jsonify({'success': True})


@app.route('/firms/<int:id>/photos/<int:photo_id>/set-primary', methods=['POST'])
def set_firm_primary_photo(id, photo_id):
    """Set a photo as the primary photo for a firm"""
    from models import FirmPhoto
    photo = FirmPhoto.query.filter_by(id=photo_id, firm_id=id).first_or_404()
    
    FirmPhoto.query.filter_by(firm_id=id, is_primary=True).update({'is_primary': False})
    photo.is_primary = True
    db.session.commit()
    
    return jsonify({'success': True})


@app.route('/api/firms/<int:id>/photos')
def get_firm_photos(id):
    """Get all photos for a firm"""
    photos = FirmPhoto.query.filter_by(firm_id=id).order_by(FirmPhoto.is_primary.desc(), FirmPhoto.created_at.desc()).all()
    return jsonify([{
        'id': p.id,
        'filename': p.filename,
        'url': f'/photos/firm/{p.id}',
        'caption': p.caption,
        'is_primary': p.is_primary
    } for p in photos])


@app.route('/photos/firm/<int:photo_id>')
def serve_firm_photo(photo_id):
    """Serve a firm photo from object storage"""
    photo = FirmPhoto.query.get_or_404(photo_id)
    
    client = get_storage_client()
    if not client:
        return "Storage not configured", 500
    
    try:
        data = client.download_as_bytes(photo.storage_path)
        return send_file(
            io.BytesIO(data),
            mimetype=photo.content_type or 'image/jpeg',
            download_name=photo.filename
        )
    except Exception as e:
        print(f"Error serving firm photo: {e}")
        return "Photo not found", 404


# Firm Documents Routes
ALLOWED_DOCUMENT_EXTENSIONS = {'pdf', 'doc', 'docx', 'xls', 'xlsx', 'ppt', 'pptx', 'txt', 'csv'}

def allowed_document_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_DOCUMENT_EXTENSIONS

def get_document_type(filename):
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''
    if ext == 'pdf':
        return 'pdf'
    elif ext in ['doc', 'docx']:
        return 'word'
    elif ext in ['xls', 'xlsx', 'csv']:
        return 'excel'
    elif ext in ['ppt', 'pptx']:
        return 'powerpoint'
    else:
        return 'other'


@app.route('/firms/<int:id>/documents', methods=['POST'])
@login_required
def upload_firm_document(id):
    """Upload a document for a firm"""
    from models import FirmDocument
    firm = Firm.query.get_or_404(id)
    
    if 'document' not in request.files:
        return jsonify({'success': False, 'error': 'No document file provided'}), 400
    
    file = request.files['document']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'}), 400
    
    if not allowed_document_file(file.filename):
        return jsonify({'success': False, 'error': 'Invalid file type. Allowed: PDF, DOC, DOCX, XLS, XLSX, PPT, PPTX, TXT, CSV'}), 400
    
    client = get_storage_client()
    if not client:
        return jsonify({'success': False, 'error': 'Object storage not configured'}), 500
    
    ext = file.filename.rsplit('.', 1)[1].lower()
    unique_id = str(uuid.uuid4())
    storage_path = f"firms/{id}/documents/{unique_id}.{ext}"
    
    try:
        file_data = file.read()
        client.upload_from_bytes(storage_path, file_data)
        
        doc = FirmDocument(
            firm_id=id,
            filename=secure_filename(file.filename),
            storage_path=storage_path,
            description=request.form.get('description', ''),
            file_size=len(file_data),
            content_type=file.content_type,
            document_type=get_document_type(file.filename)
        )
        db.session.add(doc)
        db.session.commit()
        
        return jsonify({
            'success': True,
            'document': {
                'id': doc.id,
                'filename': doc.filename,
                'description': doc.description,
                'document_type': doc.document_type,
                'file_size': doc.file_size
            }
        })
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/firms/<int:id>/documents/<int:doc_id>', methods=['DELETE'])
@login_required
def delete_firm_document(id, doc_id):
    """Delete a firm document"""
    from models import FirmDocument
    doc = FirmDocument.query.filter_by(id=doc_id, firm_id=id).first_or_404()
    
    client = get_storage_client()
    if client:
        try:
            client.delete(doc.storage_path)
        except Exception as e:
            print(f"Error deleting document from storage: {e}")
    
    db.session.delete(doc)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/firms/<int:id>/documents/<int:doc_id>/description', methods=['PUT'])
@login_required
def update_firm_document_description(id, doc_id):
    """Update a firm document description"""
    from models import FirmDocument
    doc = FirmDocument.query.filter_by(id=doc_id, firm_id=id).first_or_404()
    data = request.json
    doc.description = data.get('description', '')
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/firms/<int:id>/documents')
@login_required
def get_firm_documents(id):
    """Get all documents for a firm"""
    from models import FirmDocument
    docs = FirmDocument.query.filter_by(firm_id=id).order_by(FirmDocument.created_at.desc()).all()
    return jsonify([{
        'id': d.id,
        'filename': d.filename,
        'url': f'/documents/firm/{d.id}',
        'description': d.description,
        'document_type': d.document_type,
        'file_size': d.file_size,
        'created_at': d.created_at.isoformat() if d.created_at else None
    } for d in docs])


@app.route('/documents/firm/<int:doc_id>')
@login_required
def serve_firm_document(doc_id):
    """Serve a firm document from object storage"""
    from models import FirmDocument
    doc = FirmDocument.query.get_or_404(doc_id)
    
    client = get_storage_client()
    if not client:
        return "Storage not configured", 500
    
    try:
        data = client.download_as_bytes(doc.storage_path)
        return send_file(
            io.BytesIO(data),
            mimetype=doc.content_type or 'application/octet-stream',
            download_name=doc.filename,
            as_attachment=True
        )
    except Exception as e:
        print(f"Error serving firm document: {e}")
        return "Document not found", 404


# Marketing Photos Routes
@app.route('/marketing-photos')
def marketing_photos():
    """Marketing photos page with tag filtering"""
    from models import MarketingPhoto
    
    # Get selected tags from query string
    selected_tags = request.args.getlist('tags')
    
    # Get all photos
    query = MarketingPhoto.query
    photos = query.order_by(MarketingPhoto.created_at.desc()).all()
    
    # Get all unique tags
    all_tags = set()
    for photo in photos:
        all_tags.update(photo.get_tags_list())
    all_tags = sorted(all_tags)
    
    # Filter photos by selected tags if any (AND logic - must have ALL selected tags)
    if selected_tags:
        filtered_photos = []
        for photo in photos:
            photo_tags = [t.lower() for t in photo.get_tags_list()]
            # Check if ALL selected tags are present in photo's tags
            if all(tag.lower() in photo_tags for tag in selected_tags):
                filtered_photos.append(photo)
        photos = filtered_photos
    
    return render_template('marketing_photos.html', 
                           photos=photos, 
                           all_tags=all_tags,
                           selected_tags=selected_tags)


@app.route('/marketing-photos', methods=['POST'])
def upload_marketing_photo():
    """Upload a new marketing photo"""
    from models import MarketingPhoto
    
    if 'photo' not in request.files:
        flash('No photo file provided', 'error')
        return redirect(url_for('marketing_photos'))
    
    file = request.files['photo']
    if file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('marketing_photos'))
    
    if not allowed_image_file(file.filename):
        flash('Invalid file type. Allowed: PNG, JPG, JPEG, GIF, WEBP', 'error')
        return redirect(url_for('marketing_photos'))
    
    client = get_storage_client()
    if not client:
        flash('Object storage not configured', 'error')
        return redirect(url_for('marketing_photos'))
    
    ext = file.filename.rsplit('.', 1)[1].lower()
    unique_id = str(uuid.uuid4())
    storage_path = f"marketing/{unique_id}.{ext}"
    
    try:
        file_data = file.read()
        client.upload_from_bytes(storage_path, file_data)
        
        # Parse tags from input
        tags_input = request.form.get('tags', '')
        tags_list = [t.strip() for t in tags_input.split(',') if t.strip()]
        # Ensure tags start with #
        tags_list = ['#' + t.lstrip('#') for t in tags_list]
        
        photo = MarketingPhoto(
            filename=secure_filename(file.filename),
            storage_path=storage_path,
            caption=request.form.get('caption', ''),
            tags=','.join(tags_list),
            file_size=len(file_data),
            content_type=file.content_type
        )
        db.session.add(photo)
        db.session.commit()
        
        flash('Photo uploaded successfully', 'success')
    except Exception as e:
        flash(f'Error uploading photo: {str(e)}', 'error')
    
    return redirect(url_for('marketing_photos'))


@app.route('/photos/marketing/<int:photo_id>')
def serve_marketing_photo(photo_id):
    """Serve a marketing photo from object storage"""
    photo = MarketingPhoto.query.get_or_404(photo_id)
    
    client = get_storage_client()
    if not client:
        return "Storage not configured", 500
    
    try:
        data = client.download_as_bytes(photo.storage_path)
        return send_file(
            io.BytesIO(data),
            mimetype=photo.content_type or 'image/jpeg',
            download_name=photo.filename
        )
    except Exception as e:
        print(f"Error serving marketing photo: {e}")
        return "Photo not found", 404


@app.route('/marketing-photos/<int:id>', methods=['DELETE'])
def delete_marketing_photo(id):
    """Delete a marketing photo"""
    photo = MarketingPhoto.query.get_or_404(id)
    
    client = get_storage_client()
    if client:
        try:
            client.delete(photo.storage_path)
        except Exception as e:
            print(f"Error deleting photo from storage: {e}")
    
    db.session.delete(photo)
    db.session.commit()
    
    return jsonify({'success': True})


@app.route('/marketing-photos/<int:id>/update', methods=['POST'])
def update_marketing_photo(id):
    """Update marketing photo caption and tags"""
    from models import MarketingPhoto
    photo = MarketingPhoto.query.get_or_404(id)
    
    data = request.get_json()
    if 'caption' in data:
        photo.caption = data['caption']
    if 'tags' in data:
        tags_input = data['tags']
        tags_list = [t.strip() for t in tags_input.split(',') if t.strip()]
        tags_list = ['#' + t.lstrip('#') for t in tags_list]
        photo.tags = ','.join(tags_list)
    
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/marketing-photos')
def get_marketing_photos_api():
    """Get all marketing photos, optionally filtered by tags"""
    from models import MarketingPhoto
    
    selected_tags = request.args.getlist('tags')
    photos = MarketingPhoto.query.order_by(MarketingPhoto.created_at.desc()).all()
    
    # Filter by tags if specified
    if selected_tags:
        filtered_photos = []
        for photo in photos:
            photo_tags = photo.get_tags_list()
            if any(tag in photo_tags for tag in selected_tags):
                filtered_photos.append(photo)
        photos = filtered_photos
    
    return jsonify([{
        'id': p.id,
        'filename': p.filename,
        'url': f'/photos/marketing/{p.id}',
        'caption': p.caption,
        'tags': p.get_tags_list()
    } for p in photos])


@app.route('/api/marketing-photos/<int:photo_id>/copy-to-employee/<int:employee_id>', methods=['POST'])
def copy_marketing_to_employee(photo_id, employee_id):
    """Copy a marketing photo to an employee"""
    from models import MarketingPhoto, EmployeePhoto
    import uuid
    import traceback
    
    marketing_photo = MarketingPhoto.query.get_or_404(photo_id)
    employee = Employee.query.get_or_404(employee_id)
    
    try:
        client = get_storage_client()
        if not client:
            return jsonify({'success': False, 'error': 'Object storage not available'}), 500
        original_data = client.download_as_bytes(marketing_photo.storage_path)
        
        new_filename = f"{uuid.uuid4()}_{marketing_photo.filename}"
        new_storage_path = f"employee_photos/{new_filename}"
        client.upload_from_bytes(new_storage_path, original_data)
        
        photo = EmployeePhoto(
            employee_id=employee.id,
            filename=marketing_photo.filename,
            storage_path=new_storage_path,
            caption=marketing_photo.caption,
            file_size=marketing_photo.file_size,
            content_type=marketing_photo.content_type
        )
        db.session.add(photo)
        
        # Add employee name as tag to marketing photo
        employee_tag = f"#{employee.name.replace(' ', '')}"
        existing_tags = marketing_photo.tags or ""
        if employee_tag.lower() not in existing_tags.lower():
            if existing_tags:
                marketing_photo.tags = f"{existing_tags},{employee_tag}"
            else:
                marketing_photo.tags = employee_tag
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'photo': {
                'id': photo.id,
                'url': f'/photos/employee/{photo.id}',
                'caption': photo.caption,
                'filename': photo.filename
            }
        })
    except Exception as e:
        print(f"Error copying marketing photo to employee: {e}")
        traceback.print_exc()
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/marketing-photos/<int:photo_id>/copy-to-project/<int:project_id>', methods=['POST'])
def copy_marketing_to_project(photo_id, project_id):
    """Copy a marketing photo to a project"""
    from models import MarketingPhoto, ProjectPhoto
    import uuid
    import traceback
    
    marketing_photo = MarketingPhoto.query.get_or_404(photo_id)
    project = Project.query.get_or_404(project_id)
    
    try:
        client = get_storage_client()
        if not client:
            return jsonify({'success': False, 'error': 'Object storage not available'}), 500
        original_data = client.download_as_bytes(marketing_photo.storage_path)
        
        new_filename = f"{uuid.uuid4()}_{marketing_photo.filename}"
        new_storage_path = f"project_photos/{new_filename}"
        client.upload_from_bytes(new_storage_path, original_data)
        
        photo = ProjectPhoto(
            project_id=project.id,
            filename=marketing_photo.filename,
            storage_path=new_storage_path,
            caption=marketing_photo.caption,
            file_size=marketing_photo.file_size,
            content_type=marketing_photo.content_type
        )
        db.session.add(photo)
        
        # Add project name as tag to marketing photo
        project_tag = f"#{project.title.replace(' ', '')}"
        existing_tags = marketing_photo.tags or ""
        if project_tag.lower() not in existing_tags.lower():
            if existing_tags:
                marketing_photo.tags = f"{existing_tags},{project_tag}"
            else:
                marketing_photo.tags = project_tag
        
        db.session.commit()
        
        return jsonify({
            'success': True,
            'photo': {
                'id': photo.id,
                'url': f'/photos/project/{photo.id}',
                'caption': photo.caption,
                'filename': photo.filename
            }
        })
    except Exception as e:
        print(f"Error copying marketing photo to project: {e}")
        traceback.print_exc()
        db.session.rollback()
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/marketing-photos/tags')
def get_all_marketing_tags():
    """Get all unique tags from marketing photos"""
    from models import MarketingPhoto
    photos = MarketingPhoto.query.all()
    all_tags = set()
    for photo in photos:
        all_tags.update(photo.get_tags_list())
    return jsonify(sorted(all_tags))


@app.route('/api/marketing-photos/<int:photo_id>/add-tag', methods=['POST'])
def add_tag_to_marketing_photo(photo_id):
    """Add a tag to a marketing photo"""
    from models import MarketingPhoto
    
    marketing_photo = MarketingPhoto.query.get_or_404(photo_id)
    data = request.get_json()
    new_tag = data.get('tag', '').strip()
    
    if not new_tag:
        return jsonify({'success': False, 'error': 'Tag is required'}), 400
    
    # Ensure tag starts with #
    if not new_tag.startswith('#'):
        new_tag = f"#{new_tag}"
    
    # Remove spaces from tag
    new_tag = new_tag.replace(' ', '')
    
    existing_tags = marketing_photo.tags or ""
    if new_tag.lower() not in existing_tags.lower():
        if existing_tags:
            marketing_photo.tags = f"{existing_tags},{new_tag}"
        else:
            marketing_photo.tags = new_tag
        db.session.commit()
    
    return jsonify({
        'success': True,
        'tags': marketing_photo.get_tags_list()
    })


@app.route('/api/marketing-photos/<int:photo_id>/remove-tag', methods=['POST'])
def remove_tag_from_marketing_photo(photo_id):
    """Remove a tag from a marketing photo"""
    from models import MarketingPhoto
    
    marketing_photo = MarketingPhoto.query.get_or_404(photo_id)
    data = request.get_json()
    tag_to_remove = data.get('tag', '').strip()
    
    if not tag_to_remove:
        return jsonify({'success': False, 'error': 'Tag is required'}), 400
    
    current_tags = marketing_photo.get_tags_list()
    # Remove the tag (case-insensitive match)
    new_tags = [t for t in current_tags if t.lower() != tag_to_remove.lower()]
    marketing_photo.tags = ','.join(new_tags) if new_tags else None
    db.session.commit()
    
    return jsonify({
        'success': True,
        'tags': marketing_photo.get_tags_list()
    })


@app.route('/api/projects-list')
def get_projects_list():
    """Get all projects for dropdown (alphabetically sorted)"""
    projects = Project.query.order_by(Project.title).all()
    return jsonify([{
        'id': p.id,
        'title': p.title
    } for p in projects])


@app.route('/api/employees-list')
def get_employees_list():
    """Get all employees for dropdown (alphabetically sorted)"""
    employees = Employee.query.order_by(Employee.name).all()
    return jsonify([{
        'id': e.id,
        'name': e.name
    } for e in employees])


@app.route('/marketing-photos/scrape', methods=['POST'])
def scrape_marketing_photos():
    """Scrape images from a website (crawling multiple pages) and add them to marketing photos"""
    try:
        return _scrape_marketing_photos_impl()
    except Exception as e:
        import traceback
        error_details = traceback.format_exc()
        print(f"Scrape error: {error_details}")
        flash(f'Scraping failed: {str(e)}', 'error')
        return redirect(url_for('marketing_photos'))

def _scrape_marketing_photos_impl():
    """Implementation of image scraping"""
    try:
        from models import MarketingPhoto
        from bs4 import BeautifulSoup
        from urllib.parse import urljoin, urlparse
        import requests
    except ImportError as e:
        flash(f'Missing required library: {str(e)}', 'error')
        return redirect(url_for('marketing_photos'))
    
    url = request.form.get('website_url', '').strip()
    try:
        min_size_kb = int(request.form.get('min_size_kb', 100))
    except (ValueError, TypeError):
        min_size_kb = 100
    default_tag = request.form.get('default_tag', '').strip()
    
    if not url:
        flash('Please enter a website URL', 'error')
        return redirect(url_for('marketing_photos'))
    
    if not url.startswith(('http://', 'https://')):
        url = 'https://' + url
    
    try:
        client = get_storage_client()
        if not client:
            flash('Object storage not configured', 'error')
            return redirect(url_for('marketing_photos'))
    except Exception as e:
        flash(f'Object storage error: {str(e)}', 'error')
        return redirect(url_for('marketing_photos'))
    
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36'
        }
        
        base_domain = urlparse(url).netloc
        visited_pages = set()
        pages_to_visit = [url]
        all_img_urls = set()
        max_pages = 10  # Reduced to prevent timeout in production
        
        def extract_images_from_soup(soup, page_url):
            """Extract all image URLs from a BeautifulSoup object"""
            img_urls = set()
            for img in soup.find_all('img'):
                for attr in ['src', 'data-src', 'data-lazy-src', 'data-original', 'data-srcset']:
                    src = img.get(attr)
                    if src:
                        if attr == 'data-srcset' or attr == 'srcset':
                            for part in src.split(','):
                                s = part.strip().split()[0]
                                full_url = urljoin(page_url, s)
                                if full_url.startswith(('http://', 'https://')):
                                    img_urls.add(full_url)
                        else:
                            full_url = urljoin(page_url, src)
                            if full_url.startswith(('http://', 'https://')):
                                img_urls.add(full_url)
            
            for source in soup.find_all('source'):
                srcset = source.get('srcset')
                if srcset:
                    for part in srcset.split(','):
                        src = part.strip().split()[0]
                        full_url = urljoin(page_url, src)
                        if full_url.startswith(('http://', 'https://')):
                            img_urls.add(full_url)
            
            for div in soup.find_all(style=True):
                style = div.get('style', '')
                if 'background' in style and 'url(' in style:
                    import re
                    urls = re.findall(r'url\([\'"]?([^\'")\s]+)[\'"]?\)', style)
                    for u in urls:
                        full_url = urljoin(page_url, u)
                        if full_url.startswith(('http://', 'https://')):
                            img_urls.add(full_url)
            
            return img_urls
        
        def extract_links_from_soup(soup, page_url):
            """Extract internal links from a page"""
            links = set()
            for a in soup.find_all('a', href=True):
                href = a['href']
                full_url = urljoin(page_url, href)
                parsed = urlparse(full_url)
                if parsed.netloc == base_domain and parsed.scheme in ('http', 'https'):
                    clean_url = f"{parsed.scheme}://{parsed.netloc}{parsed.path}"
                    if not any(ext in clean_url.lower() for ext in ['.pdf', '.doc', '.zip', '.mp4', '.mp3']):
                        links.add(clean_url)
            return links
        
        pages_crawled = 0
        while pages_to_visit and pages_crawled < max_pages:
            current_page = pages_to_visit.pop(0)
            if current_page in visited_pages:
                continue
            
            visited_pages.add(current_page)
            pages_crawled += 1
            
            try:
                print(f"Crawling page {pages_crawled}: {current_page}")
                response = requests.get(current_page, headers=headers, timeout=15)
                if response.status_code != 200:
                    continue
                
                soup = BeautifulSoup(response.text, 'html.parser')
                
                page_images = extract_images_from_soup(soup, current_page)
                all_img_urls.update(page_images)
                
                page_links = extract_links_from_soup(soup, current_page)
                for link in page_links:
                    if link not in visited_pages:
                        pages_to_visit.append(link)
                        
            except Exception as e:
                print(f"Error crawling {current_page}: {e}")
                continue
        
        print(f"Crawled {pages_crawled} pages, found {len(all_img_urls)} unique image URLs")
        
        imported_count = 0
        skipped_count = 0
        error_count = 0
        min_size_bytes = min_size_kb * 1024
        
        for img_url in all_img_urls:
            try:
                img_response = requests.get(img_url, headers=headers, timeout=15)
                if img_response.status_code != 200:
                    error_count += 1
                    continue
                
                content_type = img_response.headers.get('content-type', '')
                if not content_type.startswith('image/'):
                    skipped_count += 1
                    continue
                
                img_data = img_response.content
                if len(img_data) < min_size_bytes:
                    skipped_count += 1
                    continue
                
                parsed = urlparse(img_url)
                original_filename = parsed.path.split('/')[-1] or 'image.jpg'
                ext = original_filename.rsplit('.', 1)[-1].lower() if '.' in original_filename else 'jpg'
                if ext not in ['jpg', 'jpeg', 'png', 'gif', 'webp']:
                    ext = 'jpg'
                
                unique_id = str(uuid.uuid4())
                storage_path = f"marketing/{unique_id}.{ext}"
                
                client.upload_from_bytes(storage_path, img_data)
                
                tags_list = []
                if default_tag:
                    tags_list = ['#' + t.strip().lstrip('#') for t in default_tag.split(',') if t.strip()]
                
                photo = MarketingPhoto(
                    filename=secure_filename(original_filename),
                    storage_path=storage_path,
                    caption=f"Imported from {base_domain}",
                    tags=','.join(tags_list),
                    file_size=len(img_data),
                    content_type=content_type
                )
                db.session.add(photo)
                imported_count += 1
                
            except Exception as e:
                print(f"Error downloading image {img_url}: {e}")
                error_count += 1
                continue
        
        db.session.commit()
        
        if imported_count > 0:
            flash(f'Crawled {pages_crawled} pages. Imported {imported_count} images (skipped {skipped_count} under {min_size_kb}KB, {error_count} errors)', 'success')
        else:
            flash(f'Crawled {pages_crawled} pages. No images found over {min_size_kb}KB. Found {len(all_img_urls)} images total, skipped {skipped_count} small, {error_count} errors.', 'warning')
            
    except requests.RequestException as e:
        flash(f'Error fetching website: {str(e)}', 'error')
    except Exception as e:
        flash(f'Error scraping images: {str(e)}', 'error')
    
    return redirect(url_for('marketing_photos'))


@app.route('/orgchart')
def orgchart():
    """Serve the org chart landing page (development mode)"""
    return render_template('orgchart.html')


@app.route('/orgchart-app')
@app.route('/orgchart-app/')
def orgchart_app():
    """Serve the built React org chart app (production mode)"""
    import os
    dist_path = os.path.join(os.path.dirname(__file__), 'orgchart', 'dist', 'index.html')
    if os.path.exists(dist_path):
        return send_file(dist_path)
    else:
        return "Org chart not built. Please build with: cd orgchart && npm run build", 404


@app.route('/orgchart-app/<path:filename>')
def orgchart_static(filename):
    """Serve static files for the built org chart app"""
    import os
    return send_from_directory(
        os.path.join(os.path.dirname(__file__), 'orgchart', 'dist'),
        filename
    )


@app.route('/api/proposals/list', methods=['GET'])
def api_proposals_list():
    """Get list of proposals for org chart dropdown"""
    proposals = Proposal.query.order_by(Proposal.updated_at.desc()).all()
    return jsonify([{
        'id': p.id,
        'name': p.name,
        'tracking_number': p.tracking_number,
        'has_org_chart': p.org_chart_data is not None
    } for p in proposals])


@app.route('/api/proposals/<int:proposal_id>/orgchart', methods=['GET'])
def api_get_proposal_orgchart(proposal_id):
    """Get org chart data for a proposal"""
    proposal = Proposal.query.get_or_404(proposal_id)
    return jsonify({
        'proposal_id': proposal.id,
        'proposal_name': proposal.name,
        'org_chart_data': proposal.org_chart_data,
        'org_chart_notes': proposal.org_chart_notes
    })


@app.route('/api/proposals/<int:proposal_id>/orgchart', methods=['POST'])
def api_save_proposal_orgchart(proposal_id):
    """Save org chart data to a proposal"""
    proposal = Proposal.query.get_or_404(proposal_id)
    data = request.get_json()
    
    proposal.org_chart_data = data.get('org_chart_data')
    proposal.org_chart_notes = data.get('org_chart_notes')
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'message': f'Org chart saved to proposal: {proposal.name}'
    })


@app.route('/api/projects/list', methods=['GET'])
def api_projects_list():
    """Get list of all projects for bulk assignment"""
    projects = Project.query.order_by(Project.title).all()
    return jsonify([{
        'id': p.id,
        'title': p.title,
        'location': p.location,
        'year_completed': p.year_completed_professional
    } for p in projects])


@app.route('/api/employees/bulk-info', methods=['GET'])
def api_employees_bulk_info():
    """Get basic info for multiple employees by IDs"""
    ids_param = request.args.get('ids', '').strip()
    if not ids_param:
        return jsonify({'success': False, 'error': 'No employee IDs provided'}), 400
    
    try:
        ids = [int(i) for i in ids_param.split(',') if i.strip().isdigit()]
    except (ValueError, TypeError):
        return jsonify({'success': False, 'error': 'Invalid employee IDs'}), 400
    
    if not ids:
        return jsonify({'success': False, 'error': 'No valid employee IDs provided'}), 400
    
    employees = Employee.query.filter(Employee.id.in_(ids)).all()
    return jsonify([{
        'id': e.id,
        'name': e.display_name,
        'title': e.title,
        'role': e.role,
        'years_experience': e.years_experience_total,
        'bio': e.bio
    } for e in employees])


@app.route('/api/employees/bulk-assign-projects', methods=['POST'])
def api_employees_bulk_assign_projects():
    """Bulk assign projects to multiple employees"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'Invalid request data'}), 400
    except:
        return jsonify({'success': False, 'error': 'Invalid JSON'}), 400
    
    employee_ids = data.get('employee_ids', [])
    project_ids = data.get('project_ids', [])
    
    if not employee_ids:
        return jsonify({'success': False, 'error': 'No employees selected'}), 400
    if not project_ids:
        return jsonify({'success': False, 'error': 'No projects selected'}), 400
    
    assigned = 0
    for emp_id in employee_ids:
        for proj_id in project_ids:
            try:
                existing = EmployeeProjectLink.query.filter_by(
                    employee_id=int(emp_id),
                    project_id=int(proj_id)
                ).first()
                
                if not existing:
                    link = EmployeeProjectLink(
                        employee_id=int(emp_id),
                        project_id=int(proj_id)
                    )
                    db.session.add(link)
                    assigned += 1
            except (ValueError, TypeError):
                continue
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'assigned': assigned,
        'message': f'Assigned {assigned} project links'
    })


@app.route('/api/employees/ai-response', methods=['POST'])
def api_employees_ai_response():
    """Generate AI response based on selected employees"""
    from gemini_service import client
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'Invalid request data'}), 400
    except:
        return jsonify({'success': False, 'error': 'Invalid JSON'}), 400
    
    employee_ids = data.get('employee_ids', [])
    prompt = data.get('prompt', '').strip()
    length = data.get('length', 'short')
    include_org_chart = data.get('include_org_chart', False)
    proposal_id = data.get('proposal_id')
    
    if not employee_ids:
        return jsonify({'success': False, 'error': 'No employees selected'}), 400
    if not prompt:
        return jsonify({'success': False, 'error': 'Please enter a prompt'}), 400
    
    word_limits = {
        'short': 250,
        'medium': 500,
        'long': 1000
    }
    word_limit = word_limits.get(length, 250)
    
    try:
        employees = Employee.query.filter(Employee.id.in_([int(i) for i in employee_ids])).all()
    except (ValueError, TypeError):
        return jsonify({'success': False, 'error': 'Invalid employee IDs'}), 400
    
    if not employees:
        return jsonify({'success': False, 'error': 'No employees found with the given IDs'}), 400
    
    staff_info = []
    for e in employees:
        info = f"Name: {e.display_name}\n"
        if e.title:
            info += f"Title: {e.title}\n"
        if e.role:
            info += f"Role: {e.role}\n"
        if e.years_experience_total:
            info += f"Years of Experience: {e.years_experience_total}\n"
        if e.bio:
            info += f"Bio: {e.bio[:500]}...\n" if len(e.bio or '') > 500 else f"Bio: {e.bio}\n"
        if e.education:
            info += f"Education: {e.education}\n"
        if e.registrations:
            info += f"Registrations: {e.registrations}\n"
        staff_info.append(info)
    
    org_chart_context = ""
    if include_org_chart and proposal_id:
        proposal = Proposal.query.get(proposal_id)
        if proposal and proposal.org_chart_data:
            try:
                org_data = json.loads(proposal.org_chart_data) if isinstance(proposal.org_chart_data, str) else proposal.org_chart_data
                nodes = org_data.get('nodes', [])
                org_roles = []
                for node in nodes:
                    node_data = node.get('data', {})
                    if node_data.get('assignedStaff'):
                        org_roles.append(f"- {node_data.get('assignedStaff')} serves as {node_data.get('role', 'Unknown Role')}")
                if org_roles:
                    org_chart_context = "\n\nOrganizational Chart Roles:\n" + "\n".join(org_roles)
            except:
                pass
    
    ai_style = AISettings.get_value('ai_writing_style', '')
    ai_tone = AISettings.get_value('ai_writing_tone', '')
    
    style_instructions = ""
    if ai_style or ai_tone:
        style_instructions = "\n\nWriting Style Guidelines:"
        if ai_style:
            style_instructions += f"\n- Style: {ai_style}"
        if ai_tone:
            style_instructions += f"\n- Tone: {ai_tone}"
    
    full_prompt = f"""You are a professional proposal writer. Based on the following staff information, {prompt}

Staff Information:
{'---'.join(staff_info)}
{org_chart_context}
{style_instructions}

Please write a response of approximately {word_limit} words. Be professional, highlight the team's combined qualifications, and write in a way suitable for government proposals or professional documents."""

    try:
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=full_prompt
        )
        
        return jsonify({
            'success': True,
            'response': response.text
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/employees/list', methods=['GET'])
def api_employees_list():
    """Get list of all employees for selection"""
    employees = Employee.query.order_by(Employee.name).all()
    return jsonify([{
        'id': e.id,
        'name': e.display_name,
        'title': e.title
    } for e in employees])


@app.route('/api/projects/bulk-info', methods=['GET'])
def api_projects_bulk_info():
    """Get basic info for multiple projects by IDs"""
    ids_param = request.args.get('ids', '')
    if not ids_param:
        return jsonify({'success': False, 'error': 'No project IDs provided'}), 400
    
    ids = [int(i) for i in ids_param.split(',') if i.strip().isdigit()]
    if not ids:
        return jsonify({'success': False, 'error': 'No valid project IDs provided'}), 400
    
    projects = Project.query.filter(Project.id.in_(ids)).all()
    return jsonify([{
        'id': p.id,
        'title': p.title,
        'location': p.location,
        'owner_name': p.owner_name,
        'year_completed_professional': p.year_completed_professional
    } for p in projects])


@app.route('/api/projects/bulk-assign-staff', methods=['POST'])
def api_projects_bulk_assign_staff():
    """Assign multiple employees to multiple projects"""
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'Invalid request data'}), 400
    except:
        return jsonify({'success': False, 'error': 'Invalid JSON'}), 400
    
    project_ids = data.get('project_ids', [])
    employee_ids = data.get('employee_ids', [])
    
    if not project_ids:
        return jsonify({'success': False, 'error': 'No projects selected'}), 400
    if not employee_ids:
        return jsonify({'success': False, 'error': 'No employees selected'}), 400
    
    assigned = 0
    for proj_id in project_ids:
        for emp_id in employee_ids:
            try:
                existing = EmployeeProjectExperience.query.filter_by(
                    employee_id=int(emp_id),
                    project_id=int(proj_id)
                ).first()
                
                if not existing:
                    link = EmployeeProjectExperience(
                        employee_id=int(emp_id),
                        project_id=int(proj_id)
                    )
                    db.session.add(link)
                    assigned += 1
            except (ValueError, TypeError):
                continue
    
    db.session.commit()
    
    return jsonify({
        'success': True,
        'assigned': assigned,
        'message': f'Created {assigned} staff-project assignments'
    })


@app.route('/api/projects/ai-response', methods=['POST'])
def api_projects_ai_response():
    """Generate AI response based on selected projects"""
    from gemini_service import client
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'Invalid request data'}), 400
    except:
        return jsonify({'success': False, 'error': 'Invalid JSON'}), 400
    
    project_ids = data.get('project_ids', [])
    prompt = data.get('prompt', '').strip()
    length = data.get('length', 'short')
    
    if not project_ids:
        return jsonify({'success': False, 'error': 'No projects selected'}), 400
    if not prompt:
        return jsonify({'success': False, 'error': 'Please enter a prompt'}), 400
    
    word_limits = {
        'short': 250,
        'medium': 500,
        'long': 1000
    }
    word_limit = word_limits.get(length, 250)
    
    try:
        projects = Project.query.filter(Project.id.in_([int(i) for i in project_ids])).all()
    except (ValueError, TypeError):
        return jsonify({'success': False, 'error': 'Invalid project IDs'}), 400
    
    if not projects:
        return jsonify({'success': False, 'error': 'No projects found with the given IDs'}), 400
    
    project_info = []
    for p in projects:
        info = f"Project Title: {p.title}\n"
        if p.location:
            info += f"Location: {p.location}\n"
        if p.owner_name:
            info += f"Owner/Client: {p.owner_name}\n"
        if p.year_completed_professional:
            info += f"Year Completed: {p.year_completed_professional}\n"
        if p.project_cost:
            info += f"Project Cost: {p.project_cost}\n"
        if p.brief_description:
            desc = p.brief_description[:800] + '...' if len(p.brief_description or '') > 800 else p.brief_description
            info += f"Description: {desc}\n"
        if p.relevance_writeup:
            info += f"Relevance: {p.relevance_writeup[:300]}...\n" if len(p.relevance_writeup or '') > 300 else f"Relevance: {p.relevance_writeup}\n"
        project_info.append(info)
    
    ai_style = AISettings.get_value('ai_writing_style', '')
    ai_tone = AISettings.get_value('ai_writing_tone', '')
    
    style_instructions = ""
    if ai_style or ai_tone:
        style_instructions = "\n\nWriting Style Guidelines:"
        if ai_style:
            style_instructions += f"\n- Style: {ai_style}"
        if ai_tone:
            style_instructions += f"\n- Tone: {ai_tone}"
    
    system_prompt = f"""You are a professional proposal writer specializing in SF330 government proposals for architecture and engineering firms.
You are helping to write content about the following projects:

{chr(10).join([f'--- PROJECT {i+1} ---{chr(10)}{info}' for i, info in enumerate(project_info)])}
{style_instructions}

Write a response of approximately {word_limit} words based on the user's request.
Focus on presenting the projects professionally and highlighting their relevance to government/infrastructure work.
Be specific and use actual data from the project information provided."""

    try:
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config={
                "system_instruction": system_prompt,
                "max_output_tokens": word_limit * 3
            }
        )
        
        return jsonify({
            'success': True,
            'response': response.text
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/proposals/<int:proposal_id>/ai-response', methods=['POST'])
def api_proposal_ai_response(proposal_id):
    """Generate AI response based on all proposal data with chunked processing for large proposals"""
    from gemini_service import client
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'Invalid request data'}), 400
    except:
        return jsonify({'success': False, 'error': 'Invalid JSON'}), 400
    
    prompt = data.get('prompt', '').strip()
    length = data.get('length', 'short')
    
    if not prompt:
        return jsonify({'success': False, 'error': 'Please enter a prompt'}), 400
    
    proposal = Proposal.query.get(proposal_id)
    if not proposal:
        return jsonify({'success': False, 'error': 'Proposal not found'}), 404
    
    word_limits = {'short': 250, 'medium': 500, 'long': 1000}
    word_limit = word_limits.get(length, 250)
    
    # Collect all proposal data
    def collect_proposal_data():
        sections = {}
        
        # Basic proposal info
        proposal_info = f"""PROPOSAL INFORMATION:
Name: {proposal.name}
Tracking Number: {proposal.tracking_number or 'N/A'}
Contract Title: {proposal.contract_title or 'N/A'}
Contract Location: {proposal.contract_location or 'N/A'}
Solicitation Number: {proposal.solicitation_number or 'N/A'}
Status: {proposal.status}
Win Theme: {proposal.win_theme or 'Not defined'}
"""
        sections['proposal_info'] = proposal_info
        
        # Firm info
        if proposal.firm:
            firm = proposal.firm
            firm_info = f"""FIRM INFORMATION:
Name: {firm.name}
Address: {firm.street_address or 'N/A'}
City/State: {firm.city or ''}, {firm.state or ''} {firm.zip_code or ''}
Bio: {firm.bio[:1000] if firm.bio else 'N/A'}
UEI: {firm.uei or 'N/A'}
"""
            sections['firm_info'] = firm_info
        
        # Selected personnel with roles
        selected_employees = ProposalSelectedEmployee.query.filter_by(proposal_id=proposal_id).all()
        if selected_employees:
            personnel_info = "SELECTED PERSONNEL:\n"
            for pse in selected_employees:
                emp = pse.employee
                if emp:
                    personnel_info += f"""
- {emp.display_name}
  Role in Contract: {pse.role_in_contract or 'Not specified'}
  Title: {emp.title or 'N/A'}
  Years Experience: {emp.years_experience_total or 'N/A'}
  Education: {emp.education[:200] if emp.education else 'N/A'}
  Bio: {emp.bio[:300] if emp.bio else 'N/A'}
"""
            sections['personnel'] = personnel_info
        
        # Selected projects
        selected_projects = ProposalSelectedProject.query.filter_by(proposal_id=proposal_id).all()
        if selected_projects:
            projects_info = "SELECTED PROJECTS:\n"
            for psp in selected_projects:
                proj = psp.project
                if proj:
                    desc = psp.custom_writeup or proj.brief_description or 'N/A'
                    projects_info += f"""
- {proj.title}
  Location: {proj.location or 'N/A'}
  Owner/Client: {proj.owner_name or 'N/A'}
  Year Completed: {proj.year_completed_professional or 'N/A'}
  Cost: {proj.project_cost or 'N/A'}
  Description: {desc[:400]}
"""
            sections['projects'] = projects_info
        
        # Org chart data
        if proposal.org_chart_data:
            try:
                org_data = json.loads(proposal.org_chart_data) if isinstance(proposal.org_chart_data, str) else proposal.org_chart_data
                nodes = org_data.get('nodes', [])
                if nodes:
                    org_info = "ORGANIZATIONAL CHART:\n"
                    for node in nodes:
                        node_data = node.get('data', {})
                        role = node_data.get('role', 'Unknown')
                        staff = node_data.get('assignedStaff', 'Unassigned')
                        org_info += f"- {role}: {staff}\n"
                    sections['org_chart'] = org_info
            except:
                pass
        
        # RFP content
        if proposal.rfp_text:
            rfp_text = proposal.rfp_text[:3000] if len(proposal.rfp_text) > 3000 else proposal.rfp_text
            sections['rfp'] = f"RFP/RFQ CONTENT:\n{rfp_text}\n"
        
        # Reference documents
        if proposal.reference_documents:
            refs_info = "REFERENCE DOCUMENTS:\n"
            for ref in proposal.reference_documents:
                if ref.extracted_text:
                    refs_info += f"\n--- {ref.filename} ---\n{ref.extracted_text[:1500]}\n"
            if len(refs_info) > 50:
                sections['references'] = refs_info
        
        # Intelligence documents
        if proposal.intelligence_documents:
            intel_info = "INTELLIGENCE DOCUMENTS:\n"
            for intel in proposal.intelligence_documents:
                intel_info += f"\n--- {intel.filename} ({intel.description or 'No description'}) ---\n"
                if intel.extracted_text:
                    intel_info += f"{intel.extracted_text[:1000]}\n"
            if len(intel_info) > 50:
                sections['intelligence'] = intel_info
        
        return sections
    
    sections = collect_proposal_data()
    
    # Estimate total size (rough character count)
    total_chars = sum(len(s) for s in sections.values())
    
    # Get writing style/tone
    ai_style = AISettings.get_value('ai_writing_style', '')
    ai_tone = AISettings.get_value('ai_writing_tone', '')
    
    style_instructions = ""
    if ai_style or ai_tone:
        style_instructions = "\n\nWriting Style Guidelines:"
        if ai_style:
            style_instructions += f"\n- Style: {ai_style}"
        if ai_tone:
            style_instructions += f"\n- Tone: {ai_tone}"
    
    # If total content is manageable (< 50K chars ~= 12K tokens), send directly
    if total_chars < 50000:
        full_context = "\n\n".join(sections.values())
        
        system_prompt = f"""You are a professional SF330 proposal writer helping with a government architecture/engineering proposal.

Here is all the proposal data:

{full_context}
{style_instructions}

Based on this proposal data, respond to the user's request. Write approximately {word_limit} words.
Be specific, professional, and use actual data from the proposal."""

        try:
            response = client.models.generate_content(
                model="gemini-2.5-flash",
                contents=prompt,
                config={
                    "system_instruction": system_prompt,
                    "max_output_tokens": word_limit * 3
                }
            )
            return jsonify({
                'success': True,
                'response': response.text,
                'chunked': False
            })
        except Exception as e:
            return jsonify({'success': False, 'error': str(e)}), 500
    
    # For large proposals, use chunked summarization
    try:
        summaries = {}
        
        # Summarize each section
        for section_name, section_content in sections.items():
            if len(section_content) > 500:
                summary_prompt = f"""Summarize the following {section_name.replace('_', ' ')} data concisely, keeping key facts and figures:

{section_content[:8000]}

Provide a focused summary in 150-200 words."""

                summary_response = client.models.generate_content(
                    model="gemini-2.5-flash",
                    contents=summary_prompt
                )
                summaries[section_name] = summary_response.text
            else:
                summaries[section_name] = section_content
        
        # Combine summaries for final response with size cap
        combined_parts = [f"=== {k.upper().replace('_', ' ')} ===\n{v}" for k, v in summaries.items()]
        combined_summary = "\n\n".join(combined_parts)
        
        # Cap combined summary to avoid exceeding model limits
        if len(combined_summary) > 30000:
            combined_summary = combined_summary[:30000] + "\n\n[Some content truncated due to size limits]"
        
        final_prompt = f"""You are a professional SF330 proposal writer. Based on the following summarized proposal data:

{combined_summary}
{style_instructions}

User request: {prompt}

Write a response of approximately {word_limit} words. Be specific, professional, and reference actual data from the summaries."""

        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=final_prompt,
            config={"max_output_tokens": word_limit * 3}
        )
        
        return jsonify({
            'success': True,
            'response': response.text,
            'chunked': True
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/proposals/<int:proposal_id>/ai-chat', methods=['POST'])
def api_proposal_ai_chat(proposal_id):
    """Chatbot endpoint for proposal AI assistant with conversation history"""
    from gemini_service import client
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({'success': False, 'error': 'Invalid request data'}), 400
    except:
        return jsonify({'success': False, 'error': 'Invalid JSON'}), 400
    
    message = data.get('message', '').strip()
    history = data.get('history', [])
    
    if not message:
        return jsonify({'success': False, 'error': 'Please enter a message'}), 400
    
    proposal = Proposal.query.get(proposal_id)
    if not proposal:
        return jsonify({'success': False, 'error': 'Proposal not found'}), 404
    
    # Collect proposal data (reuse the same logic)
    def collect_proposal_data():
        sections = {}
        
        proposal_info = f"""PROPOSAL INFORMATION:
Name: {proposal.name}
Tracking Number: {proposal.tracking_number or 'N/A'}
Contract Title: {proposal.contract_title or 'N/A'}
Contract Location: {proposal.contract_location or 'N/A'}
Solicitation Number: {proposal.solicitation_number or 'N/A'}
Status: {proposal.status}
Win Theme: {proposal.win_theme or 'Not defined'}
"""
        sections['proposal_info'] = proposal_info
        
        if proposal.firm:
            firm = proposal.firm
            sections['firm_info'] = f"""FIRM INFORMATION:
Name: {firm.name}
Address: {firm.street_address or 'N/A'}
City/State: {firm.city or ''}, {firm.state or ''} {firm.zip_code or ''}
Bio: {firm.bio[:1000] if firm.bio else 'N/A'}
UEI: {firm.uei or 'N/A'}
"""
        
        selected_employees = ProposalSelectedEmployee.query.filter_by(proposal_id=proposal_id).all()
        if selected_employees:
            personnel_info = "SELECTED PERSONNEL:\n"
            for pse in selected_employees:
                emp = pse.employee
                if emp:
                    personnel_info += f"\n--- {emp.display_name} ---\n"
                    personnel_info += f"Role in Contract: {pse.role_in_contract or 'N/A'}\n"
                    personnel_info += f"Title: {emp.title or 'N/A'}\n"
                    personnel_info += f"Years Experience: {emp.years_experience_total or 'N/A'} (with firm: {emp.years_experience_firm or 'N/A'})\n"
                    if emp.education:
                        personnel_info += f"Education: {emp.education}\n"
                    if emp.registrations:
                        personnel_info += f"Registrations: {emp.registrations}\n"
                    if emp.bio:
                        personnel_info += f"Bio: {emp.bio[:500]}\n"
                    
                    # Include Block 19 selected projects (ProposalEmployeeRelevantProject)
                    if pse.relevant_projects:
                        personnel_info += "Block 19 Selected Projects (for SF330 Resume):\n"
                        for rp in pse.relevant_projects:
                            proj = rp.project
                            if proj:
                                personnel_info += f"  * {proj.title}"
                                if proj.location:
                                    personnel_info += f" | {proj.location}"
                                if proj.year_completed_professional:
                                    personnel_info += f" | {proj.year_completed_professional}"
                                personnel_info += "\n"
                                # Get employee's role on this project
                                link = EmployeeProjectLink.query.filter_by(employee_id=emp.id, project_id=proj.id).first()
                                if link and link.role:
                                    personnel_info += f"    Role: {link.role}\n"
                                if proj.brief_description:
                                    personnel_info += f"    {proj.brief_description[:300]}\n"
                    
                    numbered_exps = EmployeeProjectExperience.query.filter_by(employee_id=emp.id).filter(
                        EmployeeProjectExperience.resume_order.isnot(None)
                    ).order_by(EmployeeProjectExperience.resume_order.asc()).all()
                    experiences = numbered_exps if numbered_exps else EmployeeProjectExperience.query.filter_by(employee_id=emp.id).order_by(EmployeeProjectExperience.year_completed.desc()).limit(5).all()
                    if experiences:
                        personnel_info += "Additional Project Experience:\n"
                        for exp in experiences:
                            personnel_info += f"  * {exp.project_title} ({exp.year_completed or 'N/A'}) - {exp.role_performed or 'N/A'}\n"
                            if exp.brief_description:
                                personnel_info += f"    Main: {exp.brief_description[:300]}\n"
                            for alt in exp.alternate_descriptions:
                                if alt.description:
                                    personnel_info += f"    [{alt.label}]: {alt.description[:300]}\n"
                    
                    personnel_info += "\n"
            sections['personnel'] = personnel_info
        
        selected_projects = ProposalSelectedProject.query.filter_by(proposal_id=proposal_id).all()
        if selected_projects:
            projects_info = "SELECTED PROJECTS:\n"
            for psp in selected_projects:
                proj = psp.project
                if proj:
                    desc = psp.custom_writeup or proj.brief_description or 'N/A'
                    projects_info += f"- {proj.title} | Location: {proj.location or 'N/A'} | Owner: {proj.owner_name or 'N/A'} | Year: {proj.year_completed_professional or 'N/A'}\n  Description: {desc[:300]}\n"
            sections['projects'] = projects_info
        
        if proposal.org_chart_data:
            try:
                org_data = json.loads(proposal.org_chart_data) if isinstance(proposal.org_chart_data, str) else proposal.org_chart_data
                nodes = org_data.get('nodes', [])
                if nodes:
                    org_info = "ORGANIZATIONAL CHART:\n"
                    for node in nodes:
                        node_data = node.get('data', {})
                        org_info += f"- {node_data.get('role', 'Unknown')}: {node_data.get('assignedStaff', 'Unassigned')}\n"
                    sections['org_chart'] = org_info
            except:
                pass
        
        if proposal.rfp_text:
            sections['rfp'] = f"RFP/RFQ CONTENT:\n{proposal.rfp_text[:4000]}\n"
        
        if proposal.reference_documents:
            refs_info = "REFERENCE DOCUMENTS:\n"
            for ref in proposal.reference_documents[:3]:
                if ref.extracted_text:
                    refs_info += f"--- {ref.filename} ---\n{ref.extracted_text[:1500]}\n"
            sections['references'] = refs_info
        
        if proposal.intelligence_documents:
            intel_info = "INTELLIGENCE DOCUMENTS:\n"
            for intel in proposal.intelligence_documents[:3]:
                intel_info += f"--- {intel.filename} ({intel.description or 'No description'}) ---\n"
                if intel.extracted_text:
                    intel_info += f"{intel.extracted_text[:1000]}\n"
            sections['intelligence'] = intel_info
        
        linked_responses = ProposalLinkedResponse.query.filter_by(proposal_id=proposal_id).all()
        if linked_responses:
            responses_info = "LINKED RESPONSE LIBRARY:\n"
            for link in linked_responses[:10]:
                r = link.response
                if r:
                    responses_info += f"\n--- {r.question or 'Response'} (Grade: {r.grade or 'N/A'}) ---\n"
                    responses_info += f"Client: {r.client or 'N/A'} | Year: {r.year or 'N/A'}\n"
                    if r.response:
                        responses_info += f"{r.response[:800]}\n"
            sections['linked_responses'] = responses_info
        
        linked_refs = ProposalLinkedReference.query.filter_by(proposal_id=proposal_id).all()
        if linked_refs:
            refs_context = "PERFORMANCE REFERENCES (Client Evaluations & Testimonials):\n"
            for link in linked_refs[:10]:
                ref = link.reference
                if ref:
                    refs_context += f"\n--- {ref.project_name or ref.client or 'Performance Reference'} ---\n"
                    refs_context += f"Client: {ref.client or 'N/A'} | Firm: {ref.firm or 'N/A'}\n"
                    if ref.final_score:
                        refs_context += f"Score: {ref.final_score}/10"
                        if ref.quality_score:
                            refs_context += f" | Quality: {ref.quality_score}"
                        if ref.schedule_score:
                            refs_context += f" | Schedule: {ref.schedule_score}"
                        if ref.responsiveness_score:
                            refs_context += f" | Responsiveness: {ref.responsiveness_score}"
                        refs_context += "\n"
                    if ref.evaluation_date:
                        refs_context += f"Date: {ref.evaluation_date.strftime('%B %Y')}\n"
                    if ref.consultant_pm:
                        refs_context += f"Consultant PM: {ref.consultant_pm}\n"
                    if ref.score_summary:
                        refs_context += f"Summary: {ref.score_summary[:500]}\n"
                    if ref.quotes:
                        refs_context += f"Notable Quotes: \"{ref.quotes[:500]}\"\n"
            sections['performance_references'] = refs_context
        
        return sections
    
    sections = collect_proposal_data()
    
    # Get writing style/tone
    ai_style = AISettings.get_value('ai_writing_style', '')
    ai_tone = AISettings.get_value('ai_writing_tone', '')
    
    style_instructions = ""
    if ai_style or ai_tone:
        style_instructions = "\n\nWriting Style Guidelines:"
        if ai_style:
            style_instructions += f"\n- Style: {ai_style}"
        if ai_tone:
            style_instructions += f"\n- Tone: {ai_tone}"
    
    # Build context from proposal data
    total_chars = sum(len(s) for s in sections.values())
    
    # For large proposals, summarize
    if total_chars > 40000:
        try:
            summaries = {}
            for section_name, section_content in sections.items():
                if len(section_content) > 500:
                    summary_response = client.models.generate_content(
                        model="gemini-2.5-flash",
                        contents=f"Summarize concisely, keeping key facts:\n\n{section_content[:6000]}"
                    )
                    summaries[section_name] = summary_response.text
                else:
                    summaries[section_name] = section_content
            proposal_context = "\n\n".join([f"=== {k.upper().replace('_', ' ')} ===\n{v}" for k, v in summaries.items()])
        except Exception as e:
            proposal_context = "\n\n".join(sections.values())[:30000]
    else:
        proposal_context = "\n\n".join(sections.values())
    
    # Build conversation for Gemini
    system_prompt = f"""You are a professional SF330 proposal writer helping with a government architecture/engineering proposal. You have access to all the proposal data and can help write content, answer questions, and provide strategic advice.

PROPOSAL DATA:
{proposal_context}
{style_instructions}

Be helpful, specific, and professional. Use actual data from the proposal. Provide thorough, detailed responses."""

    # Build messages including history
    conversation_context = ""
    if history:
        for msg in history[-8:]:  # Keep last 8 messages for context
            role = msg.get('role', 'user')
            content = msg.get('content', '')
            if role == 'user':
                conversation_context += f"\nUser: {content}\n"
            else:
                conversation_context += f"\nAssistant: {content}\n"
    
    full_prompt = f"{conversation_context}\nUser: {message}\n\nAssistant:"

    try:
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=full_prompt,
            config={
                "system_instruction": system_prompt,
                "max_output_tokens": 8000  # Allow long responses
            }
        )
        
        return jsonify({
            'success': True,
            'response': response.text
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@app.route('/api/proposals/<int:proposal_id>/saved-responses', methods=['GET'])
def get_proposal_saved_responses(proposal_id):
    """Get all saved AI responses for a proposal"""
    proposal = Proposal.query.get(proposal_id)
    if not proposal:
        return jsonify({'success': False, 'error': 'Proposal not found'}), 404
    
    responses = ProposalSavedResponse.query.filter_by(proposal_id=proposal_id).order_by(ProposalSavedResponse.created_at.desc()).all()
    
    return jsonify({
        'success': True,
        'responses': [{
            'id': r.id,
            'prompt': r.prompt,
            'response': r.response,
            'label': r.label,
            'created_at': r.created_at.strftime('%Y-%m-%d %H:%M')
        } for r in responses]
    })


@app.route('/api/proposals/<int:proposal_id>/saved-responses', methods=['POST'])
def save_proposal_response(proposal_id):
    """Save an AI response to the proposal"""
    proposal = Proposal.query.get(proposal_id)
    if not proposal:
        return jsonify({'success': False, 'error': 'Proposal not found'}), 404
    
    try:
        data = request.get_json()
    except:
        return jsonify({'success': False, 'error': 'Invalid JSON'}), 400
    
    response_text = data.get('response', '').strip()
    prompt = data.get('prompt', '').strip()
    label = data.get('label', '').strip()
    
    if not response_text:
        return jsonify({'success': False, 'error': 'Response text is required'}), 400
    
    saved = ProposalSavedResponse(
        proposal_id=proposal_id,
        prompt=prompt,
        response=response_text,
        label=label or None
    )
    db.session.add(saved)
    db.session.commit()
    
    return jsonify({
        'success': True,
        'id': saved.id,
        'message': 'Response saved successfully'
    })


@app.route('/api/proposals/<int:proposal_id>/saved-responses/<int:response_id>', methods=['DELETE'])
def delete_proposal_saved_response(proposal_id, response_id):
    """Delete a saved AI response"""
    saved = ProposalSavedResponse.query.filter_by(id=response_id, proposal_id=proposal_id).first()
    if not saved:
        return jsonify({'success': False, 'error': 'Saved response not found'}), 404
    
    db.session.delete(saved)
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Response deleted'})


# ========== RESPONSES ROUTES ==========

@app.route('/responses')
def responses_page():
    """Responses library page with filtering and sorting"""
    query = Response.query
    
    year_filter = request.args.get('year', '')
    client_filter = request.args.get('client', '')
    project_type_filter = request.args.get('project_type', '')
    firm_filter = request.args.get('firm', '')
    grade_filter = request.args.get('grade', '')
    
    if year_filter:
        query = query.filter(Response.year == int(year_filter))
    if client_filter:
        query = query.filter(Response.client == client_filter)
    if project_type_filter:
        query = query.filter(Response.project_type == project_type_filter)
    if firm_filter:
        query = query.filter(Response.firm == firm_filter)
    if grade_filter:
        query = query.filter(Response.grade == grade_filter)
    
    sort = request.args.get('sort', 'year')
    dir = request.args.get('dir', 'desc')
    
    if sort == 'year':
        query = query.order_by(Response.year.desc() if dir == 'desc' else Response.year.asc())
    elif sort == 'grade':
        query = query.order_by(Response.grade.asc() if dir == 'asc' else Response.grade.desc())
    else:
        query = query.order_by(Response.created_at.desc())
    
    responses = query.all()
    
    years = db.session.query(Response.year).filter(Response.year.isnot(None)).distinct().order_by(Response.year.desc()).all()
    years = [y[0] for y in years]
    
    clients = db.session.query(Response.client).filter(Response.client.isnot(None), Response.client != '').distinct().order_by(Response.client).all()
    clients = [c[0] for c in clients]
    
    project_types = db.session.query(Response.project_type).filter(Response.project_type.isnot(None), Response.project_type != '').distinct().order_by(Response.project_type).all()
    project_types = [pt[0] for pt in project_types]
    
    firms = db.session.query(Response.firm).filter(Response.firm.isnot(None), Response.firm != '').distinct().order_by(Response.firm).all()
    firms = [f[0] for f in firms]
    
    grades = db.session.query(Response.grade).filter(Response.grade.isnot(None), Response.grade != '').distinct().order_by(Response.grade).all()
    grades = [g[0] for g in grades]
    
    return render_template('responses.html', 
                           responses=responses, 
                           years=years, clients=clients, project_types=project_types, firms=firms, grades=grades,
                           sort=sort, dir=dir)


@app.route('/responses', methods=['POST'])
def create_response():
    """Create a new response"""
    response = Response(
        year=int(request.form.get('year')) if request.form.get('year') else None,
        client=request.form.get('client', '').strip(),
        project_type=request.form.get('project_type', '').strip(),
        contract=request.form.get('contract', '').strip(),
        firm=request.form.get('firm', '').strip(),
        grade=request.form.get('grade', '').strip(),
        question=request.form.get('question', '').strip(),
        response=request.form.get('response', '').strip(),
        tags=request.form.get('tags', '').strip()
    )
    db.session.add(response)
    db.session.commit()
    flash('Response created successfully', 'success')
    return redirect(url_for('responses_page'))


@app.route('/responses/<int:id>', methods=['POST'])
def update_response(id):
    """Update an existing response"""
    response = Response.query.get_or_404(id)
    
    response.year = int(request.form.get('year')) if request.form.get('year') else None
    response.client = request.form.get('client', '').strip()
    response.project_type = request.form.get('project_type', '').strip()
    response.contract = request.form.get('contract', '').strip()
    response.firm = request.form.get('firm', '').strip()
    response.grade = request.form.get('grade', '').strip()
    response.question = request.form.get('question', '').strip()
    response.response = request.form.get('response', '').strip()
    response.tags = request.form.get('tags', '').strip()
    
    db.session.commit()
    flash('Response updated successfully', 'success')
    return redirect(url_for('responses_page'))


@app.route('/responses/<int:id>', methods=['DELETE'])
def delete_response(id):
    """Delete a response"""
    response = Response.query.get_or_404(id)
    db.session.delete(response)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/responses/<int:id>')
def api_get_response(id):
    """Get response details as JSON"""
    response = Response.query.get(id)
    if not response:
        return jsonify({'success': False, 'error': 'Response not found'}), 404
    
    return jsonify({
        'success': True,
        'response': {
            'id': response.id,
            'year': response.year,
            'client': response.client,
            'project_type': response.project_type,
            'contract': response.contract,
            'firm': response.firm,
            'grade': response.grade,
            'question': response.question,
            'response': response.response,
            'tags': response.tags
        }
    })


@app.route('/responses/import', methods=['POST'])
def import_responses_csv():
    """Import responses from CSV file"""
    import csv
    import io
    
    if 'file' not in request.files:
        flash('No file uploaded', 'error')
        return redirect(url_for('responses_page'))
    
    file = request.files['file']
    if file.filename == '':
        flash('No file selected', 'error')
        return redirect(url_for('responses_page'))
    
    try:
        content = file.read().decode('utf-8-sig')
        reader = csv.DictReader(io.StringIO(content))
        
        count = 0
        for row in reader:
            year_val = row.get('Year', '').strip()
            response_text = row.get('Response', '').strip()
            
            if not response_text:
                continue
            
            response = Response(
                year=int(year_val) if year_val and year_val.isdigit() else None,
                client=row.get('Client', '').strip(),
                project_type=row.get('Project  Type', row.get('Project Type', '')).strip(),
                contract=row.get('Contract', '').strip(),
                firm=row.get('Firm', row.get('Firm ', '')).strip(),
                grade=row.get('Grade', '').strip(),
                question=row.get('Question/Section', row.get('Question', '')).strip(),
                response=response_text
            )
            db.session.add(response)
            count += 1
        
        db.session.commit()
        flash(f'Imported {count} responses successfully', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error importing CSV: {str(e)}', 'error')
    
    return redirect(url_for('responses_page'))


@app.route('/responses/export')
def export_responses_csv():
    """Export responses to CSV file"""
    import csv
    import io
    
    responses = Response.query.order_by(Response.year.desc()).all()
    
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['Year', 'Client', 'Project Type', 'Contract', 'Firm', 'Grade', 'Question/Section', 'Response'])
    
    for r in responses:
        writer.writerow([r.year or '', r.client or '', r.project_type or '', r.contract or '', 
                        r.firm or '', r.grade or '', r.question or '', r.response or ''])
    
    output.seek(0)
    return send_file(
        io.BytesIO(output.getvalue().encode('utf-8-sig')),
        mimetype='text/csv',
        as_attachment=True,
        download_name='responses_export.csv'
    )


@app.route('/api/proposals/<int:proposal_id>/linked-responses')
def get_proposal_linked_responses(proposal_id):
    """Get all linked responses for a proposal"""
    proposal = Proposal.query.get(proposal_id)
    if not proposal:
        return jsonify({'success': False, 'error': 'Proposal not found'}), 404
    
    links = ProposalLinkedResponse.query.filter_by(proposal_id=proposal_id).order_by(ProposalLinkedResponse.display_order).all()
    
    responses = []
    for link in links:
        r = link.response
        responses.append({
            'id': r.id,
            'year': r.year,
            'client': r.client,
            'project_type': r.project_type,
            'contract': r.contract,
            'firm': r.firm,
            'grade': r.grade,
            'question': r.question,
            'response': r.response
        })
    
    return jsonify({'success': True, 'responses': responses})


@app.route('/api/proposals/<int:proposal_id>/link-response', methods=['POST'])
def link_response_to_proposal(proposal_id):
    """Link a response to a proposal"""
    proposal = Proposal.query.get(proposal_id)
    if not proposal:
        return jsonify({'success': False, 'error': 'Proposal not found'}), 404
    
    data = request.get_json()
    response_id = data.get('response_id')
    
    response = Response.query.get(response_id)
    if not response:
        return jsonify({'success': False, 'error': 'Response not found'}), 404
    
    existing = ProposalLinkedResponse.query.filter_by(proposal_id=proposal_id, response_id=response_id).first()
    if existing:
        return jsonify({'success': False, 'error': 'Response already linked to this proposal'}), 400
    
    max_order = db.session.query(db.func.max(ProposalLinkedResponse.display_order)).filter_by(proposal_id=proposal_id).scalar() or 0
    
    link = ProposalLinkedResponse(
        proposal_id=proposal_id,
        response_id=response_id,
        display_order=max_order + 1
    )
    db.session.add(link)
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Response linked to proposal'})


@app.route('/api/proposals/<int:proposal_id>/unlink-response/<int:response_id>', methods=['DELETE'])
def unlink_response_from_proposal(proposal_id, response_id):
    """Unlink a response from a proposal"""
    link = ProposalLinkedResponse.query.filter_by(proposal_id=proposal_id, response_id=response_id).first()
    if not link:
        return jsonify({'success': False, 'error': 'Link not found'}), 404
    
    db.session.delete(link)
    db.session.commit()
    
    return jsonify({'success': True})


@app.route('/api/responses/<int:id>/rewrite', methods=['POST'])
def rewrite_response_with_ai(id):
    """Rewrite a response using AI"""
    from gemini_service import client
    
    response = Response.query.get(id)
    if not response:
        return jsonify({'success': False, 'error': 'Response not found'}), 404
    
    data = request.get_json() or {}
    instructions = data.get('instructions', '').strip()
    
    ai_style = AISettings.get_value('ai_writing_style', '')
    ai_tone = AISettings.get_value('ai_writing_tone', '')
    
    style_instructions = ""
    if ai_style:
        style_instructions += f"\nWriting Style: {ai_style}"
    if ai_tone:
        style_instructions += f"\nTone: {ai_tone}"
    
    prompt = f"""Rewrite the following proposal response. Keep the key information but improve clarity, professionalism, and impact.
{style_instructions}

{f"Additional instructions: {instructions}" if instructions else ""}

Original Response:
{response.response}

Rewritten Response:"""
    
    try:
        result = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt,
            config={"max_output_tokens": 4000}
        )
        return jsonify({'success': True, 'rewritten': result.text})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/responses/<int:id>/save-rewrite', methods=['POST'])
def save_rewritten_response(id):
    """Save a rewritten response as a new entry"""
    original = Response.query.get(id)
    if not original:
        return jsonify({'success': False, 'error': 'Original response not found'}), 404
    
    data = request.get_json()
    rewritten = data.get('rewritten', '').strip()
    
    if not rewritten:
        return jsonify({'success': False, 'error': 'Rewritten text is required'}), 400
    
    new_response = Response(
        year=original.year,
        client=original.client,
        project_type=original.project_type,
        contract=original.contract,
        firm=original.firm,
        grade=original.grade,
        question=original.question,
        response=rewritten,
        tags=(original.tags + ',rewritten') if original.tags else 'rewritten'
    )
    db.session.add(new_response)
    db.session.commit()
    
    return jsonify({'success': True, 'id': new_response.id})


# ============== REFERENCES SECTION ==============

@app.route('/references')
@login_required
def references_page():
    """Display the References management page"""
    clients = db.session.query(Reference.client).distinct().filter(Reference.client.isnot(None), Reference.client != '').order_by(Reference.client).all()
    firms = db.session.query(Reference.firm).distinct().filter(Reference.firm.isnot(None), Reference.firm != '').order_by(Reference.firm).all()
    employees = Employee.query.order_by(Employee.name).all()
    proposals = Proposal.query.order_by(Proposal.name).all()
    
    client_filter = request.args.get('client', '')
    firm_filter = request.args.get('firm', '')
    personnel_filter = request.args.get('personnel', '')
    sort_by = request.args.get('sort', 'date_desc')
    
    query = Reference.query
    
    if client_filter:
        query = query.filter(Reference.client == client_filter)
    if firm_filter:
        query = query.filter(Reference.firm == firm_filter)
    if personnel_filter:
        query = query.filter(Reference.personnel_tags.ilike(f'%{personnel_filter}%'))
    
    if sort_by == 'date_desc':
        query = query.order_by(Reference.evaluation_date.desc().nullslast())
    elif sort_by == 'date_asc':
        query = query.order_by(Reference.evaluation_date.asc().nullsfirst())
    elif sort_by == 'score_desc':
        query = query.order_by(Reference.final_score.desc().nullslast())
    elif sort_by == 'score_asc':
        query = query.order_by(Reference.final_score.asc().nullsfirst())
    elif sort_by == 'client':
        query = query.order_by(Reference.client.asc())
    else:
        query = query.order_by(Reference.created_at.desc())
    
    references = query.all()
    
    projects = Project.query.order_by(Project.title).all()
    
    return render_template('references.html',
                          references=references,
                          clients=[c[0] for c in clients],
                          firms=[f[0] for f in firms],
                          employees=employees,
                          proposals=proposals,
                          projects=projects,
                          client_filter=client_filter,
                          firm_filter=firm_filter,
                          personnel_filter=personnel_filter,
                          sort_by=sort_by)


@app.route('/references', methods=['POST'])
@login_required
def create_reference():
    """Create a new reference manually"""
    from datetime import datetime as dt
    
    project_id = request.form.get('project_id', '').strip()
    ref = Reference(
        project_id=int(project_id) if project_id else None,
        client=request.form.get('client', '').strip(),
        agency=request.form.get('agency', '').strip(),
        project_name=request.form.get('project_name', '').strip(),
        contract_number=request.form.get('contract_number', '').strip(),
        project_id_number=request.form.get('project_id_number', '').strip(),
        final_score=float(request.form.get('final_score')) if request.form.get('final_score') else None,
        schedule_score=float(request.form.get('schedule_score')) if request.form.get('schedule_score') else None,
        quality_score=float(request.form.get('quality_score')) if request.form.get('quality_score') else None,
        responsiveness_score=float(request.form.get('responsiveness_score')) if request.form.get('responsiveness_score') else None,
        key_staff_score=float(request.form.get('key_staff_score')) if request.form.get('key_staff_score') else None,
        dbe_score=float(request.form.get('dbe_score')) if request.form.get('dbe_score') else None,
        pm_performance_score=float(request.form.get('pm_performance_score')) if request.form.get('pm_performance_score') else None,
        score_summary=request.form.get('score_summary', '').strip(),
        quotes=request.form.get('quotes', '').strip(),
        evaluator_name=request.form.get('evaluator_name', '').strip(),
        evaluator_title=request.form.get('evaluator_title', '').strip(),
        consultant_pm=request.form.get('consultant_pm', '').strip(),
        firm=request.form.get('firm', '').strip(),
        services_description=request.form.get('services_description', '').strip(),
        activities_evaluated=request.form.get('activities_evaluated', '').strip(),
        personnel_tags=request.form.get('personnel_tags', '').strip(),
        reference_type=request.form.get('reference_type', 'evaluation'),
        is_final_evaluation=request.form.get('is_final_evaluation') == 'on'
    )
    
    eval_date_str = request.form.get('evaluation_date', '').strip()
    if eval_date_str:
        try:
            ref.evaluation_date = dt.strptime(eval_date_str, '%Y-%m-%d').date()
        except ValueError:
            pass
    
    ref.evaluation_period = request.form.get('evaluation_period', '').strip()
    
    db.session.add(ref)
    db.session.commit()
    
    flash('Reference created successfully', 'success')
    return redirect(url_for('references_page'))


@app.route('/references/<int:id>', methods=['POST'])
@login_required
def update_reference(id):
    """Update an existing reference"""
    from datetime import datetime as dt
    
    ref = Reference.query.get_or_404(id)
    
    project_id = request.form.get('project_id', '').strip()
    ref.project_id = int(project_id) if project_id else None
    ref.client = request.form.get('client', '').strip()
    ref.agency = request.form.get('agency', '').strip()
    ref.project_name = request.form.get('project_name', '').strip()
    ref.contract_number = request.form.get('contract_number', '').strip()
    ref.project_id_number = request.form.get('project_id_number', '').strip()
    ref.final_score = float(request.form.get('final_score')) if request.form.get('final_score') else None
    ref.schedule_score = float(request.form.get('schedule_score')) if request.form.get('schedule_score') else None
    ref.quality_score = float(request.form.get('quality_score')) if request.form.get('quality_score') else None
    ref.responsiveness_score = float(request.form.get('responsiveness_score')) if request.form.get('responsiveness_score') else None
    ref.key_staff_score = float(request.form.get('key_staff_score')) if request.form.get('key_staff_score') else None
    ref.dbe_score = float(request.form.get('dbe_score')) if request.form.get('dbe_score') else None
    ref.pm_performance_score = float(request.form.get('pm_performance_score')) if request.form.get('pm_performance_score') else None
    ref.score_summary = request.form.get('score_summary', '').strip()
    ref.quotes = request.form.get('quotes', '').strip()
    ref.evaluator_name = request.form.get('evaluator_name', '').strip()
    ref.evaluator_title = request.form.get('evaluator_title', '').strip()
    ref.consultant_pm = request.form.get('consultant_pm', '').strip()
    ref.firm = request.form.get('firm', '').strip()
    ref.services_description = request.form.get('services_description', '').strip()
    ref.activities_evaluated = request.form.get('activities_evaluated', '').strip()
    ref.personnel_tags = request.form.get('personnel_tags', '').strip()
    ref.reference_type = request.form.get('reference_type', 'evaluation')
    ref.is_final_evaluation = request.form.get('is_final_evaluation') == 'on'
    
    eval_date_str = request.form.get('evaluation_date', '').strip()
    if eval_date_str:
        try:
            ref.evaluation_date = dt.strptime(eval_date_str, '%Y-%m-%d').date()
        except ValueError:
            pass
    else:
        ref.evaluation_date = None
    
    ref.evaluation_period = request.form.get('evaluation_period', '').strip()
    
    db.session.commit()
    flash('Reference updated successfully', 'success')
    return redirect(url_for('references_page'))


@app.route('/references/<int:id>', methods=['DELETE'])
@login_required
def delete_performance_reference(id):
    """Delete a performance reference"""
    ref = Reference.query.get_or_404(id)
    
    if ref.pdf_object_key:
        try:
            storage = get_storage_client()
            if storage:
                storage.delete(ref.pdf_object_key)
        except Exception:
            pass
    
    db.session.delete(ref)
    db.session.commit()
    return jsonify({'success': True})


@app.route('/api/references/<int:id>')
@login_required
def get_reference_api(id):
    """Get reference data as JSON"""
    ref = Reference.query.get_or_404(id)
    return jsonify({
        'success': True,
        'reference': {
            'id': ref.id,
            'project_id': ref.project_id,
            'client': ref.client,
            'agency': ref.agency,
            'project_name': ref.project_name,
            'contract_number': ref.contract_number,
            'project_id_number': ref.project_id_number,
            'evaluation_date': ref.evaluation_date.isoformat() if ref.evaluation_date else None,
            'evaluation_period': ref.evaluation_period,
            'final_score': ref.final_score,
            'schedule_score': ref.schedule_score,
            'quality_score': ref.quality_score,
            'responsiveness_score': ref.responsiveness_score,
            'key_staff_score': ref.key_staff_score,
            'dbe_score': ref.dbe_score,
            'pm_performance_score': ref.pm_performance_score,
            'score_summary': ref.score_summary,
            'quotes': ref.quotes,
            'evaluator_name': ref.evaluator_name,
            'evaluator_title': ref.evaluator_title,
            'consultant_pm': ref.consultant_pm,
            'firm': ref.firm,
            'services_description': ref.services_description,
            'activities_evaluated': ref.activities_evaluated,
            'personnel_tags': ref.personnel_tags,
            'reference_type': ref.reference_type,
            'is_final_evaluation': ref.is_final_evaluation,
            'pdf_filename': ref.pdf_filename
        }
    })


@app.route('/references/<int:id>/upload-pdf', methods=['POST'])
@login_required
def upload_reference_pdf(id):
    """Upload a PDF for a reference"""
    ref = Reference.query.get_or_404(id)
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'}), 400
    
    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'success': False, 'error': 'Only PDF files are allowed'}), 400
    
    if ref.pdf_object_key:
        try:
            storage = get_storage_client()
            if storage:
                storage.delete(ref.pdf_object_key)
        except Exception:
            pass
    
    filename = secure_filename(file.filename)
    object_key = f"references/{ref.id}/{uuid.uuid4().hex}_{filename}"
    
    try:
        storage = get_storage_client()
        if not storage:
            return jsonify({'success': False, 'error': 'Object storage not configured'}), 500
        file_data = file.read()
        storage.upload_from_bytes(object_key, file_data)
        
        ref.pdf_filename = filename
        ref.pdf_object_key = object_key
        db.session.commit()
        
        return jsonify({'success': True, 'filename': filename})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/references/<int:id>/download-pdf')
@login_required
def download_reference_pdf(id):
    """Download the PDF for a reference"""
    ref = Reference.query.get_or_404(id)
    
    if not ref.pdf_object_key:
        flash('No PDF file attached to this reference', 'error')
        return redirect(url_for('references_page'))
    
    try:
        storage = get_storage_client()
        if not storage:
            flash('Object storage not configured', 'error')
            return redirect(url_for('references_page'))
        file_data = storage.download_as_bytes(ref.pdf_object_key)
        return send_file(
            io.BytesIO(file_data),
            download_name=ref.pdf_filename or 'reference.pdf',
            as_attachment=True,
            mimetype='application/pdf'
        )
    except Exception as e:
        flash(f'Error downloading file: {str(e)}', 'error')
        return redirect(url_for('references_page'))


@app.route('/references/upload-and-parse', methods=['POST'])
@login_required
def upload_and_parse_reference_pdf():
    """Upload a PDF and use AI to extract reference data"""
    from gemini_service import client
    from datetime import datetime as dt
    
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': 'No file selected'}), 400
    
    if not file.filename.lower().endswith('.pdf'):
        return jsonify({'success': False, 'error': 'Only PDF files are allowed'}), 400
    
    filename = secure_filename(file.filename)
    file_data = file.read()
    
    try:
        from document_parser import extract_text_from_file
        text = extract_text_from_file(filename, file_data)
    except Exception as e:
        return jsonify({'success': False, 'error': f'Failed to extract text: {str(e)}'}), 500
    
    prompt = """Analyze this consultant performance evaluation document and extract the following information in JSON format:

{
    "client": "The client/agency name (e.g., SCDOT)",
    "agency": "Full agency name if different from client",
    "project_name": "Name of the project",
    "contract_number": "Contract number(s)",
    "project_id_number": "Project ID number",
    "evaluation_date": "Date in YYYY-MM-DD format",
    "evaluation_period": "April or October or other period",
    "final_score": number (0-10 scale),
    "schedule_score": number or null,
    "quality_score": number or null,
    "responsiveness_score": number or null,
    "key_staff_score": number or null,
    "dbe_score": number or null,
    "pm_performance_score": number or null,
    "score_summary": "Brief summary of the performance explanation section",
    "quotes": "Extract 2-3 notable quotes from the evaluation that praise the work",
    "evaluator_name": "Name of the person who completed the evaluation",
    "evaluator_title": "Title of the evaluator if available",
    "consultant_pm": "Consultant Project Manager name",
    "firm": "Consultant firm name",
    "services_description": "Description of project services",
    "activities_evaluated": "Activities being evaluated",
    "is_final_evaluation": true or false
}

Return ONLY valid JSON, no other text.

Document text:
""" + text[:15000]
    
    try:
        response = client.models.generate_content(
            model="gemini-2.5-flash",
            contents=prompt
        )
        
        response_text = response.text.strip()
        if response_text.startswith('```'):
            response_text = response_text.split('\n', 1)[1]
            response_text = response_text.rsplit('```', 1)[0]
        
        data = json.loads(response_text)
        
        ref = Reference(
            client=data.get('client', ''),
            agency=data.get('agency', ''),
            project_name=data.get('project_name', ''),
            contract_number=data.get('contract_number', ''),
            project_id_number=data.get('project_id_number', ''),
            final_score=data.get('final_score'),
            schedule_score=data.get('schedule_score'),
            quality_score=data.get('quality_score'),
            responsiveness_score=data.get('responsiveness_score'),
            key_staff_score=data.get('key_staff_score'),
            dbe_score=data.get('dbe_score'),
            pm_performance_score=data.get('pm_performance_score'),
            score_summary=data.get('score_summary', ''),
            quotes=data.get('quotes', ''),
            evaluator_name=data.get('evaluator_name', ''),
            evaluator_title=data.get('evaluator_title', ''),
            consultant_pm=data.get('consultant_pm', ''),
            firm=data.get('firm', ''),
            services_description=data.get('services_description', ''),
            activities_evaluated=data.get('activities_evaluated', ''),
            evaluation_period=data.get('evaluation_period', ''),
            is_final_evaluation=data.get('is_final_evaluation', False),
            reference_type='evaluation'
        )
        
        eval_date_str = data.get('evaluation_date', '')
        if eval_date_str:
            try:
                ref.evaluation_date = dt.strptime(eval_date_str, '%Y-%m-%d').date()
            except ValueError:
                pass
        
        db.session.add(ref)
        db.session.flush()
        
        object_key = f"references/{ref.id}/{uuid.uuid4().hex}_{filename}"
        storage = get_storage_client()
        if storage:
            storage.upload_from_bytes(object_key, file_data)
        else:
            object_key = None
        
        ref.pdf_filename = filename
        ref.pdf_object_key = object_key
        db.session.commit()
        
        return jsonify({
            'success': True,
            'id': ref.id,
            'data': data
        })
    except json.JSONDecodeError as e:
        return jsonify({'success': False, 'error': f'Failed to parse AI response: {str(e)}'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/proposals/<int:proposal_id>/link-reference', methods=['POST'])
@login_required
def link_reference_to_proposal(proposal_id):
    """Link a reference to a proposal"""
    proposal = Proposal.query.get(proposal_id)
    if not proposal:
        return jsonify({'success': False, 'error': 'Proposal not found'}), 404
    
    data = request.get_json()
    reference_id = data.get('reference_id')
    
    ref = Reference.query.get(reference_id)
    if not ref:
        return jsonify({'success': False, 'error': 'Reference not found'}), 404
    
    existing = ProposalLinkedReference.query.filter_by(proposal_id=proposal_id, reference_id=reference_id).first()
    if existing:
        return jsonify({'success': False, 'error': 'Reference already linked to this proposal'}), 400
    
    max_order = db.session.query(db.func.max(ProposalLinkedReference.display_order)).filter_by(proposal_id=proposal_id).scalar() or 0
    
    link = ProposalLinkedReference(
        proposal_id=proposal_id,
        reference_id=reference_id,
        display_order=max_order + 1
    )
    db.session.add(link)
    db.session.commit()
    
    return jsonify({'success': True, 'message': 'Reference linked to proposal'})


@app.route('/api/proposals/<int:proposal_id>/unlink-reference/<int:reference_id>', methods=['DELETE'])
@login_required
def unlink_reference_from_proposal(proposal_id, reference_id):
    """Unlink a reference from a proposal"""
    link = ProposalLinkedReference.query.filter_by(proposal_id=proposal_id, reference_id=reference_id).first()
    if not link:
        return jsonify({'success': False, 'error': 'Link not found'}), 404
    
    db.session.delete(link)
    db.session.commit()
    
    return jsonify({'success': True})


@app.route('/api/proposals/<int:proposal_id>/linked-references')
@login_required
def get_proposal_linked_references(proposal_id):
    """Get all references linked to a proposal"""
    links = ProposalLinkedReference.query.filter_by(proposal_id=proposal_id).order_by(ProposalLinkedReference.display_order).all()
    
    refs = []
    for link in links:
        ref = link.reference
        refs.append({
            'id': ref.id,
            'client': ref.client,
            'project_name': ref.project_name,
            'final_score': ref.final_score,
            'evaluation_date': ref.evaluation_date.isoformat() if ref.evaluation_date else None,
            'score_summary': ref.score_summary,
            'quotes': ref.quotes,
            'consultant_pm': ref.consultant_pm,
            'firm': ref.firm,
            'personnel_tags': ref.personnel_tags,
            'pdf_filename': ref.pdf_filename
        })
    
    return jsonify({'success': True, 'references': refs})
