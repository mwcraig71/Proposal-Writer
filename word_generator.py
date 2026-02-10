"""
SF330 Word Document Generator
Generates filled SF330 Word documents from proposal data
"""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy
import io
import os

TEMPLATE_DIR = 'templates/sf330_word'

def get_cell_text(table, row, col):
    """Safely get cell text"""
    try:
        return table.rows[row].cells[col].text.strip()
    except:
        return ""

def set_cell_text(table, row, col, text):
    """Set cell text, preserving formatting"""
    try:
        if row >= len(table.rows) or col >= len(table.rows[row].cells):
            return
        cell = table.rows[row].cells[col]
        text_str = str(text) if text else ""
        if cell.paragraphs:
            for run in cell.paragraphs[0].runs:
                run.text = ""
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].text = text_str
            else:
                cell.paragraphs[0].add_run(text_str)
        else:
            cell.text = text_str
    except Exception as e:
        print(f"Error setting cell [{row}][{col}]: {e}")

def generate_section_a_c(proposal, firms):
    """
    Generate Section A (Contract Info), B (POC), C (Proposed Team)
    Returns: Document object
    """
    doc = Document(os.path.join(TEMPLATE_DIR, 'section_a_c.docx'))
    table = doc.tables[0]
    
    # Section A - Contract Information
    # Row 3: Title and Location (spans full width, data goes in second row after label)
    title_location = f"{proposal.contract_title or ''}"
    if proposal.contract_location:
        title_location += f" ({proposal.contract_location})"
    
    # Find the cells for each field by examining the structure
    # Row 3 has the label, need to find where to put data
    # Based on SF330 structure, data typically goes in cells below or adjacent to labels
    
    # For this template, we need to identify exact cell positions
    # Row 3: Title and Location label - data goes below
    # Row 4: Public Notice Date | Solicitation Number
    
    # Set contract title/location (in the row content area)
    set_cell_text(table, 3, 0, f"1.  TITLE AND LOCATION (City and State)\n{title_location}")
    
    # Row 4: Public Notice Date and Solicitation Number
    if proposal.public_notice_date:
        if hasattr(proposal.public_notice_date, 'strftime'):
            notice_date = proposal.public_notice_date.strftime('%m/%d/%Y')
        else:
            notice_date = str(proposal.public_notice_date)
    else:
        notice_date = ""
    set_cell_text(table, 4, 0, f"2.  PUBLIC NOTICE DATE\n{notice_date}")
    set_cell_text(table, 4, 8, f"3.  SOLICITATION OR PROJECT NUMBER\n{proposal.solicitation_number or ''}")
    
    # Section B - Point of Contact (primary firm)
    if proposal.firm:
        firm = proposal.firm
        poc_name = getattr(firm, 'point_of_contact_name', '') or getattr(firm, 'contact_name', '') or ''
        poc_title = getattr(firm, 'point_of_contact_title', '') or ''
        poc_display = f"{poc_name}, {poc_title}" if poc_title else poc_name
        # Row 5-6: POC section
        set_cell_text(table, 6, 0, f"4.  NAME AND TITLE\n{poc_display}")
        set_cell_text(table, 7, 0, f"5.  NAME OF FIRM\n{firm.name or ''}")
        set_cell_text(table, 8, 0, f"6.  TELEPHONE NUMBER\n{getattr(firm, 'phone', '') or ''}")
        set_cell_text(table, 8, 4, f"7.  FAX NUMBER\n{getattr(firm, 'fax', '') or ''}")
        set_cell_text(table, 8, 7, f"8.  E-MAIL ADDRESS\n{getattr(firm, 'email', '') or ''}")
    
    # Section C - Proposed Team (rows 12-17 typically for team members)
    # The table has rows for up to 6 firms (a through f)
    team_start_row = 12  # Approximate start for team entries
    
    for idx, firm_data in enumerate(firms[:6]):
        row_offset = team_start_row + idx
        if row_offset < len(table.rows):
            firm = firm_data.get('firm')
            role = firm_data.get('role', '')
            is_prime = firm_data.get('is_prime', False)
            
            if firm:
                # Set firm name
                address = f"{firm.city}, {firm.state}" if firm.city and firm.state else ""
                set_cell_text(table, row_offset, 4, firm.name or '')
                set_cell_text(table, row_offset, 7, address)
                set_cell_text(table, row_offset, 9, role)
    
    return doc


def generate_section_e(employee, proposal, project_experiences):
    """
    Generate Section E (Resume) for one employee using SF330 Section E template.
    Supports unlimited project entries by dynamically cloning/removing project blocks.
    Returns: Document object
    """
    from docx.oxml.ns import qn

    template_source = None
    try:
        from replit.object_storage import Client as StorageClient
        client = StorageClient()
        template_bytes = client.download_as_bytes('templates/sf330_resume_template_custom.docx')
        if template_bytes:
            template_source = io.BytesIO(template_bytes)
    except:
        pass
    if template_source is None:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        for path in [
            os.path.join(base_dir, 'attached_assets', 'sf330_section_e_template.docx'),
            os.path.join(base_dir, 'attached_assets', '330_Section_E_Standards_template_1770398969209.docx'),
        ]:
            if os.path.exists(path):
                template_source = path
                break
    if template_source is None:
        template_source = os.path.join(TEMPLATE_DIR, 'section_e.docx')

    if isinstance(template_source, io.BytesIO):
        template_source.seek(0)
        doc = Document(template_source)
    else:
        doc = Document(template_source)

    num_projects = len(project_experiences)

    def _get_cell_text_e(tc_element):
        text = ''
        for p in tc_element.findall(qn('w:p')):
            for r in p.findall(qn('w:r')):
                for t in r.findall(qn('w:t')):
                    text += (t.text or '')
        return text

    def _replace_text_in_row_e(row_el, old_text, new_text):
        for tc in row_el.findall(qn('w:tc')):
            for p_el in tc.findall(qn('w:p')):
                runs = p_el.findall(qn('w:r'))
                if not runs:
                    continue
                full = ''
                for r in runs:
                    for t in r.findall(qn('w:t')):
                        full += (t.text or '')
                if old_text not in full:
                    continue
                new_full = full.replace(old_text, new_text)
                first_set = False
                for r in runs:
                    for t in r.findall(qn('w:t')):
                        if not first_set:
                            t.text = new_full
                            t.set(qn('xml:space'), 'preserve')
                            first_set = True
                        else:
                            t.text = ''
                            t.set(qn('xml:space'), 'preserve')

    def _get_alpha_label_e(idx):
        result = ''
        n = idx
        while True:
            result = chr(ord('a') + n % 26) + result
            n = n // 26 - 1
            if n < 0:
                break
        return result + '.'

    table = doc.tables[0] if doc.tables else None
    if table is not None and num_projects != 5:
        tbl_element = table._tbl
        all_tr = list(tbl_element.findall(qn('w:tr')))

        label_rows = []
        for i, tr in enumerate(all_tr):
            first_tc = tr.find(qn('w:tc'))
            if first_tc is not None:
                text = _get_cell_text_e(first_tc).strip()
                label = text.rstrip('.')
                if label and len(label) <= 2 and label.isalpha() and label.islower():
                    label_rows.append(i)

        if len(label_rows) >= 2:
            label_in_block_offset = 0
            if label_rows[0] > 0:
                prev_text = ''
                for tc in all_tr[label_rows[0] - 1].findall(qn('w:tc')):
                    prev_text += _get_cell_text_e(tc)
                if '(1)' in prev_text:
                    label_in_block_offset = 1

            block_starts = [lr - label_in_block_offset for lr in label_rows]
            rows_per_block = block_starts[1] - block_starts[0]
            template_block_count = len(block_starts)
            last_block_start_idx = block_starts[-1]
            last_block_project_num = template_block_count

            last_block_end = last_block_start_idx + rows_per_block
            footer_tr_elements = all_tr[last_block_end:]
            insert_before = footer_tr_elements[0] if footer_tr_elements else None

            if num_projects > template_block_count:
                template_block_rows = all_tr[last_block_start_idx:last_block_end]
                label_row_in_block = template_block_rows[label_in_block_offset]
                old_label_text = _get_cell_text_e(label_row_in_block.find(qn('w:tc'))).strip()

                for extra_idx in range(template_block_count, num_projects):
                    project_num = extra_idx + 1
                    new_label = _get_alpha_label_e(extra_idx)

                    for row_el in template_block_rows:
                        new_row = copy.deepcopy(row_el)
                        _replace_text_in_row_e(
                            new_row,
                            f'PROJECT_EXPERIENCE_{last_block_project_num}',
                            f'PROJECT_EXPERIENCE_{project_num}'
                        )
                        first_tc = new_row.find(qn('w:tc'))
                        if first_tc is not None:
                            cell_text = _get_cell_text_e(first_tc).strip()
                            if cell_text == old_label_text:
                                for p_el in first_tc.findall(qn('w:p')):
                                    for r in p_el.findall(qn('w:r')):
                                        for t in r.findall(qn('w:t')):
                                            if t.text and t.text.strip().rstrip('.') == old_label_text.rstrip('.'):
                                                t.text = t.text.replace(old_label_text.rstrip('.'), new_label.rstrip('.')).replace(old_label_text, new_label)

                        if insert_before is not None:
                            tbl_element.insert(list(tbl_element).index(insert_before), new_row)
                        else:
                            tbl_element.append(new_row)

            elif num_projects < template_block_count:
                for block_idx in range(template_block_count - 1, max(num_projects - 1, -1), -1):
                    if block_idx >= len(block_starts):
                        continue
                    block_start = block_starts[block_idx]
                    rows_to_remove = all_tr[block_start:block_start + rows_per_block]
                    for tr in rows_to_remove:
                        if tr.getparent() is not None:
                            tbl_element.remove(tr)

    def _replace_placeholder_in_element(element, placeholder, value):
        for p_el in element.findall(qn('w:p')):
            runs = p_el.findall(qn('w:r'))
            if not runs:
                continue
            full = ''
            for r in runs:
                for t in r.findall(qn('w:t')):
                    full += (t.text or '')
            if placeholder not in full:
                continue
            new_full = full.replace(placeholder, value)
            first_set = False
            for r in runs:
                for t in r.findall(qn('w:t')):
                    if not first_set:
                        t.text = new_full
                        t.set(qn('xml:space'), 'preserve')
                        first_set = True
                    else:
                        t.text = ''
                        t.set(qn('xml:space'), 'preserve')

    role = ''
    if hasattr(employee, 'proposal_role'):
        role = employee.proposal_role or ''

    other_quals = ''
    if hasattr(employee, 'training') and employee.training:
        other_quals += employee.training
    if hasattr(employee, 'other_qualifications') and employee.other_qualifications:
        if other_quals:
            other_quals += '\n'
        other_quals += employee.other_qualifications

    firm_name = ''
    firm_location = ''
    if proposal.firm:
        firm_name = proposal.firm.name or ''
        parts = [proposal.firm.city or '', proposal.firm.state or '']
        firm_location = ', '.join(p for p in parts if p)

    placeholders = {
        '{{EMPLOYEE_NAME}}': employee.name or '',
        '{{EMPLOYEE_ROLE}}': role,
        '{{YEARS_EXPERIENCE_TOTAL}}': str(employee.years_experience or '') if hasattr(employee, 'years_experience') else '',
        '{{YEARS_EXPERIENCE_FIRM}}': str(employee.years_with_firm or '') if hasattr(employee, 'years_with_firm') else '',
        '{{FIRM_NAME}}': firm_name,
        '{{FIRM_LOCATION}}': firm_location,
        '{{FIRM_NAME_LOCATION}}': f"{firm_name}, {firm_location}" if firm_location else firm_name,
        '{{EDUCATION}}': employee.education or '',
        '{{REGISTRATIONS}}': employee.registrations or '',
        '{{OTHER_QUALIFICATIONS}}': other_quals,
    }

    for idx, exp in enumerate(project_experiences):
        proj_num = idx + 1
        title = exp.title or ''
        location = exp.location or ''
        year = str(exp.year_completed or '')
        description = exp.description or ''
        exp_role = exp.role or ''

        placeholders[f'{{{{PROJECT_EXPERIENCE_{proj_num}_TITLE}}}}'] = title
        placeholders[f'{{{{PROJECT_EXPERIENCE_{proj_num}_LOCATION}}}}'] = location
        placeholders[f'{{{{PROJECT_EXPERIENCE_{proj_num}_YEAR}}}}'] = year
        placeholders[f'{{{{PROJECT_EXPERIENCE_{proj_num}_DESCRIPTION}}}}'] = description
        placeholders[f'{{{{PROJECT_EXPERIENCE_{proj_num}_ROLE}}}}'] = exp_role
        full_block = title
        if location:
            full_block += f", {location}"
        if year:
            full_block += f" ({year})"
        if exp_role:
            full_block += f"\nRole: {exp_role}"
        if description:
            full_block += f"\n{description}"
        placeholders[f'{{{{PROJECT_EXPERIENCE_{proj_num}}}}}'] = full_block

    if table is not None:
        tbl_element = table._tbl
        for tr in tbl_element.findall(qn('w:tr')):
            for tc in tr.findall(qn('w:tc')):
                for placeholder, value in placeholders.items():
                    _replace_placeholder_in_element(tc, placeholder, str(value))

    for para in doc.paragraphs:
        for placeholder, value in placeholders.items():
            runs = para._element.findall(qn('w:r'))
            if not runs:
                continue
            full = ''
            for r in runs:
                for t in r.findall(qn('w:t')):
                    full += (t.text or '')
            if placeholder not in full:
                continue
            new_full = full.replace(placeholder, str(value))
            first_set = False
            for r in runs:
                for t in r.findall(qn('w:t')):
                    if not first_set:
                        t.text = new_full
                        t.set(qn('xml:space'), 'preserve')
                        first_set = True
                    else:
                        t.text = ''
                        t.set(qn('xml:space'), 'preserve')

    import re
    project_exp_pattern = re.compile(r'\{\{PROJECT_EXPERIENCE_\d+(_[A-Z]+)?\}\}')
    if table is not None:
        tbl_element = table._tbl
        for tr in tbl_element.findall(qn('w:tr')):
            for tc in tr.findall(qn('w:tc')):
                for p_el in tc.findall(qn('w:p')):
                    for r in p_el.findall(qn('w:r')):
                        for t in r.findall(qn('w:t')):
                            if t.text and project_exp_pattern.search(t.text):
                                t.text = project_exp_pattern.sub('', t.text)

    return doc


def generate_section_f(project, proposal, project_key_number, firms_involved=None):
    """
    Generate Section F (Example Project) for one project
    Returns: Document object
    """
    doc = Document(os.path.join(TEMPLATE_DIR, 'section_f.docx'))
    
    # Table 0: Main project info
    table = doc.tables[0]
    
    # Project Key Number (Block 20)
    set_cell_text(table, 0, 7, str(project_key_number))
    
    # Title and Location (Block 21)
    title_loc = project.title or ''
    if project.location:
        title_loc += f", {project.location}"
    set_cell_text(table, 1, 0, title_loc)
    
    # Year Completed (Block 22) - use professional services year
    year_completed = getattr(project, 'year_completed_professional', '') or getattr(project, 'year_completed', '') or ''
    set_cell_text(table, 1, 5, str(year_completed))
    
    # Project Owner Info (Block 23)
    set_cell_text(table, 3, 0, getattr(project, 'owner_name', '') or '')
    set_cell_text(table, 3, 3, getattr(project, 'owner_contact_name', '') or '')
    set_cell_text(table, 3, 6, getattr(project, 'owner_contact_phone', '') or '')
    
    # Brief Description (Block 24) - large text area
    description = project.brief_description or ''
    set_cell_text(table, 4, 0, description)
    
    # Table 1: Firms involved (Block 25)
    if len(doc.tables) > 1 and firms_involved:
        firms_table = doc.tables[1]
        for idx, firm_info in enumerate(firms_involved[:6]):
            row_idx = idx + 1  # Skip header row
            if row_idx < len(firms_table.rows):
                firm = firm_info.get('firm')
                if firm:
                    set_cell_text(firms_table, row_idx, 1, firm.name or '')
                    loc = f"{firm.city}, {firm.state}" if firm.city else ""
                    set_cell_text(firms_table, row_idx, 3, loc)
                    set_cell_text(firms_table, row_idx, 5, firm_info.get('role', ''))
    
    return doc


def _clone_row(table, source_row_idx):
    """Clone a table row and insert it after the source row. Returns the new row element."""
    from lxml import etree
    source_tr = table.rows[source_row_idx]._tr
    new_tr = copy.deepcopy(source_tr)
    source_tr.addnext(new_tr)
    return new_tr


def _replace_cell_text(cell, old_text, new_text):
    """Replace text in a cell, handling runs properly."""
    for para in cell.paragraphs:
        full = para.text
        if old_text not in full:
            continue
        replaced = False
        for run in para.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                replaced = True
        if not replaced and para.runs:
            new_full = full.replace(old_text, new_text)
            para.runs[0].text = new_full
            for run in para.runs[1:]:
                run.text = ''


def _replace_all_placeholders_in_doc(doc, placeholders):
    """Replace all placeholders throughout the entire document."""
    seen_tcs = set()
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                tc = cell._tc
                if tc in seen_tcs:
                    continue
                seen_tcs.add(tc)
                for para in cell.paragraphs:
                    full_text = para.text
                    if '{{' not in full_text:
                        continue
                    for placeholder, value in placeholders.items():
                        if placeholder in full_text:
                            replaced = False
                            for run in para.runs:
                                if placeholder in run.text:
                                    run.text = run.text.replace(placeholder, value)
                                    replaced = True
                            if not replaced and para.runs:
                                new_text = full_text
                                for ph, val in placeholders.items():
                                    new_text = new_text.replace(ph, val)
                                para.runs[0].text = new_text
                                for run in para.runs[1:]:
                                    run.text = ''
                                break
                            full_text = para.text


def generate_section_g(employees_with_roles, projects, matrix, template_doc=None):
    """
    Generate Section G (Key Personnel Participation Matrix) using template.
    Dynamically expands rows to fit any number of employees.
    employees_with_roles: list of dicts with 'employee' and 'role' keys
    projects: list of project objects (max 10)
    matrix: dict of {employee_id: set(project_ids)}
    template_doc: optional Document object (custom template); falls back to default
    Returns: Document object
    """
    if template_doc:
        doc = template_doc
    else:
        default_template = os.path.join(os.path.dirname(__file__), 'attached_assets',
                                        '330_Section_G_Standards_08-2026-mwc_1770486074175.docx')
        if os.path.exists(default_template):
            doc = Document(default_template)
        else:
            doc = Document(os.path.join(TEMPLATE_DIR, 'section_g.docx'))

    table = doc.tables[0]
    TEMPLATE_EMP_ROWS = 13
    num_employees = len(employees_with_roles)

    first_emp_row_idx = None
    for row_idx, row in enumerate(table.rows):
        seen = set()
        for cell in row.cells:
            cid = id(cell._tc)
            if cid in seen:
                continue
            seen.add(cid)
            if '{{EMPLOYEE_NAME_1}}' in cell.text:
                first_emp_row_idx = row_idx
                break
        if first_emp_row_idx is not None:
            break

    if first_emp_row_idx is None:
        first_emp_row_idx = 3

    last_template_emp_row_idx = first_emp_row_idx + TEMPLATE_EMP_ROWS - 1

    if num_employees > TEMPLATE_EMP_ROWS:
        extra_rows_needed = num_employees - TEMPLATE_EMP_ROWS
        for i in range(extra_rows_needed):
            _clone_row(table, last_template_emp_row_idx)
    elif num_employees < TEMPLATE_EMP_ROWS:
        from lxml import etree
        rows_to_remove = TEMPLATE_EMP_ROWS - num_employees
        for i in range(rows_to_remove):
            remove_idx = first_emp_row_idx + num_employees
            if remove_idx < len(table.rows):
                tr = table.rows[remove_idx]._tr
                tr.getparent().remove(tr)

    placeholders = {}

    for emp_idx, emp_data in enumerate(employees_with_roles):
        emp_num = emp_idx + 1
        employee = emp_data.get('employee')
        role = emp_data.get('role', '')
        placeholders[f'{{{{EMPLOYEE_NAME_{emp_num}}}}}'] = employee.name if employee else ''
        placeholders[f'{{{{EMPLOYEE_ROLE_{emp_num}}}}}'] = role

        if employee:
            emp_projects = matrix.get(employee.id, set())
            for proj_idx, project in enumerate(projects[:10]):
                proj_num = proj_idx + 1
                mark = 'X' if project.id in emp_projects else ''
                placeholders[f'{{{{EMPLOYEE_PROJECT_{emp_num}_{proj_num}}}}}'] = mark

        for proj_num in range(len(projects) + 1, 11):
            placeholders[f'{{{{EMPLOYEE_PROJECT_{emp_num}_{proj_num}}}}}'] = ''

    if num_employees > TEMPLATE_EMP_ROWS:
        src_num = TEMPLATE_EMP_ROWS
        for emp_num in range(TEMPLATE_EMP_ROWS + 1, num_employees + 1):
            row_idx = first_emp_row_idx + emp_num - 1
            if row_idx < len(table.rows):
                row = table.rows[row_idx]
                seen_tcs = set()
                for cell in row.cells:
                    tc = cell._tc
                    if tc in seen_tcs:
                        continue
                    seen_tcs.add(tc)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.text = run.text.replace(
                                f'{{{{EMPLOYEE_NAME_{src_num}}}}}',
                                f'{{{{EMPLOYEE_NAME_{emp_num}}}}}'
                            )
                            run.text = run.text.replace(
                                f'{{{{EMPLOYEE_ROLE_{src_num}}}}}',
                                f'{{{{EMPLOYEE_ROLE_{emp_num}}}}}'
                            )
                            for pn in range(1, 11):
                                run.text = run.text.replace(
                                    f'{{{{EMPLOYEE_PROJECT_{src_num}_{pn}}}}}',
                                    f'{{{{EMPLOYEE_PROJECT_{emp_num}_{pn}}}}}'
                                )

    for proj_idx, project in enumerate(projects[:10]):
        proj_num = proj_idx + 1
        placeholders[f'{{{{PROJECT_TITLE_{proj_num}}}}}'] = project.title or ''

    for proj_num in range(len(projects) + 1, 11):
        placeholders[f'{{{{PROJECT_TITLE_{proj_num}}}}}'] = ''

    _replace_all_placeholders_in_doc(doc, placeholders)

    return doc


def generate_section_h(proposal, additional_info=""):
    """
    Generate Section H (Additional Information) and I (Signature)
    Returns: Document object
    """
    doc = Document(os.path.join(TEMPLATE_DIR, 'section_h_i.docx'))
    
    # Section H is mostly in paragraphs, not tables
    # Find and fill the additional info area
    for para in doc.paragraphs:
        if "30." in para.text or "ADDITIONAL INFORMATION" in para.text:
            # Add content after this paragraph
            continue
    
    # The written sections go here
    content = proposal.written_sections or additional_info or ""
    if doc.paragraphs:
        # Find appropriate paragraph to add content
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip() == "" and i > 5:  # Find empty paragraph after header
                para.text = content
                break
    
    # Table 1 has signature info
    if len(doc.tables) > 1:
        sig_table = doc.tables[1]
        # Signature date
        from datetime import datetime
        set_cell_text(sig_table, 0, 1, datetime.now().strftime('%m/%d/%Y'))
        
        # Name and title
        if proposal.firm:
            poc_name = getattr(proposal.firm, 'point_of_contact_name', '') or ''
            poc_title = getattr(proposal.firm, 'point_of_contact_title', '') or ''
            poc_display = f"{poc_name}, {poc_title}" if poc_title else poc_name
            set_cell_text(sig_table, 2, 0, poc_display)
    
    return doc


def generate_part_ii(firm):
    """
    Generate Part II (General Qualifications)
    Returns: Document object
    """
    doc = Document(os.path.join(TEMPLATE_DIR, 'part_ii.docx'))
    table = doc.tables[0]
    
    # Row 1: Firm name | Year Established | UEI
    set_cell_text(table, 2, 0, firm.name or '')  # Block 2a
    set_cell_text(table, 2, 7, str(firm.year_established or ''))  # Block 3
    set_cell_text(table, 2, 10, firm.uei or '')  # Block 4
    
    # Row 3-5: Street, City, State, ZIP
    set_cell_text(table, 4, 0, firm.street_address or '')  # Block 2b
    set_cell_text(table, 6, 0, firm.city or '')  # Block 2c
    set_cell_text(table, 6, 4, firm.state or '')  # Block 2d
    set_cell_text(table, 6, 7, firm.zip_code or '')  # Block 2e
    
    # Ownership type
    set_cell_text(table, 4, 7, firm.ownership_type or '')  # Block 5a
    
    # Point of Contact (Block 6)
    poc_name = getattr(firm, 'point_of_contact_name', '') or ''
    poc_title = getattr(firm, 'point_of_contact_title', '') or ''
    poc_display = f"{poc_name}, {poc_title}" if poc_title else poc_name
    set_cell_text(table, 10, 0, poc_display)  # Block 6a
    set_cell_text(table, 12, 0, getattr(firm, 'phone', '') or '')  # Block 6b
    set_cell_text(table, 12, 7, getattr(firm, 'email', '') or '')  # Block 6c
    
    # Former firm names (Block 8) - if applicable
    # Employee disciplines (Block 9) - would need employee data
    # Experience profile (Block 10) - would need profile data
    
    return doc


def generate_full_sf330_word(proposal, selected_employees, selected_projects, matrix):
    """
    Generate complete SF330 as multiple Word documents or combined document
    Returns: dict with section names and document bytes
    """
    from models import Firm
    
    documents = {}
    
    # Get primary firm
    firm = proposal.firm
    
    # Section A/C - Contract Info and Team
    firms_data = [{'firm': firm, 'role': 'Prime', 'is_prime': True}] if firm else []
    doc_a = generate_section_a_c(proposal, firms_data)
    
    buffer = io.BytesIO()
    doc_a.save(buffer)
    documents['section_a_c'] = buffer.getvalue()
    
    # Section E - One per employee
    section_e_docs = []
    for emp_data in selected_employees:
        employee = emp_data.get('employee')
        experiences = emp_data.get('experiences', [])
        role = emp_data.get('role', '')
        
        if employee:
            employee.proposal_role = role
            doc_e = generate_section_e(employee, proposal, experiences)
            buffer = io.BytesIO()
            doc_e.save(buffer)
            section_e_docs.append({
                'name': employee.name,
                'data': buffer.getvalue()
            })
    documents['section_e'] = section_e_docs
    
    # Section F - One per project
    section_f_docs = []
    for idx, proj_data in enumerate(selected_projects):
        project = proj_data.get('project')
        if project:
            doc_f = generate_section_f(project, proposal, idx + 1)
            buffer = io.BytesIO()
            doc_f.save(buffer)
            section_f_docs.append({
                'name': project.title,
                'data': buffer.getvalue()
            })
    documents['section_f'] = section_f_docs
    
    # Section G - Matrix
    employees_with_roles = [
        {'employee': ed.get('employee'), 'role': ed.get('role', '')}
        for ed in selected_employees
    ]
    projects = [pd.get('project') for pd in selected_projects if pd.get('project')]
    
    doc_g = generate_section_g(employees_with_roles, projects, matrix)
    buffer = io.BytesIO()
    doc_g.save(buffer)
    documents['section_g'] = buffer.getvalue()
    
    # Section H - Additional Info
    doc_h = generate_section_h(proposal)
    buffer = io.BytesIO()
    doc_h.save(buffer)
    documents['section_h_i'] = buffer.getvalue()
    
    # Part II - Firm Qualifications
    if firm:
        doc_ii = generate_part_ii(firm)
        buffer = io.BytesIO()
        doc_ii.save(buffer)
        documents['part_ii'] = buffer.getvalue()
    
    return documents


def combine_documents_as_zip(documents):
    """
    Package multiple Word documents into a ZIP file
    Returns: bytes of ZIP archive
    """
    import zipfile
    
    buffer = io.BytesIO()
    
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
        sections_order = ['section_a_c', 'section_e', 'section_f', 'section_g', 'section_h_i', 'part_ii']
        
        for section in sections_order:
            if section not in documents:
                continue
                
            data = documents[section]
            
            if isinstance(data, list):
                # Multiple docs (Section E, F)
                for idx, item in enumerate(data, 1):
                    name = item.get('name', f'item_{idx}')
                    # Sanitize filename
                    safe_name = "".join(c for c in name if c.isalnum() or c in (' ', '-', '_')).strip()[:50]
                    filename = f"{section}_{idx:02d}_{safe_name}.docx"
                    zf.writestr(filename, item['data'])
            else:
                # Single doc
                zf.writestr(f"{section}.docx", data)
    
    buffer.seek(0)
    return buffer.getvalue()


def generate_simple_sf330(proposal, employees_data, projects_data, firms_data, matrix_data):
    """
    Generate a simple Word document with all SF330 info (no template)
    Returns: bytes of the Word document
    """
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Set heading styles to left align
    for i in range(4):
        try:
            h_style = doc.styles[f'Heading {i+1}']
            h_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        except:
            pass
    
    def add_heading(text, level=1):
        h = doc.add_heading(text, level=level)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return h
    
    def add_field(label, value):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.add_run(f"{label}: ").bold = True
        p.add_run(str(value) if value else "")
        return p
    
    # Title
    title = doc.add_heading('SF330 - ARCHITECT-ENGINEER QUALIFICATIONS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Part I - Contract-Specific Qualifications
    add_heading('PART I - CONTRACT-SPECIFIC QUALIFICATIONS', 1)
    
    # Section A - Contract Information
    add_heading('A. CONTRACT INFORMATION', 2)
    add_field('1. Title and Location', f"{proposal.contract_title or ''}, {proposal.contract_location or ''}")
    
    if proposal.public_notice_date:
        if hasattr(proposal.public_notice_date, 'strftime'):
            notice_date = proposal.public_notice_date.strftime('%m/%d/%Y')
        else:
            notice_date = str(proposal.public_notice_date)
    else:
        notice_date = ""
    add_field('2. Public Notice Date', notice_date)
    add_field('3. Solicitation or Project Number', proposal.solicitation_number or '')
    
    # Section B - Point of Contact
    add_heading('B. ARCHITECT-ENGINEER POINT OF CONTACT', 2)
    if proposal.firm:
        firm = proposal.firm
        poc_name = getattr(firm, 'point_of_contact_name', '') or ''
        poc_title = getattr(firm, 'point_of_contact_title', '') or ''
        add_field('4. Name and Title', f"{poc_name}, {poc_title}" if poc_title else poc_name)
        add_field('5. Name of Firm', firm.name or '')
        add_field('6. Telephone Number', getattr(firm, 'phone', '') or '')
        add_field('7. Fax Number', getattr(firm, 'fax', '') or '')
        add_field('8. E-mail Address', getattr(firm, 'email', '') or '')
    
    # Section C - Proposed Team
    add_heading('C. PROPOSED TEAM', 2)
    if firms_data:
        for idx, firm_info in enumerate(firms_data):
            firm = firm_info.get('firm')
            if firm:
                role = firm_info.get('role', '')
                loc = f"{firm.city}, {firm.state}" if firm.city else ""
                p = doc.add_paragraph()
                p.add_run(f"{chr(97+idx)}. ").bold = True
                p.add_run(f"{firm.name} - {loc}")
                if role:
                    p.add_run(f" ({role})")
    
    # Section D - Org Chart placeholder
    add_heading('D. ORGANIZATIONAL CHART OF PROPOSED TEAM', 2)
    doc.add_paragraph('(Attach organizational chart)')
    
    # Section E - Resumes
    add_heading('E. RESUMES OF KEY PERSONNEL PROPOSED FOR THIS CONTRACT', 2)
    
    for emp_data in employees_data:
        emp = emp_data.get('employee')
        if not emp:
            continue
        
        doc.add_paragraph()
        h = doc.add_heading(f"{emp.name}", 3)
        
        add_field('12. Name', emp.name)
        add_field('13. Role in This Contract', emp_data.get('role', '') or getattr(emp, 'role', ''))
        
        firm_name = ""
        if proposal.firm:
            firm_loc = f"{proposal.firm.city}, {proposal.firm.state}" if proposal.firm.city else ""
            firm_name = f"{proposal.firm.name}, {firm_loc}"
        add_field('14. Years Experience: a. Total', getattr(emp, 'years_experience_total', ''))
        add_field('14. Years Experience: b. With Current Firm', getattr(emp, 'years_experience_firm', ''))
        add_field('15. Firm Name and Location', firm_name)
        add_field('16. Education', getattr(emp, 'education', ''))
        add_field('17. Current Professional Registration', getattr(emp, 'registrations', ''))
        add_field('18. Other Professional Qualifications', getattr(emp, 'other_qualifications', ''))
        
        # Training
        if getattr(emp, 'training', ''):
            add_field('Training', emp.training)
        
        # Project Experience
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run('19. RELEVANT PROJECTS').bold = True
        
        experiences = emp_data.get('experiences', [])
        for exp_idx, exp in enumerate(experiences[:5]):
            exp_p = doc.add_paragraph()
            exp_p.add_run(f"({chr(97+exp_idx)}) ").bold = True
            title = getattr(exp, 'project_title', '') or getattr(exp, 'title', '') or ''
            location = getattr(exp, 'location', '') or ''
            year = getattr(exp, 'year_completed', '') or ''
            exp_p.add_run(f"{title}")
            if location:
                exp_p.add_run(f", {location}")
            if year:
                exp_p.add_run(f" ({year})")
            
            desc = getattr(exp, 'brief_description', '') or getattr(exp, 'description', '') or ''
            role = getattr(exp, 'role_performed', '') or getattr(exp, 'role', '') or ''
            if desc:
                doc.add_paragraph(desc)
            if role:
                add_field('Role', role)
        
        doc.add_page_break()
    
    # Section F - Example Projects
    add_heading('F. EXAMPLE PROJECTS WHICH BEST ILLUSTRATE PROPOSED TEAM\'S QUALIFICATIONS', 2)
    
    for idx, proj_data in enumerate(projects_data):
        proj = proj_data.get('project')
        if not proj:
            continue
        
        doc.add_paragraph()
        h = doc.add_heading(f"Project {idx + 1}: {proj.title}", 3)
        
        add_field('21. Title and Location', f"{proj.title}, {proj.location or ''}")
        year = getattr(proj, 'year_completed_professional', '') or ''
        add_field('22. Year Completed: Professional Services', year)
        year_const = getattr(proj, 'year_completed_construction', '') or ''
        if year_const:
            add_field('22. Year Completed: Construction', year_const)
        
        add_field('23a. Project Owner', getattr(proj, 'owner_name', '') or '')
        add_field('23b. Point of Contact Name', getattr(proj, 'owner_contact_name', '') or '')
        add_field('23c. Point of Contact Telephone', getattr(proj, 'owner_contact_phone', '') or '')
        
        # Get description (alternate or main)
        desc = proj_data.get('description') or getattr(proj, 'brief_description', '') or ''
        add_field('24. Brief Description', '')
        if desc:
            doc.add_paragraph(desc)
        
        # Firms involved
        if hasattr(proj, 'firm_involvements') and proj.firm_involvements:
            doc.add_paragraph()
            p = doc.add_paragraph()
            p.add_run('25. Firms Involved:').bold = True
            for fi in proj.firm_involvements:
                if fi.firm:
                    doc.add_paragraph(f"  - {fi.firm.name}: {fi.role or ''}")
        
        doc.add_page_break()
    
    # Section G - Key Personnel Participation Matrix
    add_heading('G. KEY PERSONNEL PARTICIPATION IN EXAMPLE PROJECTS', 2)
    
    if employees_data and projects_data:
        # Create table
        num_cols = min(len(projects_data), 10) + 2  # Name, Role, + projects
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Name'
        hdr_cells[1].text = 'Role'
        for idx, proj_data in enumerate(projects_data[:10]):
            proj = proj_data.get('project')
            if proj and idx + 2 < num_cols:
                hdr_cells[idx + 2].text = str(idx + 1)
        
        # Data rows
        for emp_data in employees_data:
            emp = emp_data.get('employee')
            if not emp:
                continue
            row = table.add_row().cells
            row[0].text = emp.name or ''
            row[1].text = emp_data.get('role', '') or ''
            
            # Check matrix for participation
            for proj_idx, proj_data in enumerate(projects_data[:10]):
                proj = proj_data.get('project')
                if proj and proj_idx + 2 < num_cols:
                    # Check if employee participated in this project
                    key = (emp.id, proj.id)
                    if matrix_data and key in matrix_data:
                        row[proj_idx + 2].text = 'X'
    
    doc.add_page_break()
    
    # Section H - Additional Information
    add_heading('H. ADDITIONAL INFORMATION', 2)
    written_sections = getattr(proposal, 'written_sections', '') or ''
    if written_sections:
        doc.add_paragraph(written_sections)
    else:
        doc.add_paragraph('(No additional information provided)')
    
    # Section I - Authorized Representative
    add_heading('I. AUTHORIZED REPRESENTATIVE', 2)
    if proposal.firm:
        poc_name = getattr(proposal.firm, 'point_of_contact_name', '') or ''
        poc_title = getattr(proposal.firm, 'point_of_contact_title', '') or ''
        add_field('Signature of Authorized Representative', '')
        add_field('Name', poc_name)
        add_field('Title', poc_title)
        from datetime import datetime
        add_field('Date', datetime.now().strftime('%m/%d/%Y'))
    
    doc.add_page_break()
    
    # Part II - General Qualifications
    add_heading('PART II - GENERAL QUALIFICATIONS', 1)
    
    if proposal.firm:
        firm = proposal.firm
        add_field('2a. Firm Name', firm.name or '')
        add_field('2b. Street Address', getattr(firm, 'street_address', '') or '')
        add_field('2c. City', getattr(firm, 'city', '') or '')
        add_field('2d. State', getattr(firm, 'state', '') or '')
        add_field('2e. Zip Code', getattr(firm, 'zip_code', '') or '')
        add_field('3. Year Established', getattr(firm, 'year_established', '') or '')
        add_field('4. UEI Number', getattr(firm, 'uei', '') or '')
        add_field('5a. Ownership Type', getattr(firm, 'ownership_type', '') or '')
        
        poc_name = getattr(firm, 'point_of_contact_name', '') or ''
        poc_title = getattr(firm, 'point_of_contact_title', '') or ''
        add_field('6a. Point of Contact Name and Title', f"{poc_name}, {poc_title}" if poc_title else poc_name)
        add_field('6b. Telephone Number', getattr(firm, 'phone', '') or '')
        add_field('6c. E-mail Address', getattr(firm, 'email', '') or '')
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
